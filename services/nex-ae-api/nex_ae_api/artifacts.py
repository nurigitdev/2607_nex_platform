from __future__ import annotations

import base64
import hashlib
import html
import json
import os
import re
import textwrap
from dataclasses import dataclass, field
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, Protocol
from uuid import NAMESPACE_URL, uuid5

import httpx
from docx import Document
from fastapi import FastAPI, Header, Request
from fastapi.responses import JSONResponse
from sqlalchemy import text
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session, sessionmaker

from nex_runtime import (
    DEFAULT_SERVICE_SCOPE,
    issue_mock_service_token,
    problem_response,
    request_id_from_headers,
    trace_id_from_headers,
    validate_authorization_header,
)


DEFAULT_TENANT_ID = "local-tenant"
DEFAULT_OWNER_USER_ID = "local-user"
DEFAULT_TARGET_FORMATS = ["MD", "HTML_PREVIEW"]
DEFAULT_RETENTION_POLICY_REF = "generated-artifact-retention-local-v1"
DEFAULT_ARTIFACT_TYPE = "generated_document"
MARKDOWN_RENDERER_POLICY_ID = "ae-markdown-renderer-v1"
MULTI_FORMAT_RENDER_STAGE_ORDER = (
    "HANDOFF_VALIDATING",
    "MARKDOWN_RENDERING",
    "HTML_PREVIEW_RENDERING",
    "DOCX_RENDERING",
    "PDF_RENDERING",
    "LINK_CREATING",
    "FINALIZING",
)
ARTIFACT_TRANSFORMER_CATALOG = {
    "MD": {
        "format": "MD",
        "mime_type": "text/markdown",
        "extension": "md",
        "render_stage": "MARKDOWN_RENDERING",
        "content_kind": "text",
        "materializer": "markdown_renderer",
        "implemented": True,
    },
    "HTML_PREVIEW": {
        "format": "HTML_PREVIEW",
        "mime_type": "text/html",
        "extension": "html",
        "render_stage": "HTML_PREVIEW_RENDERING",
        "content_kind": "text",
        "materializer": "html_preview_transformer",
        "implemented": True,
    },
    "DOCX": {
        "format": "DOCX",
        "mime_type": (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        "extension": "docx",
        "render_stage": "DOCX_RENDERING",
        "content_kind": "binary",
        "materializer": "docx_export_transformer",
        "implemented": True,
    },
    "PDF": {
        "format": "PDF",
        "mime_type": "application/pdf",
        "extension": "pdf",
        "render_stage": "PDF_RENDERING",
        "content_kind": "binary",
        "materializer": "pdf_export_transformer",
        "implemented": True,
    },
}
IMPLEMENTED_RENDER_FORMATS = {
    target_format
    for target_format, spec in ARTIFACT_TRANSFORMER_CATALOG.items()
    if spec["implemented"]
}
ARTIFACT_HANDOFF_JSON_FIELDS = (
    "actor_claims_ref",
    "workspace_ref",
    "target_formats",
    "quality_summary",
)
ARTIFACT_RECORD_JSON_FIELDS = (
    "owner_actor_ref",
    "workspace_ref",
    "target_formats",
    "template_ref",
    "handoff_ref",
)
ARTIFACT_SOURCE_REF_JSON_FIELDS = ("quality_summary",)
ARTIFACT_VERSION_JSON_FIELDS = ("rendered_formats", "validation_snapshot")
ARTIFACT_LINK_JSON_FIELDS = ("created_by_actor_ref",)
SUPPORTED_ARTIFACT_INTENTS = {
    "preview_only",
    "create_artifact",
    "create_and_export",
}
SUPPORTED_ARTIFACT_TYPES = {"generated_document", "summary", "answer_export"}
SUPPORTED_ARTIFACT_STATUSES = {
    "DRAFT",
    "RENDERING",
    "READY",
    "FAILED",
    "ARCHIVED",
    "DELETED",
}
AE_ARTIFACT_LIFECYCLE_ACTION_SCHEMA_VERSION = "ae_artifact_lifecycle_action.v1"
AE_ARTIFACT_LIFECYCLE_ACTION_RESULT_SCHEMA_VERSION = (
    "ae_artifact_lifecycle_action_result.v1"
)
AE_ARTIFACT_RETENTION_POLICY_SCHEMA_VERSION = "ae_artifact_retention_policy.v1"
AE_ARTIFACT_RETENTION_CANDIDATE_COLLECTION_SCHEMA_VERSION = (
    "ae_artifact_retention_candidate_collection.v1"
)
AE_ARTIFACT_RETENTION_CANDIDATE_ITEM_SCHEMA_VERSION = (
    "ae_artifact_retention_candidate_item.v1"
)
AE_ARTIFACT_RETENTION_EXECUTION_SCHEMA_VERSION = (
    "ae_artifact_retention_execution.v1"
)
AE_ARTIFACT_RETENTION_EXECUTION_HISTORY_SCHEMA_VERSION = (
    "ae_artifact_retention_execution_history.v1"
)
ARTIFACT_RETENTION_EXECUTION_HISTORY_JSON_FIELDS = (
    "deleted_counts",
    "requested_by",
    "error",
    "audit",
    "metadata",
    "execution",
)
SUPPORTED_ARTIFACT_LIFECYCLE_ACTIONS = {"ARCHIVE", "RESTORE", "MARK_DELETED"}
ARCHIVABLE_ARTIFACT_STATUSES = {"DRAFT", "READY", "FAILED"}
DELETABLE_ARTIFACT_STATUSES = {"DRAFT", "READY", "FAILED", "ARCHIVED"}
RESTORABLE_ARTIFACT_STATUSES = {"DRAFT", "READY", "FAILED"}
ARTIFACT_LIFECYCLE_DEFAULT_RESTORE_STATUS = "READY"
ARTIFACT_LIFECYCLE_DEFAULT_REASON_CODE = "user_requested"
DEFAULT_ARTIFACT_RETENTION_POLICY_ID = "ae-artifact-logical-purge-30d-local-v1"
DEFAULT_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE = 30
MIN_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE = 1
MAX_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE = 365
SUPPORTED_ARTIFACT_RETENTION_DAY_PRESETS = (15, 30)
ARTIFACT_RETENTION_LOGICAL_PURGE_STATUS = "DELETED"
ARTIFACT_RETENTION_LOGICAL_PURGE_FIELD = "artifact_status"
ARTIFACT_RETENTION_BATCH_TIMEZONE = "Asia/Seoul"
ARTIFACT_RETENTION_BATCH_WINDOW_START = "02:00"
ARTIFACT_RETENTION_BATCH_WINDOW_END = "05:00"
ARTIFACT_RETENTION_EXECUTION_MODES = ("DRY_RUN", "EXECUTE")
ARTIFACT_RETENTION_EXECUTION_STATUSES = (
    "PLANNED",
    "SUCCEEDED",
    "BLOCKED",
    "FAILED",
)
DEFAULT_ARTIFACT_RETENTION_MAX_DELETE_COUNT = 20
MAX_ARTIFACT_RETENTION_MAX_DELETE_COUNT = 100
SUPPORTED_TARGET_FORMATS = {"MD", "HTML_PREVIEW", "DOCX", "PDF"}
ARTIFACT_COLLECTION_SCHEMA_VERSION = "ae_artifact_collection.v1"
ARTIFACT_COLLECTION_ITEM_SCHEMA_VERSION = "ae_artifact_collection_item.v1"
DEFAULT_ARTIFACT_COLLECTION_LIMIT = 20
MAX_ARTIFACT_COLLECTION_LIMIT = 100


class CxArtifactSourceClient(Protocol):
    def get_generation(
        self,
        cx_generation_id: str,
        *,
        request_id: str,
        trace_id: str,
    ) -> dict[str, Any]:
        ...

    def get_structured_draft(
        self,
        cx_generation_id: str,
        *,
        request_id: str,
        trace_id: str,
    ) -> dict[str, Any]:
        ...


class RenderedArtifactStorage(Protocol):
    def save_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
        payload: bytes,
    ) -> str:
        ...

    def get_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bytes | None:
        ...

    def delete_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bool:
        ...

    def save_markdown(self, artifact_file: dict[str, Any], markdown: str) -> str:
        ...

    def get_markdown(self, artifact_file: dict[str, Any]) -> str | None:
        ...


@dataclass(frozen=True)
class HttpCxArtifactSourceClient:
    base_url: str = "http://127.0.0.1:8104"
    service_token: str | None = None
    timeout_seconds: float = 5.0

    def get_generation(
        self,
        cx_generation_id: str,
        *,
        request_id: str,
        trace_id: str,
    ) -> dict[str, Any]:
        return self._get_json(
            f"/api/v1/generations/{cx_generation_id}",
            request_id=request_id,
            trace_id=trace_id,
        )

    def get_structured_draft(
        self,
        cx_generation_id: str,
        *,
        request_id: str,
        trace_id: str,
    ) -> dict[str, Any]:
        return self._get_json(
            f"/api/v1/generations/{cx_generation_id}/structured-draft",
            request_id=request_id,
            trace_id=trace_id,
        )

    def _get_json(
        self,
        path: str,
        *,
        request_id: str,
        trace_id: str,
    ) -> dict[str, Any]:
        token = self.service_token or issue_mock_service_token(
            service_id="nex-ae-api",
            audience="nex-cx",
        ).access_token
        response = httpx.get(
            f"{self.base_url}{path}",
            headers={
                "Authorization": f"Bearer {token}",
                "X-Request-ID": request_id,
                "traceparent": f"00-{trace_id}-00f067aa0ba902b7-01",
                "X-Service-ID": "nex-ae-api",
            },
            timeout=self.timeout_seconds,
        )
        if response.status_code >= 400:
            body = _safe_response_json(response)
            raise ArtifactHandoffError(
                status_code=response.status_code,
                error_code=body.get("error_code", "cx.artifact_source_request_failed"),
                detail=body.get("detail", "CX artifact source request failed."),
                retryable=body.get("retryable", False),
            )
        return response.json()


@dataclass
class ArtifactHandoffStore:
    records: dict[str, dict[str, Any]] = field(default_factory=dict)

    def save(self, record: dict[str, Any]) -> dict[str, Any]:
        self.records[record["artifact_handoff_id"]] = record
        return record

    def get(self, artifact_handoff_id: str) -> dict[str, Any] | None:
        return self.records.get(artifact_handoff_id)


@dataclass
class ArtifactRecordStore:
    records: dict[str, dict[str, Any]] = field(default_factory=dict)
    render_jobs: dict[str, dict[str, Any]] = field(default_factory=dict)
    rendered_markdown: dict[str, str] = field(default_factory=dict)
    rendered_artifact_files: dict[str, bytes] = field(default_factory=dict)
    artifact_files: dict[str, dict[str, Any]] = field(default_factory=dict)
    artifact_links: dict[str, dict[str, Any]] = field(default_factory=dict)

    def create(self, record: dict[str, Any]) -> dict[str, Any]:
        existing = self.records.get(record["artifact_id"])
        if existing is not None:
            return existing
        self.records[record["artifact_id"]] = record
        return record

    def save(self, record: dict[str, Any]) -> dict[str, Any]:
        self.records[record["artifact_id"]] = record
        return record

    def get(self, artifact_id: str) -> dict[str, Any] | None:
        return self.records.get(artifact_id)

    def list_versions(self, artifact_id: str) -> list[dict[str, Any]] | None:
        record = self.get(artifact_id)
        if record is None:
            return None
        return list(record["versions"])

    def list_artifacts(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        status: str | None = None,
        limit: int | str | None = None,
    ) -> dict[str, Any]:
        collection_filter = build_artifact_collection_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            status=status,
            limit=limit,
        )
        records = [
            record
            for record in self.records.values()
            if artifact_record_matches_collection_filter(record, collection_filter)
        ]
        records.sort(
            key=lambda record: (
                str(record.get("updated_at") or ""),
                str(record.get("artifact_id") or ""),
            ),
            reverse=True,
        )
        return build_artifact_collection(
            records[: collection_filter["limit"]],
            collection_filter=collection_filter,
        )

    def list_retention_candidates(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        retention_days: int | str | None = None,
        as_of: str | None = None,
        limit: int | str | None = None,
    ) -> dict[str, Any]:
        candidate_filter = build_artifact_retention_candidate_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            retention_days=retention_days,
            as_of=as_of,
            limit=limit,
        )
        records = [
            record
            for record in self.records.values()
            if artifact_record_matches_retention_candidate_filter(
                record,
                candidate_filter,
            )
        ]
        records.sort(
            key=lambda record: (
                str(record.get("updated_at") or ""),
                str(record.get("artifact_id") or ""),
            )
        )
        return build_artifact_retention_candidate_collection(
            records[: candidate_filter["limit"]],
            candidate_filter=candidate_filter,
        )

    def purge_retention_candidates(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        retention_days: int | str | None = None,
        as_of: str | None = None,
        checked_at: str | None = None,
        scan_limit: int | str | None = None,
        max_delete_count: int | str | None = None,
        dry_run: bool = True,
        delete_enabled: bool = False,
        storage_mutation_enabled: bool = False,
        database_row_delete_enabled: bool = False,
        requested_by: dict[str, Any] | None = None,
        idempotency_key: str | None = None,
        trace_id: str | None = None,
        request_id: str | None = None,
    ) -> dict[str, Any]:
        candidate_filter = build_artifact_retention_candidate_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            retention_days=retention_days,
            as_of=as_of,
            limit=scan_limit,
        )
        records = [
            record
            for record in self.records.values()
            if artifact_record_matches_retention_candidate_filter(
                record,
                candidate_filter,
            )
        ]
        records.sort(
            key=lambda record: (
                str(record.get("updated_at") or ""),
                str(record.get("artifact_id") or ""),
            )
        )
        records = records[: candidate_filter["limit"]]
        max_delete = normalize_artifact_retention_delete_limit(max_delete_count)
        selected = records[:max_delete]
        if dry_run:
            return build_artifact_retention_execution(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                owner_user_id=owner_user_id,
                mode="DRY_RUN",
                execution_status="SUCCEEDED",
                retention_days=retention_days,
                as_of=candidate_filter["as_of"],
                checked_at=checked_at,
                scan_limit=candidate_filter["limit"],
                max_delete_count=max_delete,
                candidate_count=len(records),
                selected_count=len(selected),
                delete_enabled=delete_enabled,
                storage_mutation_enabled=storage_mutation_enabled,
                database_row_delete_enabled=database_row_delete_enabled,
                requested_by=requested_by,
                idempotency_key=idempotency_key,
                trace_id=trace_id,
                request_id=request_id,
            )
        if not (
            delete_enabled
            and storage_mutation_enabled
            and database_row_delete_enabled
        ):
            return build_artifact_retention_execution(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                owner_user_id=owner_user_id,
                mode="EXECUTE",
                execution_status="BLOCKED",
                retention_days=retention_days,
                as_of=candidate_filter["as_of"],
                checked_at=checked_at,
                scan_limit=candidate_filter["limit"],
                max_delete_count=max_delete,
                candidate_count=len(records),
                selected_count=0,
                delete_enabled=delete_enabled,
                storage_mutation_enabled=storage_mutation_enabled,
                database_row_delete_enabled=database_row_delete_enabled,
                requested_by=requested_by,
                idempotency_key=idempotency_key,
                trace_id=trace_id,
                request_id=request_id,
                blocked_reason="delete_not_enabled",
            )
        deleted_counts = self._delete_retention_selected_records(selected)
        return build_artifact_retention_execution(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            mode="EXECUTE",
            execution_status="SUCCEEDED",
            retention_days=retention_days,
            as_of=candidate_filter["as_of"],
            checked_at=checked_at,
            scan_limit=candidate_filter["limit"],
            max_delete_count=max_delete,
            candidate_count=len(records),
            selected_count=len(selected),
            deleted_counts=deleted_counts,
            delete_enabled=True,
            storage_mutation_enabled=True,
            database_row_delete_enabled=True,
            requested_by=requested_by,
            idempotency_key=idempotency_key,
            trace_id=trace_id,
            request_id=request_id,
        )

    def _delete_retention_selected_records(
        self,
        records: list[dict[str, Any]],
    ) -> dict[str, int]:
        counts = _empty_artifact_retention_deleted_counts()
        for record in records:
            artifact_id = record["artifact_id"]
            for artifact_link in _list_of_mappings(record.get("links")):
                if self.artifact_links.pop(artifact_link["artifact_link_id"], None):
                    counts["links"] += 1
            for artifact_file in _list_of_mappings(record.get("files")):
                if self.artifact_files.pop(artifact_file["artifact_file_id"], None):
                    counts["files"] += 1
                if self.rendered_artifact_files.pop(
                    artifact_file["artifact_file_id"],
                    None,
                ) is not None:
                    counts["storage_files"] += 1
                if artifact_file.get("format") == "MD":
                    self.rendered_markdown.pop(
                        artifact_file["artifact_version_id"],
                        None,
                    )
            for render_job in _list_of_mappings(record.get("render_jobs")):
                if self.render_jobs.pop(render_job["render_job_id"], None):
                    counts["render_jobs"] += 1
            counts["versions"] += len(_list_of_mappings(record.get("versions")))
            counts["source_refs"] += len(_list_of_mappings(record.get("source_refs")))
            if self.records.pop(artifact_id, None) is not None:
                counts["artifacts"] += 1
        return counts

    def get_render_job(self, render_job_id: str) -> dict[str, Any] | None:
        return self.render_jobs.get(render_job_id)

    def get_rendered_markdown(self, artifact_version_id: str) -> str | None:
        return self.rendered_markdown.get(artifact_version_id)

    def get_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bytes | None:
        payload = self.rendered_artifact_files.get(artifact_file["artifact_file_id"])
        if payload is not None:
            return payload
        markdown = self.get_rendered_markdown(artifact_file["artifact_version_id"])
        if artifact_file["format"] == "MD" and markdown is not None:
            return markdown.encode("utf-8")
        return None

    def get_file(self, artifact_file_id: str) -> dict[str, Any] | None:
        return self.artifact_files.get(artifact_file_id)

    def get_file_link(
        self,
        artifact_file_id: str,
        link_type: str,
    ) -> dict[str, Any] | None:
        for link in self.artifact_links.values():
            if (
                link["artifact_file_id"] == artifact_file_id
                and link["link_type"] == link_type
            ):
                return link
        return None

    def apply_markdown_render(
        self,
        *,
        artifact_id: str,
        artifact_version: dict[str, Any],
        render_job: dict[str, Any],
        markdown: str,
        artifact_files: list[dict[str, Any]],
        artifact_links: list[dict[str, Any]],
        rendered_payloads: dict[str, bytes] | None = None,
    ) -> dict[str, Any]:
        record = self.records[artifact_id]
        payloads = rendered_payloads or {"MD": markdown.encode("utf-8")}
        record["versions"].append(artifact_version)
        record["render_jobs"].append(render_job)
        record["files"].extend(artifact_files)
        record["links"].extend(artifact_links)
        record["artifact_status"] = "READY"
        record["current_version_id"] = artifact_version["artifact_version_id"]
        record["updated_at"] = render_job["completed_at"]
        self.render_jobs[render_job["render_job_id"]] = render_job
        self.rendered_markdown[artifact_version["artifact_version_id"]] = markdown
        for artifact_file in artifact_files:
            self.artifact_files[artifact_file["artifact_file_id"]] = artifact_file
            payload = rendered_payload_for_file(artifact_file, payloads)
            if payload is not None:
                self.rendered_artifact_files[artifact_file["artifact_file_id"]] = payload
        for artifact_link in artifact_links:
            self.artifact_links[artifact_link["artifact_link_id"]] = artifact_link
        return record

    def apply_lifecycle_action(
        self,
        artifact_id: str,
        action_request: dict[str, Any],
    ) -> dict[str, Any]:
        record = self.get(artifact_id)
        if record is None:
            raise ArtifactHandoffError(
                status_code=404,
                error_code="ae.artifact_not_found",
                detail=f"Artifact was not found: {artifact_id}",
            )
        updated_record = apply_artifact_lifecycle_action(
            artifact_record=record,
            action_request=action_request,
        )
        self.save(updated_record)
        return build_artifact_lifecycle_action_result(
            action_request=action_request,
            artifact_record=updated_record,
        )


@dataclass
class ArtifactRetentionExecutionHistoryStore:
    records: dict[str, dict[str, Any]] = field(default_factory=dict)

    def save(self, execution: dict[str, Any]) -> dict[str, Any]:
        record = build_artifact_retention_execution_history_record(execution)
        existing = self.get_by_idempotency_key(
            tenant_id=record["tenant_id"],
            workspace_id=record["workspace_id"],
            owner_user_id=record["owner_user_id"],
            idempotency_key=record["idempotency_key"],
        )
        if existing is not None:
            return existing
        existing_by_id = self.records.get(record["retention_execution_id"])
        if existing_by_id is not None:
            return existing_by_id
        self.records[record["retention_execution_id"]] = record
        return record

    def get(self, retention_execution_id: str) -> dict[str, Any] | None:
        return self.records.get(retention_execution_id)

    def get_by_idempotency_key(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        idempotency_key: str | None,
    ) -> dict[str, Any] | None:
        normalized_idempotency_key = optional_text(idempotency_key)
        if normalized_idempotency_key is None:
            return None
        normalized_tenant_id = _required_collection_scope_text(tenant_id, "tenant_id")
        normalized_workspace_id = _required_collection_scope_text(
            workspace_id,
            "workspace_id",
        )
        normalized_owner_user_id = _required_collection_scope_text(
            owner_user_id,
            "owner_user_id",
        )
        for record in self.records.values():
            if (
                record.get("tenant_id") == normalized_tenant_id
                and record.get("workspace_id") == normalized_workspace_id
                and record.get("owner_user_id") == normalized_owner_user_id
                and record.get("idempotency_key") == normalized_idempotency_key
            ):
                return record
        return None

    def list_executions(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        mode: str | None = None,
        execution_status: str | None = None,
        limit: int | str | None = None,
    ) -> list[dict[str, Any]]:
        normalized_filter = _artifact_retention_execution_history_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            mode=mode,
            execution_status=execution_status,
            limit=limit,
        )
        records = [
            record
            for record in self.records.values()
            if _artifact_retention_execution_history_matches(record, normalized_filter)
        ]
        records.sort(
            key=lambda record: (
                str(record.get("checked_at") or ""),
                str(record.get("retention_execution_id") or ""),
            ),
            reverse=True,
        )
        return records[: normalized_filter["limit"]]


@dataclass
class InMemoryRenderedArtifactStorage:
    rendered_markdown: dict[str, str] = field(default_factory=dict)
    rendered_artifact_files: dict[str, bytes] = field(default_factory=dict)

    def save_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
        payload: bytes,
    ) -> str:
        self.rendered_artifact_files[artifact_file["artifact_file_id"]] = payload
        if artifact_file["format"] == "MD":
            self.rendered_markdown[artifact_file["artifact_version_id"]] = (
                payload.decode("utf-8")
            )
        return artifact_file["storage_ref"]

    def get_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bytes | None:
        payload = self.rendered_artifact_files.get(artifact_file["artifact_file_id"])
        if payload is not None:
            return payload
        markdown = self.rendered_markdown.get(artifact_file["artifact_version_id"])
        if artifact_file["format"] == "MD" and markdown is not None:
            return markdown.encode("utf-8")
        return None

    def delete_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bool:
        artifact_file_id = artifact_file["artifact_file_id"]
        artifact_version_id = artifact_file["artifact_version_id"]
        removed = self.rendered_artifact_files.pop(artifact_file_id, None) is not None
        if artifact_file.get("format") == "MD":
            self.rendered_markdown.pop(artifact_version_id, None)
        return removed

    def save_markdown(self, artifact_file: dict[str, Any], markdown: str) -> str:
        return self.save_rendered_artifact_file(artifact_file, markdown.encode("utf-8"))

    def get_markdown(self, artifact_file: dict[str, Any]) -> str | None:
        payload = self.get_rendered_artifact_file(artifact_file)
        if payload is None:
            return None
        return payload.decode("utf-8")


@dataclass(frozen=True)
class LocalRenderedArtifactStorage:
    root: Path

    def save_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
        payload: bytes,
    ) -> str:
        path = self.path_for_storage_ref(artifact_file["storage_ref"])
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(payload)
        return artifact_file["storage_ref"]

    def get_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bytes | None:
        path = self.path_for_storage_ref(artifact_file["storage_ref"])
        if not path.exists():
            return None
        return path.read_bytes()

    def delete_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bool:
        path = self.path_for_storage_ref(artifact_file["storage_ref"])
        if not path.exists():
            return False
        path.unlink()
        return True

    def save_markdown(self, artifact_file: dict[str, Any], markdown: str) -> str:
        return self.save_rendered_artifact_file(
            artifact_file,
            markdown.encode("utf-8"),
        )

    def get_markdown(self, artifact_file: dict[str, Any]) -> str | None:
        payload = self.get_rendered_artifact_file(artifact_file)
        if payload is None:
            return None
        return payload.decode("utf-8")

    def path_for_storage_ref(self, storage_ref: str) -> Path:
        relative = _storage_ref_relative_path(storage_ref)
        root = self.root.resolve()
        path = root.joinpath(*relative.split("/")).resolve(strict=False)
        if not path.is_relative_to(root):
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_storage_ref_invalid",
                detail="Artifact storage ref escapes the configured storage root.",
            )
        return path


def build_default_rendered_artifact_storage(
    environ: dict[str, str] | None = None,
) -> RenderedArtifactStorage:
    env = environ if environ is not None else os.environ
    root = optional_text(env.get("NEX_AE_ARTIFACT_STORAGE_ROOT"))
    if root is None:
        return InMemoryRenderedArtifactStorage()
    return LocalRenderedArtifactStorage(Path(root))


class SqlAlchemyArtifactHandoffStore:
    def __init__(self, session_factory: sessionmaker[Session]) -> None:
        self._session_factory = session_factory

    def save(self, record: dict[str, Any]) -> dict[str, Any]:
        validate_artifact_handoff_record(record)
        try:
            with self._session_factory() as session:
                session.execute(
                    text(_artifact_handoff_upsert_sql(_dialect_name(session))),
                    _artifact_handoff_params(record),
                )
                session.commit()
            return record
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_handoff_store_unavailable",
                detail="AE artifact handoff store is unavailable.",
                retryable=True,
            ) from exc

    def get(self, artifact_handoff_id: str) -> dict[str, Any] | None:
        try:
            with self._session_factory() as session:
                row = (
                    session.execute(
                        text(_artifact_handoff_select_sql()),
                        {"artifact_handoff_id": artifact_handoff_id},
                    )
                    .mappings()
                    .first()
                )
            return _artifact_handoff_from_row(row) if row is not None else None
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_handoff_store_unavailable",
                detail="AE artifact handoff store is unavailable.",
                retryable=True,
            ) from exc

    def delete(self, artifact_handoff_id: str) -> int:
        try:
            with self._session_factory() as session:
                result = session.execute(
                    text(
                        """
                        DELETE FROM ae_artifact_handoffs
                        WHERE artifact_handoff_id = :artifact_handoff_id
                        """
                    ),
                    {"artifact_handoff_id": artifact_handoff_id},
                )
                session.commit()
                return int(result.rowcount or 0)
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_handoff_store_unavailable",
                detail="AE artifact handoff store is unavailable.",
                retryable=True,
            ) from exc


class SqlAlchemyArtifactRecordStore:
    def __init__(
        self,
        session_factory: sessionmaker[Session],
        *,
        rendered_storage: RenderedArtifactStorage | None = None,
    ) -> None:
        self._session_factory = session_factory
        self._rendered_storage = rendered_storage or InMemoryRenderedArtifactStorage()
        self.rendered_markdown = getattr(self._rendered_storage, "rendered_markdown", {})
        self.rendered_artifact_files = getattr(
            self._rendered_storage,
            "rendered_artifact_files",
            {},
        )

    def create(self, record: dict[str, Any]) -> dict[str, Any]:
        existing = self.get(record["artifact_id"])
        if existing is not None:
            return existing
        return self.save(record)

    def save(self, record: dict[str, Any]) -> dict[str, Any]:
        try:
            with self._session_factory() as session:
                _persist_artifact_record(session, record)
                session.commit()
            return record
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def get(self, artifact_id: str) -> dict[str, Any] | None:
        try:
            with self._session_factory() as session:
                return _load_artifact_record(session, artifact_id)
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def list_versions(self, artifact_id: str) -> list[dict[str, Any]] | None:
        record = self.get(artifact_id)
        if record is None:
            return None
        return list(record["versions"])

    def list_artifacts(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        status: str | None = None,
        limit: int | str | None = None,
    ) -> dict[str, Any]:
        collection_filter = build_artifact_collection_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            status=status,
            limit=limit,
        )
        try:
            with self._session_factory() as session:
                records = _load_artifact_records_for_collection(
                    session,
                    collection_filter,
                )
            return build_artifact_collection(
                records,
                collection_filter=collection_filter,
            )
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def list_retention_candidates(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        retention_days: int | str | None = None,
        as_of: str | None = None,
        limit: int | str | None = None,
    ) -> dict[str, Any]:
        candidate_filter = build_artifact_retention_candidate_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            retention_days=retention_days,
            as_of=as_of,
            limit=limit,
        )
        try:
            with self._session_factory() as session:
                records = _load_artifact_records_for_retention_candidates(
                    session,
                    candidate_filter,
                )
            return build_artifact_retention_candidate_collection(
                records,
                candidate_filter=candidate_filter,
            )
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def purge_retention_candidates(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        retention_days: int | str | None = None,
        as_of: str | None = None,
        checked_at: str | None = None,
        scan_limit: int | str | None = None,
        max_delete_count: int | str | None = None,
        dry_run: bool = True,
        delete_enabled: bool = False,
        storage_mutation_enabled: bool = False,
        database_row_delete_enabled: bool = False,
        requested_by: dict[str, Any] | None = None,
        idempotency_key: str | None = None,
        trace_id: str | None = None,
        request_id: str | None = None,
    ) -> dict[str, Any]:
        candidate_filter = build_artifact_retention_candidate_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            retention_days=retention_days,
            as_of=as_of,
            limit=scan_limit,
        )
        max_delete = normalize_artifact_retention_delete_limit(max_delete_count)
        try:
            with self._session_factory() as session:
                records = _load_artifact_records_for_retention_candidates(
                    session,
                    candidate_filter,
                )
                selected = records[:max_delete]
                if dry_run:
                    return build_artifact_retention_execution(
                        tenant_id=tenant_id,
                        workspace_id=workspace_id,
                        owner_user_id=owner_user_id,
                        mode="DRY_RUN",
                        execution_status="SUCCEEDED",
                        retention_days=retention_days,
                        as_of=candidate_filter["as_of"],
                        checked_at=checked_at,
                        scan_limit=candidate_filter["limit"],
                        max_delete_count=max_delete,
                        candidate_count=len(records),
                        selected_count=len(selected),
                        delete_enabled=delete_enabled,
                        storage_mutation_enabled=storage_mutation_enabled,
                        database_row_delete_enabled=database_row_delete_enabled,
                        requested_by=requested_by,
                        idempotency_key=idempotency_key,
                        trace_id=trace_id,
                        request_id=request_id,
                    )
                if not (
                    delete_enabled
                    and storage_mutation_enabled
                    and database_row_delete_enabled
                ):
                    return build_artifact_retention_execution(
                        tenant_id=tenant_id,
                        workspace_id=workspace_id,
                        owner_user_id=owner_user_id,
                        mode="EXECUTE",
                        execution_status="BLOCKED",
                        retention_days=retention_days,
                        as_of=candidate_filter["as_of"],
                        checked_at=checked_at,
                        scan_limit=candidate_filter["limit"],
                        max_delete_count=max_delete,
                        candidate_count=len(records),
                        selected_count=0,
                        delete_enabled=delete_enabled,
                        storage_mutation_enabled=storage_mutation_enabled,
                        database_row_delete_enabled=database_row_delete_enabled,
                        requested_by=requested_by,
                        idempotency_key=idempotency_key,
                        trace_id=trace_id,
                        request_id=request_id,
                        blocked_reason="delete_not_enabled",
                    )
                deleted_counts = _delete_retention_selected_artifact_records(
                    session,
                    selected,
                    rendered_storage=self._rendered_storage,
                )
                execution = build_artifact_retention_execution(
                    tenant_id=tenant_id,
                    workspace_id=workspace_id,
                    owner_user_id=owner_user_id,
                    mode="EXECUTE",
                    execution_status="SUCCEEDED",
                    retention_days=retention_days,
                    as_of=candidate_filter["as_of"],
                    checked_at=checked_at,
                    scan_limit=candidate_filter["limit"],
                    max_delete_count=max_delete,
                    candidate_count=len(records),
                    selected_count=len(selected),
                    deleted_counts=deleted_counts,
                    delete_enabled=True,
                    storage_mutation_enabled=True,
                    database_row_delete_enabled=True,
                    requested_by=requested_by,
                    idempotency_key=idempotency_key,
                    trace_id=trace_id,
                    request_id=request_id,
                )
                session.commit()
                return execution
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def get_render_job(self, render_job_id: str) -> dict[str, Any] | None:
        try:
            with self._session_factory() as session:
                row = (
                    session.execute(
                        text(_artifact_render_job_select_sql("render_job_id = :render_job_id")),
                        {"render_job_id": render_job_id},
                    )
                    .mappings()
                    .first()
                )
            return _artifact_render_job_from_row(row) if row is not None else None
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def get_rendered_markdown(self, artifact_version_id: str) -> str | None:
        artifact_file = self._get_markdown_file_for_version(artifact_version_id)
        if artifact_file is None:
            return None
        return self._rendered_storage.get_markdown(artifact_file)

    def get_rendered_artifact_file(
        self,
        artifact_file: dict[str, Any],
    ) -> bytes | None:
        return self._rendered_storage.get_rendered_artifact_file(artifact_file)

    def get_file(self, artifact_file_id: str) -> dict[str, Any] | None:
        try:
            with self._session_factory() as session:
                row = (
                    session.execute(
                        text(_artifact_file_select_sql("artifact_file_id = :artifact_file_id")),
                        {"artifact_file_id": artifact_file_id},
                    )
                    .mappings()
                    .first()
                )
            return _artifact_file_from_row(row) if row is not None else None
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def get_file_link(
        self,
        artifact_file_id: str,
        link_type: str,
    ) -> dict[str, Any] | None:
        try:
            with self._session_factory() as session:
                row = (
                    session.execute(
                        text(
                            _artifact_link_select_sql(
                                "artifact_file_id = :artifact_file_id "
                                "AND link_type = :link_type"
                            )
                        ),
                        {
                            "artifact_file_id": artifact_file_id,
                            "link_type": link_type,
                        },
                    )
                    .mappings()
                    .first()
                )
            return _artifact_link_from_row(row) if row is not None else None
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def apply_markdown_render(
        self,
        *,
        artifact_id: str,
        artifact_version: dict[str, Any],
        render_job: dict[str, Any],
        markdown: str,
        artifact_files: list[dict[str, Any]],
        artifact_links: list[dict[str, Any]],
        rendered_payloads: dict[str, bytes] | None = None,
    ) -> dict[str, Any]:
        record = self.get(artifact_id)
        if record is None:
            raise ArtifactHandoffError(
                status_code=404,
                error_code="ae.artifact_not_found",
                detail=f"Artifact was not found: {artifact_id}",
            )
        record["versions"].append(artifact_version)
        record["render_jobs"].append(render_job)
        record["files"].extend(artifact_files)
        record["links"].extend(artifact_links)
        record["artifact_status"] = "READY"
        record["current_version_id"] = artifact_version["artifact_version_id"]
        record["updated_at"] = render_job["completed_at"]
        payloads = rendered_payloads or {"MD": markdown.encode("utf-8")}
        for artifact_file in artifact_files:
            payload = rendered_payload_for_file(artifact_file, payloads)
            if payload is not None:
                self._rendered_storage.save_rendered_artifact_file(
                    artifact_file,
                    payload,
                )
        return self.save(record)

    def apply_lifecycle_action(
        self,
        artifact_id: str,
        action_request: dict[str, Any],
    ) -> dict[str, Any]:
        record = self.get(artifact_id)
        if record is None:
            raise ArtifactHandoffError(
                status_code=404,
                error_code="ae.artifact_not_found",
                detail=f"Artifact was not found: {artifact_id}",
            )
        updated_record = apply_artifact_lifecycle_action(
            artifact_record=record,
            action_request=action_request,
        )
        self.save(updated_record)
        return build_artifact_lifecycle_action_result(
            action_request=action_request,
            artifact_record=updated_record,
        )

    def delete(self, artifact_id: str) -> int:
        try:
            with self._session_factory() as session:
                result = session.execute(
                    text(
                        """
                        DELETE FROM ae_artifacts
                        WHERE artifact_id = :artifact_id
                        """
                    ),
                    {"artifact_id": artifact_id},
                )
                session.commit()
                return int(result.rowcount or 0)
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc

    def _get_markdown_file_for_version(
        self,
        artifact_version_id: str,
    ) -> dict[str, Any] | None:
        try:
            with self._session_factory() as session:
                row = (
                    session.execute(
                        text(
                            _artifact_file_select_sql(
                                "artifact_version_id = :artifact_version_id "
                                "AND format = 'MD'"
                            )
                        ),
                        {"artifact_version_id": artifact_version_id},
                    )
                    .mappings()
                    .first()
                )
            return _artifact_file_from_row(row) if row is not None else None
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_store_unavailable",
                detail="AE artifact store is unavailable.",
                retryable=True,
            ) from exc


class SqlAlchemyArtifactRetentionExecutionHistoryStore:
    def __init__(self, session_factory: sessionmaker[Session]) -> None:
        self._session_factory = session_factory

    def save(self, execution: dict[str, Any]) -> dict[str, Any]:
        record = build_artifact_retention_execution_history_record(execution)
        existing = self.get_by_idempotency_key(
            tenant_id=record["tenant_id"],
            workspace_id=record["workspace_id"],
            owner_user_id=record["owner_user_id"],
            idempotency_key=record["idempotency_key"],
        )
        if existing is not None:
            return existing
        try:
            with self._session_factory() as session:
                session.execute(
                    text(
                        _artifact_retention_execution_history_upsert_sql(
                            _dialect_name(session)
                        )
                    ),
                    _artifact_retention_execution_history_params(record),
                )
                session.commit()
            return record
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_retention_history_store_unavailable",
                detail="AE artifact retention execution history store is unavailable.",
                retryable=True,
            ) from exc

    def get(self, retention_execution_id: str) -> dict[str, Any] | None:
        try:
            with self._session_factory() as session:
                row = (
                    session.execute(
                        text(
                            _artifact_retention_execution_history_select_sql(
                                "retention_execution_id = :retention_execution_id"
                            )
                        ),
                        {"retention_execution_id": retention_execution_id},
                    )
                    .mappings()
                    .first()
                )
            return (
                _artifact_retention_execution_history_from_row(row)
                if row is not None
                else None
            )
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_retention_history_store_unavailable",
                detail="AE artifact retention execution history store is unavailable.",
                retryable=True,
            ) from exc

    def get_by_idempotency_key(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        idempotency_key: str | None,
    ) -> dict[str, Any] | None:
        normalized_idempotency_key = optional_text(idempotency_key)
        if normalized_idempotency_key is None:
            return None
        params = {
            "tenant_id": _required_collection_scope_text(tenant_id, "tenant_id"),
            "workspace_id": _required_collection_scope_text(
                workspace_id,
                "workspace_id",
            ),
            "owner_user_id": _required_collection_scope_text(
                owner_user_id,
                "owner_user_id",
            ),
            "idempotency_key": normalized_idempotency_key,
        }
        try:
            with self._session_factory() as session:
                row = (
                    session.execute(
                        text(
                            _artifact_retention_execution_history_select_sql(
                                "tenant_id = :tenant_id "
                                "AND workspace_id = :workspace_id "
                                "AND owner_user_id = :owner_user_id "
                                "AND idempotency_key = :idempotency_key"
                            )
                        ),
                        params,
                    )
                    .mappings()
                    .first()
                )
            return (
                _artifact_retention_execution_history_from_row(row)
                if row is not None
                else None
            )
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_retention_history_store_unavailable",
                detail="AE artifact retention execution history store is unavailable.",
                retryable=True,
            ) from exc

    def list_executions(
        self,
        *,
        tenant_id: str,
        workspace_id: str,
        owner_user_id: str,
        mode: str | None = None,
        execution_status: str | None = None,
        limit: int | str | None = None,
    ) -> list[dict[str, Any]]:
        normalized_filter = _artifact_retention_execution_history_filter(
            tenant_id=tenant_id,
            workspace_id=workspace_id,
            owner_user_id=owner_user_id,
            mode=mode,
            execution_status=execution_status,
            limit=limit,
        )
        try:
            with self._session_factory() as session:
                rows = (
                    session.execute(
                        text(
                            _artifact_retention_execution_history_list_sql(
                                normalized_filter
                            )
                        ),
                        _artifact_retention_execution_history_filter_params(
                            normalized_filter
                        ),
                    )
                    .mappings()
                    .all()
                )
            return [
                _artifact_retention_execution_history_from_row(row) for row in rows
            ]
        except SQLAlchemyError as exc:
            raise ArtifactHandoffError(
                status_code=503,
                error_code="ae.artifact_retention_history_store_unavailable",
                detail="AE artifact retention execution history store is unavailable.",
                retryable=True,
            ) from exc


@dataclass(frozen=True)
class ArtifactHandoffError(Exception):
    status_code: int
    error_code: str
    detail: str
    retryable: bool = False


DEFAULT_ARTIFACT_HANDOFF_STORE = ArtifactHandoffStore()
DEFAULT_ARTIFACT_RECORD_STORE = ArtifactRecordStore()
DEFAULT_ARTIFACT_RETENTION_HISTORY_STORE = ArtifactRetentionExecutionHistoryStore()


def build_default_artifact_handoff_store(app: Any) -> Any:
    persistence = getattr(app.state, "nex_persistence", None)
    session_factory = getattr(persistence, "api_session_factory", None)
    if session_factory is not None:
        return SqlAlchemyArtifactHandoffStore(session_factory)
    return DEFAULT_ARTIFACT_HANDOFF_STORE


def build_default_artifact_record_store(app: Any) -> Any:
    persistence = getattr(app.state, "nex_persistence", None)
    session_factory = getattr(persistence, "api_session_factory", None)
    if session_factory is not None:
        return SqlAlchemyArtifactRecordStore(
            session_factory,
            rendered_storage=build_default_rendered_artifact_storage(),
        )
    return DEFAULT_ARTIFACT_RECORD_STORE


def build_default_artifact_retention_execution_history_store(app: Any) -> Any:
    persistence = getattr(app.state, "nex_persistence", None)
    session_factory = getattr(persistence, "api_session_factory", None)
    if session_factory is not None:
        return SqlAlchemyArtifactRetentionExecutionHistoryStore(session_factory)
    return DEFAULT_ARTIFACT_RETENTION_HISTORY_STORE


def build_default_cx_artifact_source_client() -> HttpCxArtifactSourceClient:
    return HttpCxArtifactSourceClient(
        base_url=os.getenv("NEX_CX_BASE_URL", "http://127.0.0.1:8104"),
        service_token=os.getenv("NEX_AE_TO_CX_SERVICE_TOKEN"),
    )


def register_artifact_handoff_routes(
    app: FastAPI,
    *,
    store: Any | None = None,
    artifact_store: Any | None = None,
    cx_client: CxArtifactSourceClient | None = None,
) -> None:
    handoff_store = store or build_default_artifact_handoff_store(app)
    artifact_record_store = artifact_store or build_default_artifact_record_store(app)
    client = cx_client or build_default_cx_artifact_source_client()

    @app.post("/api/v1/artifact-handoffs", response_model=None)
    def create_artifact_handoff(
        payload: dict[str, Any],
        request: Request,
        authorization: str | None = Header(default=None),
        idempotency_key: str | None = Header(default=None, alias="Idempotency-Key"),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        request_id = request_id_from_headers(request)
        trace_id = payload.get("trace_id") or trace_id_from_headers(request)
        try:
            cx_generation_id = required_string(
                payload,
                "cx_generation_id",
                "ae.cx_generation_id_required",
            )
            generation_record = client.get_generation(
                cx_generation_id,
                request_id=request_id,
                trace_id=trace_id,
            )
            structured_draft = client.get_structured_draft(
                cx_generation_id,
                request_id=request_id,
                trace_id=trace_id,
            )
            return handoff_store.save(
                build_artifact_handoff_record(
                    source_payload=payload,
                    generation_record=generation_record,
                    structured_draft=structured_draft,
                    artifact_request_id=idempotency_key
                    or optional_text(payload.get("artifact_request_id")),
                    request_id=request_id,
                    trace_id=trace_id,
                )
            )
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.get("/api/v1/artifact-handoffs/{artifact_handoff_id}", response_model=None)
    def get_artifact_handoff(
        artifact_handoff_id: str,
        request: Request,
        authorization: str | None = Header(default=None),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        record = handoff_store.get(artifact_handoff_id)
        if record is None:
            return _artifact_problem_response(
                request,
                ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.artifact_handoff_not_found",
                    detail=f"Artifact handoff was not found: {artifact_handoff_id}",
                ),
        )
        return record

    @app.post("/api/v1/artifacts", response_model=None)
    def create_artifact(
        payload: dict[str, Any],
        request: Request,
        authorization: str | None = Header(default=None),
        idempotency_key: str | None = Header(default=None, alias="Idempotency-Key"),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            artifact_handoff_id = required_string(
                payload,
                "artifact_handoff_id",
                "ae.artifact_handoff_id_required",
            )
            handoff_record = handoff_store.get(artifact_handoff_id)
            if handoff_record is None:
                raise ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.artifact_handoff_not_found",
                    detail=f"Artifact handoff was not found: {artifact_handoff_id}",
                )
            record = build_artifact_record_from_handoff(
                source_payload=payload,
                handoff_record=handoff_record,
                artifact_request_id=idempotency_key
                or optional_text(payload.get("artifact_request_id")),
                request_id=request_id_from_headers(request),
                trace_id=payload.get("trace_id") or trace_id_from_headers(request),
            )
            return artifact_record_store.create(record)
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.get("/api/v1/artifacts", response_model=None)
    def list_artifacts(
        request: Request,
        authorization: str | None = Header(default=None),
        tenant_id: str | None = None,
        workspace_id: str | None = None,
        owner_user_id: str | None = None,
        status: str | None = None,
        limit: str | None = None,
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            return artifact_record_store.list_artifacts(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                owner_user_id=owner_user_id,
                status=status,
                limit=limit,
            )
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.get("/api/v1/artifact-retention/candidates", response_model=None)
    def list_artifact_retention_candidates(
        request: Request,
        authorization: str | None = Header(default=None),
        tenant_id: str | None = None,
        workspace_id: str | None = None,
        owner_user_id: str | None = None,
        retention_days: str | None = None,
        as_of: str | None = None,
        limit: str | None = None,
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            return artifact_record_store.list_retention_candidates(
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                owner_user_id=owner_user_id,
                retention_days=retention_days,
                as_of=as_of,
                limit=limit,
            )
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.post("/api/v1/artifact-retention/purge", response_model=None)
    def purge_artifact_retention_candidates(
        payload: dict[str, Any],
        request: Request,
        authorization: str | None = Header(default=None),
        idempotency_key: str | None = Header(default=None, alias="Idempotency-Key"),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            return artifact_record_store.purge_retention_candidates(
                tenant_id=payload.get("tenant_id"),
                workspace_id=payload.get("workspace_id"),
                owner_user_id=payload.get("owner_user_id"),
                retention_days=payload.get("retention_days"),
                as_of=payload.get("as_of"),
                checked_at=payload.get("checked_at"),
                scan_limit=payload.get("scan_limit"),
                max_delete_count=payload.get("max_delete_count"),
                dry_run=_boolean_from_payload(payload, "dry_run", default=True),
                delete_enabled=_boolean_from_payload(
                    payload,
                    "delete_enabled",
                    default=False,
                ),
                storage_mutation_enabled=_boolean_from_payload(
                    payload,
                    "storage_mutation_enabled",
                    default=False,
                ),
                database_row_delete_enabled=_boolean_from_payload(
                    payload,
                    "database_row_delete_enabled",
                    default=False,
                ),
                requested_by=payload.get("requested_by"),
                idempotency_key=idempotency_key
                or optional_text(payload.get("idempotency_key")),
                trace_id=payload.get("trace_id") or trace_id_from_headers(request),
                request_id=request_id_from_headers(request),
            )
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.get("/api/v1/artifacts/{artifact_id}", response_model=None)
    def get_artifact(
        artifact_id: str,
        request: Request,
        authorization: str | None = Header(default=None),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        record = artifact_record_store.get(artifact_id)
        if record is None:
            return _artifact_problem_response(
                request,
                ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.artifact_not_found",
                    detail=f"Artifact was not found: {artifact_id}",
                ),
            )
        return record

    @app.get("/api/v1/artifacts/{artifact_id}/versions", response_model=None)
    def list_artifact_versions(
        artifact_id: str,
        request: Request,
        authorization: str | None = Header(default=None),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        record = artifact_record_store.get(artifact_id)
        if record is None:
            return _artifact_problem_response(
                request,
                ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.artifact_not_found",
                    detail=f"Artifact was not found: {artifact_id}",
                ),
            )
        return {
            "artifact_id": record["artifact_id"],
            "current_version_id": record["current_version_id"],
            "versions": artifact_record_store.list_versions(artifact_id) or [],
        }

    @app.post("/api/v1/artifacts/{artifact_id}/lifecycle-actions", response_model=None)
    def create_artifact_lifecycle_action(
        artifact_id: str,
        payload: dict[str, Any],
        request: Request,
        authorization: str | None = Header(default=None),
        idempotency_key: str | None = Header(default=None, alias="Idempotency-Key"),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            record = artifact_record_store.get(artifact_id)
            if record is None:
                raise ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.artifact_not_found",
                    detail=f"Artifact was not found: {artifact_id}",
                )
            action_request = build_artifact_lifecycle_action_request(
                payload=payload,
                artifact_record=record,
                request_id=request_id_from_headers(request),
                trace_id=payload.get("trace_id") or trace_id_from_headers(request),
                idempotency_key=idempotency_key,
            )
            return artifact_record_store.apply_lifecycle_action(
                artifact_id,
                action_request,
            )
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.post("/api/v1/artifacts/{artifact_id}/render-jobs", response_model=None)
    def create_artifact_render_job(
        artifact_id: str,
        payload: dict[str, Any],
        request: Request,
        authorization: str | None = Header(default=None),
        idempotency_key: str | None = Header(default=None, alias="Idempotency-Key"),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            record = artifact_record_store.get(artifact_id)
            if record is None:
                raise ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.artifact_not_found",
                    detail=f"Artifact was not found: {artifact_id}",
                )
            render_request_id = idempotency_key or required_string(
                payload,
                "render_request_id",
                "ae.render_request_id_required",
            )
            render_job_id = deterministic_render_job_id(artifact_id, render_request_id)
            existing_job = artifact_record_store.get_render_job(render_job_id)
            if existing_job is not None:
                return {
                    "render_result_schema_version": "ae_markdown_render_result.v1",
                    "render_job": existing_job,
                    "artifact": record,
                }
            target_formats = render_target_formats_from_payload(payload, record)
            source_ref = record["source_refs"][0]
            structured_draft = client.get_structured_draft(
                source_ref["cx_generation_id"],
                request_id=request_id_from_headers(request),
                trace_id=payload.get("trace_id") or trace_id_from_headers(request),
            )
            render_result = build_markdown_render_result(
                artifact_record=record,
                structured_draft=structured_draft,
                target_formats=target_formats,
                render_request_id=render_request_id,
                render_job_id=render_job_id,
            )
            updated_record = artifact_record_store.apply_markdown_render(
                artifact_id=artifact_id,
                artifact_version=render_result["artifact_version"],
                render_job=render_result["render_job"],
                markdown=render_result["markdown"],
                artifact_files=render_result["artifact_files"],
                artifact_links=render_result["artifact_links"],
                rendered_payloads=render_result["rendered_payloads"],
            )
            return {
                "render_result_schema_version": "ae_markdown_render_result.v1",
                "render_job": render_result["render_job"],
                "artifact": updated_record,
            }
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.get("/api/v1/artifact-render-jobs/{render_job_id}", response_model=None)
    def get_artifact_render_job(
        render_job_id: str,
        request: Request,
        authorization: str | None = Header(default=None),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        render_job = artifact_record_store.get_render_job(render_job_id)
        if render_job is None:
            return _artifact_problem_response(
                request,
                ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.render_job_not_found",
                    detail=f"Artifact render job was not found: {render_job_id}",
                ),
            )
        return render_job

    @app.get("/api/v1/artifact-files/{artifact_file_id}", response_model=None)
    def get_artifact_file(
        artifact_file_id: str,
        request: Request,
        authorization: str | None = Header(default=None),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        artifact_file = artifact_record_store.get_file(artifact_file_id)
        if artifact_file is None:
            return _artifact_problem_response(
                request,
                ArtifactHandoffError(
                    status_code=404,
                    error_code="ae.artifact_file_not_found",
                    detail=f"Artifact file was not found: {artifact_file_id}",
                ),
            )
        return artifact_file

    @app.get("/api/v1/artifact-files/{artifact_file_id}/preview", response_model=None)
    def preview_artifact_file(
        artifact_file_id: str,
        request: Request,
        authorization: str | None = Header(default=None),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            artifact_file, preview_link, payload = resolve_rendered_artifact_file_payload(
                artifact_record_store,
                artifact_file_id=artifact_file_id,
                link_type="preview",
            )
            rendered_text = rendered_text_from_payload(artifact_file, payload)
            preview_text = rendered_text[:2000]
            return {
                "preview_schema_version": "ae_artifact_file_preview.v1",
                "artifact_file": artifact_file,
                "artifact_link": preview_link,
                "content_type": artifact_file["mime_type"],
                "text_preview": preview_text,
                "truncated": len(rendered_text) > len(preview_text),
            }
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)

    @app.get("/api/v1/artifact-files/{artifact_file_id}/download", response_model=None)
    def download_artifact_file(
        artifact_file_id: str,
        request: Request,
        authorization: str | None = Header(default=None),
    ):
        auth_problem = _authorize_ae_request(request, authorization)
        if auth_problem is not None:
            return auth_problem

        try:
            artifact_file, download_link, payload = resolve_rendered_artifact_file_payload(
                artifact_record_store,
                artifact_file_id=artifact_file_id,
                link_type="download",
            )
            return {
                "download_schema_version": "ae_artifact_file_download.v1",
                "artifact_file": artifact_file,
                "artifact_link": download_link,
                "download_file_name": artifact_file["file_name"],
                "content_type": artifact_file["mime_type"],
                "content_hash": artifact_file["file_hash"],
                **rendered_download_fields_from_payload(artifact_file, payload),
            }
        except ArtifactHandoffError as exc:
            return _artifact_problem_response(request, exc)


def build_artifact_handoff_record(
    *,
    source_payload: dict[str, Any],
    generation_record: dict[str, Any],
    structured_draft: dict[str, Any],
    artifact_request_id: str | None,
    request_id: str,
    trace_id: str,
) -> dict[str, Any]:
    validate_source_generation(generation_record, structured_draft)
    chat_document_id = required_string(
        source_payload,
        "chat_document_id",
        "ae.chat_document_id_required",
    )
    interaction_id = required_string(
        source_payload,
        "interaction_id",
        "ae.interaction_id_required",
    )
    normalized_artifact_request_id = artifact_request_id or str(
        uuid5(
            NAMESPACE_URL,
            (
                "ae-artifact-request:"
                f"{chat_document_id}:{interaction_id}:{generation_record['cx_generation_id']}"
            ),
        )
    )
    artifact_handoff_id = source_payload.get("artifact_handoff_id") or str(
        uuid5(NAMESPACE_URL, f"ae-artifact-handoff:{normalized_artifact_request_id}")
    )
    now = _utc_now()
    citation_claims_hash = sha256_json({"citations": structured_draft["citations"]})
    validation_result_hash = sha256_json(structured_draft["validation"])
    request_metadata = generation_record.get("request_metadata", {})
    return {
        "handoff_schema_version": "ae_artifact_handoff.v1",
        "artifact_handoff_id": artifact_handoff_id,
        "artifact_request_id": normalized_artifact_request_id,
        "handoff_status": "READY_FOR_RENDERING",
        "trace_id": trace_id,
        "request_id": request_id,
        "chat_document_id": chat_document_id,
        "interaction_id": interaction_id,
        "actor_claims_ref": actor_claims_ref_from_payload(source_payload),
        "workspace_ref": workspace_ref_from_payload(source_payload),
        "cx_generation_id": generation_record["cx_generation_id"],
        "structured_draft_id": structured_draft["structured_draft_id"],
        "draft_schema_version": structured_draft["structured_draft_schema_version"],
        "structured_draft_content_hash": structured_draft["content_hash"],
        "citation_claims_hash": citation_claims_hash,
        "validation_result_hash": validation_result_hash,
        "template_id": optional_text(source_payload.get("template_id")),
        "template_version": optional_text(source_payload.get("template_version")),
        "rendering_template_id": optional_text(source_payload.get("rendering_template_id")),
        "artifact_intent": artifact_intent_from_payload(source_payload),
        "target_formats": target_formats_from_payload(source_payload),
        "artifact_title": artifact_title_from_payload(source_payload, structured_draft),
        "language": language_from_payload(source_payload),
        "retention_policy_ref": optional_text(source_payload.get("retention_policy_ref"))
        or DEFAULT_RETENTION_POLICY_REF,
        "quality_summary": {
            "citation_status": structured_draft["validation"]["citation_status"],
            "citation_count": len(structured_draft["citations"]),
            "validation_error_count": len(structured_draft["validation"]["errors"]),
            "warning_count": len(structured_draft["validation"]["warnings"]),
            "grounding_required": bool(request_metadata.get("grounding_required")),
            "retrieval_package_id": request_metadata.get("retrieval_package_id"),
            "retrieval_package_hash": request_metadata.get("retrieval_package_hash"),
            "evidence_ref_count": int(request_metadata.get("selected_evidence_count") or 0),
        },
        "created_at": now,
        "updated_at": now,
    }


def build_artifact_record_from_handoff(
    *,
    source_payload: dict[str, Any],
    handoff_record: dict[str, Any],
    artifact_request_id: str | None,
    request_id: str,
    trace_id: str,
) -> dict[str, Any]:
    validate_artifact_handoff_record(handoff_record)
    normalized_artifact_request_id = artifact_request_id or required_string(
        source_payload,
        "artifact_request_id",
        "ae.artifact_request_id_required",
    )
    artifact_id = optional_text(source_payload.get("artifact_id")) or str(
        uuid5(
            NAMESPACE_URL,
            (
                "ae-artifact:"
                f"{normalized_artifact_request_id}:{handoff_record['artifact_handoff_id']}"
            ),
        )
    )
    now = _utc_now()
    return {
        "artifact_schema_version": "ae_artifact_record.v1",
        "artifact_id": artifact_id,
        "artifact_type": artifact_type_from_payload(source_payload),
        "artifact_status": "DRAFT",
        "current_version_id": None,
        "trace_id": trace_id,
        "request_id": request_id,
        "chat_document_id": handoff_record["chat_document_id"],
        "interaction_id": handoff_record["interaction_id"],
        "owner_actor_ref": dict(handoff_record["actor_claims_ref"]),
        "workspace_ref": dict(handoff_record["workspace_ref"]),
        "display_title": optional_text(source_payload.get("display_title"))
        or handoff_record["artifact_title"],
        "language": handoff_record["language"],
        "artifact_intent": handoff_record["artifact_intent"],
        "target_formats": list(handoff_record["target_formats"]),
        "retention_policy_ref": handoff_record["retention_policy_ref"],
        "template_ref": {
            "template_id": handoff_record["template_id"],
            "template_version": handoff_record["template_version"],
            "rendering_template_id": handoff_record["rendering_template_id"],
        },
        "handoff_ref": {
            "artifact_handoff_id": handoff_record["artifact_handoff_id"],
            "artifact_request_id": handoff_record["artifact_request_id"],
            "handoff_schema_version": handoff_record["handoff_schema_version"],
        },
        "source_refs": [artifact_source_ref_from_handoff(artifact_id, handoff_record)],
        "versions": [],
        "render_jobs": [],
        "files": [],
        "links": [],
        "created_at": now,
        "updated_at": now,
    }


def artifact_source_ref_from_handoff(
    artifact_id: str,
    handoff_record: dict[str, Any],
) -> dict[str, Any]:
    quality_summary = dict(handoff_record["quality_summary"])
    evidence_ref_count = int(quality_summary.get("evidence_ref_count") or 0)
    source_ref_id = str(
        uuid5(
            NAMESPACE_URL,
            (
                "ae-artifact-source-ref:"
                f"{artifact_id}:{handoff_record['cx_generation_id']}:"
                f"{handoff_record['structured_draft_id']}"
            ),
        )
    )
    return {
        "source_ref_id": source_ref_id,
        "cx_generation_id": handoff_record["cx_generation_id"],
        "structured_draft_id": handoff_record["structured_draft_id"],
        "draft_schema_version": handoff_record["draft_schema_version"],
        "structured_draft_content_hash": handoff_record[
            "structured_draft_content_hash"
        ],
        "citation_claims_hash": handoff_record["citation_claims_hash"],
        "validation_result_hash": handoff_record["validation_result_hash"],
        "retrieval_package_id": quality_summary["retrieval_package_id"],
        "retrieval_package_hash": quality_summary["retrieval_package_hash"],
        "evidence_ref_count": evidence_ref_count,
        "source_anchor_count": evidence_ref_count,
        "quality_summary": quality_summary,
    }


def build_artifact_collection_filter(
    *,
    tenant_id: str,
    workspace_id: str,
    owner_user_id: str,
    status: str | None = None,
    limit: int | str | None = None,
) -> dict[str, Any]:
    normalized_status = optional_text(status)
    if normalized_status is not None:
        normalized_status = normalized_status.upper()
        if normalized_status not in SUPPORTED_ARTIFACT_STATUSES:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_collection_status_invalid",
                detail=f"Unsupported artifact collection status: {status}",
            )
    return {
        "tenant_id": _required_collection_scope_text(
            tenant_id,
            "tenant_id",
        ),
        "workspace_id": _required_collection_scope_text(
            workspace_id,
            "workspace_id",
        ),
        "owner_user_id": _required_collection_scope_text(
            owner_user_id,
            "owner_user_id",
        ),
        "status": normalized_status,
        "limit": normalize_artifact_collection_limit(limit),
    }


def normalize_artifact_collection_limit(limit: int | str | None) -> int:
    if limit is None:
        return DEFAULT_ARTIFACT_COLLECTION_LIMIT
    if isinstance(limit, bool):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_collection_limit_invalid",
            detail="Artifact collection limit must be an integer.",
        )
    try:
        normalized = int(limit)
    except (TypeError, ValueError) as exc:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_collection_limit_invalid",
            detail="Artifact collection limit must be an integer.",
        ) from exc
    if normalized < 1 or normalized > MAX_ARTIFACT_COLLECTION_LIMIT:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_collection_limit_invalid",
            detail=(
                "Artifact collection limit must be between 1 and "
                f"{MAX_ARTIFACT_COLLECTION_LIMIT}."
            ),
        )
    return normalized


def artifact_record_matches_collection_filter(
    record: dict[str, Any],
    collection_filter: dict[str, Any],
) -> bool:
    owner = record.get("owner_actor_ref", {})
    workspace = record.get("workspace_ref", {})
    return (
        isinstance(owner, dict)
        and isinstance(workspace, dict)
        and owner.get("tenant_id") == collection_filter["tenant_id"]
        and owner.get("actor_id") == collection_filter["owner_user_id"]
        and workspace.get("workspace_id") == collection_filter["workspace_id"]
        and (
            collection_filter["status"] is None
            or record.get("artifact_status") == collection_filter["status"]
        )
    )


def build_artifact_collection(
    records: list[dict[str, Any]],
    *,
    collection_filter: dict[str, Any],
) -> dict[str, Any]:
    items = [build_artifact_collection_item(record) for record in records]
    collection = {
        "artifact_collection_schema_version": ARTIFACT_COLLECTION_SCHEMA_VERSION,
        "filter": dict(collection_filter),
        "count": len(items),
        "limit": collection_filter["limit"],
        "next_cursor": None,
        "items": items,
    }
    assert_artifact_collection_payload_safe(collection)
    return collection


def build_artifact_collection_item(record: dict[str, Any]) -> dict[str, Any]:
    owner = record.get("owner_actor_ref", {})
    workspace = record.get("workspace_ref", {})
    source_ref = _first_mapping(record.get("source_refs"))
    versions = _list_of_mappings(record.get("versions"))
    render_jobs = _list_of_mappings(record.get("render_jobs"))
    files = _list_of_mappings(record.get("files"))
    links = _list_of_mappings(record.get("links"))
    item = {
        "artifact_collection_item_schema_version": (
            ARTIFACT_COLLECTION_ITEM_SCHEMA_VERSION
        ),
        "artifact_id": record["artifact_id"],
        "artifact_type": record["artifact_type"],
        "artifact_status": record["artifact_status"],
        "display_title": record["display_title"],
        "language": record["language"],
        "artifact_intent": record["artifact_intent"],
        "target_formats": list(record.get("target_formats", [])),
        "available_formats": _available_artifact_formats(files),
        "downloadable_formats": _linked_artifact_formats(files, links, "download"),
        "previewable_formats": _linked_artifact_formats(files, links, "preview"),
        "current_version_id": record.get("current_version_id"),
        "current_version_no": _current_version_no(
            versions,
            record.get("current_version_id"),
        ),
        "version_count": len(versions),
        "file_count": len(files),
        "link_count": len(links),
        "render_job_count": len(render_jobs),
        "latest_render_job": _latest_render_job_summary(render_jobs),
        "source_summary": _artifact_collection_source_summary(source_ref),
        "quality_summary": _artifact_collection_quality_summary(
            source_ref.get("quality_summary"),
        ),
        "routes": {
            "detail": f"/api/v1/artifacts/{record['artifact_id']}",
            "versions": f"/api/v1/artifacts/{record['artifact_id']}/versions",
        },
        "tenant_id": owner.get("tenant_id") if isinstance(owner, dict) else None,
        "workspace_id": (
            workspace.get("workspace_id") if isinstance(workspace, dict) else None
        ),
        "owner_user_id": owner.get("actor_id") if isinstance(owner, dict) else None,
        "chat_document_id": record["chat_document_id"],
        "interaction_id": record["interaction_id"],
        "created_at": record["created_at"],
        "updated_at": record["updated_at"],
    }
    assert_artifact_collection_payload_safe(item)
    return item


def assert_artifact_collection_payload_safe(payload: dict[str, Any]) -> None:
    serialized = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    blocked_tokens = (
        "storage_ref",
        "storage_root",
        "content_base64",
        "rendered_payloads",
        "rendered_markdown",
        "/data/nex-platform",
        "postgresql://",
        "postgresql+psycopg://",
        "nuri1004",
    )
    for token in blocked_tokens:
        if token in serialized:
            raise ArtifactHandoffError(
                status_code=500,
                error_code="ae.artifact_collection_payload_unsafe",
                detail="Artifact collection payload contains private material.",
            )


def _required_collection_scope_text(value: str, field_name: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_collection_scope_required",
            detail=f"{field_name} is required for artifact collection queries.",
        )
    return value.strip()


def _first_mapping(raw_value: Any) -> dict[str, Any]:
    values = _list_of_mappings(raw_value)
    return values[0] if values else {}


def _list_of_mappings(raw_value: Any) -> list[dict[str, Any]]:
    if not isinstance(raw_value, list):
        return []
    return [dict(value) for value in raw_value if isinstance(value, dict)]


def _available_artifact_formats(files: list[dict[str, Any]]) -> list[str]:
    return sorted(
        {
            str(artifact_file["format"])
            for artifact_file in files
            if artifact_file.get("format")
        }
    )


def _linked_artifact_formats(
    files: list[dict[str, Any]],
    links: list[dict[str, Any]],
    link_type: str,
) -> list[str]:
    linked_file_ids = {
        link.get("artifact_file_id")
        for link in links
        if link.get("link_type") == link_type and link.get("link_route")
    }
    return sorted(
        {
            str(artifact_file["format"])
            for artifact_file in files
            if artifact_file.get("artifact_file_id") in linked_file_ids
            and artifact_file.get("format")
        }
    )


def _current_version_no(
    versions: list[dict[str, Any]],
    current_version_id: Any,
) -> int | None:
    for version in versions:
        if version.get("artifact_version_id") == current_version_id:
            return int(version["version_no"])
    return None


def _latest_render_job_summary(
    render_jobs: list[dict[str, Any]],
) -> dict[str, Any] | None:
    if not render_jobs:
        return None
    render_job = render_jobs[-1]
    return {
        "render_job_id": render_job.get("render_job_id"),
        "job_status": render_job.get("job_status"),
        "current_stage": render_job.get("current_stage"),
        "progress_percent": render_job.get("progress_percent"),
        "retryable": bool(render_job.get("retryable")),
        "failure_code": render_job.get("failure_code"),
    }


def _artifact_collection_source_summary(
    source_ref: dict[str, Any],
) -> dict[str, Any]:
    return {
        "cx_generation_id": source_ref.get("cx_generation_id"),
        "structured_draft_id": source_ref.get("structured_draft_id"),
        "retrieval_package_id": source_ref.get("retrieval_package_id"),
        "retrieval_package_hash": source_ref.get("retrieval_package_hash"),
        "evidence_ref_count": int(source_ref.get("evidence_ref_count") or 0),
        "source_anchor_count": int(source_ref.get("source_anchor_count") or 0),
    }


def _artifact_collection_quality_summary(raw_value: Any) -> dict[str, Any]:
    if not isinstance(raw_value, dict):
        return {}
    allowed_keys = (
        "citation_status",
        "citation_count",
        "validation_error_count",
        "warning_count",
        "grounding_required",
        "evidence_ref_count",
    )
    return {
        key: _collection_json_safe_value(raw_value.get(key))
        for key in allowed_keys
        if raw_value.get(key) is not None
    }


def _collection_json_safe_value(raw_value: Any) -> Any:
    if isinstance(raw_value, (str, int, float, bool)) or raw_value is None:
        return raw_value
    if isinstance(raw_value, list):
        return [_collection_json_safe_value(value) for value in raw_value]
    if isinstance(raw_value, dict):
        return {
            str(key): _collection_json_safe_value(value)
            for key, value in raw_value.items()
            if value is not None
        }
    return str(raw_value)


def normalize_artifact_lifecycle_action(value: Any) -> str:
    normalized = optional_text(value)
    if normalized is None:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_lifecycle_action_required",
            detail="Artifact lifecycle action is required.",
        )
    normalized = normalized.replace("-", "_").upper()
    if normalized not in SUPPORTED_ARTIFACT_LIFECYCLE_ACTIONS:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_lifecycle_action_invalid",
            detail=f"Unsupported artifact lifecycle action: {value}",
        )
    return normalized


def normalize_artifact_restore_status(value: Any) -> str | None:
    normalized = optional_text(value)
    if normalized is None:
        return None
    normalized = normalized.upper()
    if normalized not in RESTORABLE_ARTIFACT_STATUSES:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_lifecycle_restore_status_invalid",
            detail=f"Unsupported artifact restore status: {value}",
        )
    return normalized


def artifact_lifecycle_target_status(
    *,
    current_status: str,
    action: str,
    restore_status: str | None = None,
) -> str:
    normalized_current_status = optional_text(current_status)
    if normalized_current_status not in SUPPORTED_ARTIFACT_STATUSES:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_lifecycle_current_status_invalid",
            detail=f"Unsupported artifact current status: {current_status}",
        )
    normalized_action = normalize_artifact_lifecycle_action(action)
    normalized_restore_status = normalize_artifact_restore_status(restore_status)

    if normalized_action == "ARCHIVE":
        if normalized_current_status == "ARCHIVED":
            return "ARCHIVED"
        if normalized_current_status not in ARCHIVABLE_ARTIFACT_STATUSES:
            raise ArtifactHandoffError(
                status_code=409,
                error_code="ae.artifact_lifecycle_transition_invalid",
                detail=f"Artifact status cannot be archived: {current_status}",
            )
        return "ARCHIVED"

    if normalized_action == "MARK_DELETED":
        if normalized_current_status == "DELETED":
            return "DELETED"
        if normalized_current_status not in DELETABLE_ARTIFACT_STATUSES:
            raise ArtifactHandoffError(
                status_code=409,
                error_code="ae.artifact_lifecycle_transition_invalid",
                detail=f"Artifact status cannot be marked deleted: {current_status}",
            )
        return "DELETED"

    if normalized_current_status not in {"ARCHIVED", "DELETED"}:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_lifecycle_transition_invalid",
            detail=f"Artifact status cannot be restored: {current_status}",
        )
    return normalized_restore_status or ARTIFACT_LIFECYCLE_DEFAULT_RESTORE_STATUS


def build_artifact_lifecycle_action_request(
    *,
    payload: dict[str, Any],
    artifact_record: dict[str, Any],
    request_id: str,
    trace_id: str,
    idempotency_key: str | None = None,
) -> dict[str, Any]:
    artifact_id = required_string(
        artifact_record,
        "artifact_id",
        "ae.artifact_id_required",
    )
    payload_artifact_id = optional_text(payload.get("artifact_id"))
    if payload_artifact_id is not None and payload_artifact_id != artifact_id:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_lifecycle_artifact_mismatch",
            detail="Lifecycle action artifact id does not match the route artifact.",
        )
    action = normalize_artifact_lifecycle_action(
        payload.get("action") or payload.get("lifecycle_action")
    )
    restore_status = normalize_artifact_restore_status(payload.get("restore_status"))
    previous_status = required_string(
        artifact_record,
        "artifact_status",
        "ae.artifact_status_required",
    )
    target_status = artifact_lifecycle_target_status(
        current_status=previous_status,
        action=action,
        restore_status=restore_status,
    )
    normalized_idempotency_key = (
        optional_text(idempotency_key)
        or optional_text(payload.get("idempotency_key"))
        or optional_text(payload.get("lifecycle_action_request_id"))
        or f"{request_id}:{artifact_id}:{action}:{target_status}"
    )
    comment_text = optional_text(payload.get("comment"))
    actor_ref = lifecycle_actor_ref_from_payload(payload, artifact_record)
    request = {
        "artifact_lifecycle_action_schema_version": (
            AE_ARTIFACT_LIFECYCLE_ACTION_SCHEMA_VERSION
        ),
        "lifecycle_action_id": str(
            uuid5(
                NAMESPACE_URL,
                f"ae-artifact-lifecycle-action:{artifact_id}:{normalized_idempotency_key}",
            )
        ),
        "artifact_id": artifact_id,
        "action": action,
        "previous_status": previous_status,
        "target_status": target_status,
        "restore_status": restore_status,
        "reason_code": (
            optional_text(payload.get("reason_code"))
            or ARTIFACT_LIFECYCLE_DEFAULT_REASON_CODE
        ),
        "comment_hash": sha256_text(comment_text) if comment_text else None,
        "comment_length": len(comment_text) if comment_text else 0,
        "actor_ref": actor_ref,
        "request_id": request_id,
        "trace_id": trace_id,
        "idempotency_key": normalized_idempotency_key,
        "metadata": {
            "physical_delete_requested": False,
            "storage_mutation_requested": False,
            "raw_comment_included": False,
        },
    }
    assert_artifact_lifecycle_payload_safe(request)
    return request


def build_artifact_lifecycle_action_result(
    *,
    action_request: dict[str, Any],
    artifact_record: dict[str, Any],
) -> dict[str, Any]:
    validate_artifact_lifecycle_action_request(action_request)
    result = {
        "artifact_lifecycle_action_result_schema_version": (
            AE_ARTIFACT_LIFECYCLE_ACTION_RESULT_SCHEMA_VERSION
        ),
        "lifecycle_action": dict(action_request),
        "artifact_id": artifact_record["artifact_id"],
        "artifact_status": artifact_record["artifact_status"],
        "previous_status": action_request["previous_status"],
        "target_status": action_request["target_status"],
        "transition_applied": artifact_record["artifact_status"]
        == action_request["target_status"],
        "routes": {
            "artifact": f"/api/v1/artifacts/{artifact_record['artifact_id']}",
            "collection": "/api/v1/artifacts",
        },
        "updated_at": artifact_record["updated_at"],
        "metadata": {
            "rendered_payload_included": False,
            "storage_location_included": False,
            "physical_delete_executed": False,
        },
    }
    assert_artifact_lifecycle_payload_safe(result)
    return result


def apply_artifact_lifecycle_action(
    *,
    artifact_record: dict[str, Any],
    action_request: dict[str, Any],
) -> dict[str, Any]:
    validate_artifact_lifecycle_action_request(action_request)
    if artifact_record["artifact_id"] != action_request["artifact_id"]:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_lifecycle_artifact_mismatch",
            detail="Lifecycle action artifact id does not match the artifact record.",
        )
    if artifact_record["artifact_status"] != action_request["previous_status"]:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_lifecycle_status_changed",
            detail="Artifact status changed before lifecycle action application.",
        )
    updated_record = dict(artifact_record)
    updated_record["artifact_status"] = action_request["target_status"]
    updated_record["updated_at"] = _utc_now()
    return updated_record


def validate_artifact_lifecycle_action_request(action_request: dict[str, Any]) -> None:
    if (
        action_request.get("artifact_lifecycle_action_schema_version")
        != AE_ARTIFACT_LIFECYCLE_ACTION_SCHEMA_VERSION
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_lifecycle_schema_invalid",
            detail="Artifact lifecycle action schema version is invalid.",
        )
    action = normalize_artifact_lifecycle_action(action_request.get("action"))
    previous_status = required_string(
        action_request,
        "previous_status",
        "ae.artifact_lifecycle_previous_status_required",
    )
    target_status = required_string(
        action_request,
        "target_status",
        "ae.artifact_lifecycle_target_status_required",
    )
    expected_target = artifact_lifecycle_target_status(
        current_status=previous_status,
        action=action,
        restore_status=action_request.get("restore_status"),
    )
    if target_status != expected_target:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_lifecycle_target_status_mismatch",
            detail="Artifact lifecycle target status does not match the action.",
        )
    assert_artifact_lifecycle_payload_safe(action_request)


def build_artifact_retention_policy(
    payload: dict[str, Any] | None = None,
) -> dict[str, Any]:
    source = payload or {}
    retention_days = normalize_artifact_retention_days(
        source.get("retention_days_after_logical_purge")
        if "retention_days_after_logical_purge" in source
        else source.get("retention_days")
    )
    policy = {
        "artifact_retention_policy_schema_version": (
            AE_ARTIFACT_RETENTION_POLICY_SCHEMA_VERSION
        ),
        "retention_policy_id": (
            optional_text(source.get("retention_policy_id"))
            or DEFAULT_ARTIFACT_RETENTION_POLICY_ID
        ),
        "logical_purge": {
            "enabled": True,
            "flag_field": ARTIFACT_RETENTION_LOGICAL_PURGE_FIELD,
            "flag_value": ARTIFACT_RETENTION_LOGICAL_PURGE_STATUS,
            "first_action": "MARK_DELETED",
            "reversible_before_physical_purge": True,
        },
        "physical_purge": {
            "enabled": False,
            "execution_mode": "scheduled_batch_after_retention",
            "retention_days_after_logical_purge": retention_days,
            "supported_retention_day_presets": list(
                SUPPORTED_ARTIFACT_RETENTION_DAY_PRESETS
            ),
            "batch_window": {
                "timezone": ARTIFACT_RETENTION_BATCH_TIMEZONE,
                "start_local_time": ARTIFACT_RETENTION_BATCH_WINDOW_START,
                "end_local_time": ARTIFACT_RETENTION_BATCH_WINDOW_END,
            },
            "dry_run_required": True,
            "storage_mutation_enabled": False,
            "database_row_delete_enabled": False,
        },
        "candidate_query": {
            "status": ARTIFACT_RETENTION_LOGICAL_PURGE_STATUS,
            "timestamp_field": "updated_at",
            "metadata_only": True,
            "default_limit": DEFAULT_ARTIFACT_COLLECTION_LIMIT,
            "max_limit": MAX_ARTIFACT_COLLECTION_LIMIT,
        },
    }
    validate_artifact_retention_policy(policy)
    return policy


def normalize_artifact_retention_days(value: Any) -> int:
    if value is None:
        return DEFAULT_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE
    if isinstance(value, bool):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_days_invalid",
            detail="Artifact retention days must be an integer.",
        )
    try:
        normalized = int(value)
    except (TypeError, ValueError) as exc:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_days_invalid",
            detail="Artifact retention days must be an integer.",
        ) from exc
    if (
        normalized < MIN_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE
        or normalized > MAX_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_days_invalid",
            detail=(
                "Artifact retention days must be between "
                f"{MIN_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE} and "
                f"{MAX_ARTIFACT_RETENTION_DAYS_AFTER_LOGICAL_PURGE}."
            ),
        )
    return normalized


def validate_artifact_retention_policy(policy: dict[str, Any]) -> None:
    if (
        policy.get("artifact_retention_policy_schema_version")
        != AE_ARTIFACT_RETENTION_POLICY_SCHEMA_VERSION
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_policy_schema_invalid",
            detail="Artifact retention policy schema version is invalid.",
        )
    logical_purge = policy.get("logical_purge")
    physical_purge = policy.get("physical_purge")
    candidate_query = policy.get("candidate_query")
    if not isinstance(logical_purge, dict) or not isinstance(physical_purge, dict):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_policy_invalid",
            detail="Artifact retention policy purge sections are required.",
        )
    if logical_purge.get("flag_field") != ARTIFACT_RETENTION_LOGICAL_PURGE_FIELD:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_policy_invalid",
            detail="Artifact retention logical purge flag field is invalid.",
        )
    if logical_purge.get("flag_value") != ARTIFACT_RETENTION_LOGICAL_PURGE_STATUS:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_policy_invalid",
            detail="Artifact retention logical purge flag value is invalid.",
        )
    if physical_purge.get("enabled") is not False:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_policy_invalid",
            detail="Artifact physical purge must stay disabled in this policy.",
        )
    if physical_purge.get("dry_run_required") is not True:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_policy_invalid",
            detail="Artifact retention candidate scans must require dry-run first.",
        )
    normalize_artifact_retention_days(
        physical_purge.get("retention_days_after_logical_purge")
    )
    if not isinstance(candidate_query, dict) or candidate_query.get("metadata_only") is not True:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_policy_invalid",
            detail="Artifact retention candidate query must be metadata-only.",
        )
    assert_artifact_retention_payload_safe(policy)


def assert_artifact_retention_payload_safe(payload: dict[str, Any]) -> None:
    serialized = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    blocked_tokens = (
        "raw_prompt",
        "raw_text",
        "raw_source",
        "rendered_payloads",
        "rendered_markdown",
        "content_base64",
        "storage_ref",
        "storage_root",
        "/data/nex-platform",
        "postgresql://",
        "postgresql+psycopg://",
        "nuri1004",
        "ed6@c496em",
    )
    for token in blocked_tokens:
        if token in serialized:
            raise ArtifactHandoffError(
                status_code=500,
                error_code="ae.artifact_retention_payload_unsafe",
                detail="Artifact retention payload contains private material.",
            )


def build_artifact_retention_candidate_filter(
    *,
    tenant_id: str,
    workspace_id: str,
    owner_user_id: str,
    retention_days: int | str | None = None,
    as_of: str | None = None,
    limit: int | str | None = None,
) -> dict[str, Any]:
    normalized_days = normalize_artifact_retention_days(retention_days)
    as_of_at = parse_artifact_retention_timestamp(as_of, field_name="as_of")
    cutoff_at = as_of_at - timedelta(days=normalized_days)
    candidate_filter = {
        "tenant_id": _required_collection_scope_text(tenant_id, "tenant_id"),
        "workspace_id": _required_collection_scope_text(
            workspace_id,
            "workspace_id",
        ),
        "owner_user_id": _required_collection_scope_text(
            owner_user_id,
            "owner_user_id",
        ),
        "status": ARTIFACT_RETENTION_LOGICAL_PURGE_STATUS,
        "retention_days_after_logical_purge": normalized_days,
        "as_of": format_artifact_retention_timestamp(as_of_at),
        "cutoff_at": format_artifact_retention_timestamp(cutoff_at),
        "limit": normalize_artifact_collection_limit(limit),
        "dry_run": True,
    }
    assert_artifact_retention_payload_safe(candidate_filter)
    return candidate_filter


def artifact_record_matches_retention_candidate_filter(
    record: dict[str, Any],
    candidate_filter: dict[str, Any],
) -> bool:
    owner = record.get("owner_actor_ref", {})
    workspace = record.get("workspace_ref", {})
    if not (
        isinstance(owner, dict)
        and isinstance(workspace, dict)
        and owner.get("tenant_id") == candidate_filter["tenant_id"]
        and owner.get("actor_id") == candidate_filter["owner_user_id"]
        and workspace.get("workspace_id") == candidate_filter["workspace_id"]
        and record.get("artifact_status") == ARTIFACT_RETENTION_LOGICAL_PURGE_STATUS
    ):
        return False
    updated_at = parse_artifact_retention_timestamp(
        record.get("updated_at"),
        field_name="updated_at",
    )
    cutoff_at = parse_artifact_retention_timestamp(
        candidate_filter["cutoff_at"],
        field_name="cutoff_at",
    )
    return updated_at <= cutoff_at


def build_artifact_retention_candidate_collection(
    records: list[dict[str, Any]],
    *,
    candidate_filter: dict[str, Any],
) -> dict[str, Any]:
    items = [
        build_artifact_retention_candidate_item(
            record,
            candidate_filter=candidate_filter,
        )
        for record in records
    ]
    collection = {
        "artifact_retention_candidate_collection_schema_version": (
            AE_ARTIFACT_RETENTION_CANDIDATE_COLLECTION_SCHEMA_VERSION
        ),
        "policy": build_artifact_retention_policy(
            {
                "retention_days_after_logical_purge": candidate_filter[
                    "retention_days_after_logical_purge"
                ]
            }
        ),
        "filter": dict(candidate_filter),
        "count": len(items),
        "limit": candidate_filter["limit"],
        "next_cursor": None,
        "items": items,
        "metadata": {
            "dry_run": True,
            "physical_delete_executed": False,
            "storage_mutation_executed": False,
            "database_row_delete_executed": False,
        },
    }
    assert_artifact_retention_payload_safe(collection)
    return collection


def build_artifact_retention_candidate_item(
    record: dict[str, Any],
    *,
    candidate_filter: dict[str, Any],
) -> dict[str, Any]:
    owner = record.get("owner_actor_ref", {})
    workspace = record.get("workspace_ref", {})
    logical_purged_at = parse_artifact_retention_timestamp(
        record.get("updated_at"),
        field_name="updated_at",
    )
    as_of_at = parse_artifact_retention_timestamp(
        candidate_filter["as_of"],
        field_name="as_of",
    )
    retention_days = int(candidate_filter["retention_days_after_logical_purge"])
    purge_eligible_at = logical_purged_at + timedelta(days=retention_days)
    item = {
        "artifact_retention_candidate_item_schema_version": (
            AE_ARTIFACT_RETENTION_CANDIDATE_ITEM_SCHEMA_VERSION
        ),
        "artifact_id": record["artifact_id"],
        "artifact_status": record["artifact_status"],
        "display_title": record["display_title"],
        "tenant_id": owner.get("tenant_id") if isinstance(owner, dict) else None,
        "workspace_id": (
            workspace.get("workspace_id") if isinstance(workspace, dict) else None
        ),
        "owner_user_id": owner.get("actor_id") if isinstance(owner, dict) else None,
        "retention_policy_ref": record["retention_policy_ref"],
        "current_version_id": record.get("current_version_id"),
        "version_count": len(_list_of_mappings(record.get("versions"))),
        "file_count": len(_list_of_mappings(record.get("files"))),
        "link_count": len(_list_of_mappings(record.get("links"))),
        "render_job_count": len(_list_of_mappings(record.get("render_jobs"))),
        "logical_purged_at": format_artifact_retention_timestamp(logical_purged_at),
        "purge_eligible_at": format_artifact_retention_timestamp(purge_eligible_at),
        "age_days_after_logical_purge": max(0, (as_of_at - logical_purged_at).days),
        "routes": {
            "artifact": f"/api/v1/artifacts/{record['artifact_id']}",
            "lifecycle": (
                f"/api/v1/artifacts/{record['artifact_id']}/lifecycle-actions"
            ),
        },
        "purge_plan": {
            "dry_run": True,
            "logical_purge_already_applied": True,
            "physical_delete_deferred": True,
            "storage_mutation_enabled": False,
            "database_row_delete_enabled": False,
            "planned_execution": "scheduled_batch_after_retention",
        },
    }
    assert_artifact_retention_payload_safe(item)
    return item


def parse_artifact_retention_timestamp(value: Any, *, field_name: str) -> datetime:
    normalized = optional_text(value)
    if normalized is None:
        if field_name == "as_of":
            return datetime.now(UTC)
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_timestamp_invalid",
            detail=f"{field_name} must be an ISO-8601 timestamp.",
        )
    try:
        parsed = datetime.fromisoformat(normalized.replace("Z", "+00:00"))
    except ValueError as exc:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_timestamp_invalid",
            detail=f"{field_name} must be an ISO-8601 timestamp.",
        ) from exc
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=UTC)
    return parsed.astimezone(UTC)


def format_artifact_retention_timestamp(value: datetime) -> str:
    return value.astimezone(UTC).isoformat().replace("+00:00", "Z")


def build_artifact_retention_execution(
    *,
    tenant_id: str,
    workspace_id: str,
    owner_user_id: str,
    mode: str = "DRY_RUN",
    execution_status: str = "PLANNED",
    retention_days: int | str | None = None,
    as_of: str | None = None,
    checked_at: str | None = None,
    scan_limit: int | str | None = None,
    max_delete_count: int | str | None = None,
    candidate_count: int = 0,
    selected_count: int = 0,
    deleted_counts: dict[str, int] | None = None,
    delete_enabled: bool = False,
    storage_mutation_enabled: bool = False,
    database_row_delete_enabled: bool = False,
    requested_by: dict[str, Any] | None = None,
    idempotency_key: str | None = None,
    trace_id: str | None = None,
    request_id: str | None = None,
    blocked_reason: str | None = None,
    error: dict[str, Any] | None = None,
    execution_id: str | None = None,
) -> dict[str, Any]:
    candidate_filter = build_artifact_retention_candidate_filter(
        tenant_id=tenant_id,
        workspace_id=workspace_id,
        owner_user_id=owner_user_id,
        retention_days=retention_days,
        as_of=as_of,
        limit=scan_limit,
    )
    normalized_checked_at = format_artifact_retention_timestamp(
        parse_artifact_retention_timestamp(
            checked_at or _utc_now(),
            field_name="checked_at",
        )
    )
    normalized_mode = _normalize_artifact_retention_execution_mode(mode)
    normalized_status = _normalize_artifact_retention_execution_status(
        execution_status
    )
    normalized_deleted_counts = _normalize_artifact_retention_deleted_counts(
        deleted_counts
    )
    normalized_max_delete_count = normalize_artifact_retention_delete_limit(
        max_delete_count
    )
    normalized_execution_id = execution_id or _artifact_retention_execution_id(
        mode=normalized_mode,
        execution_status=normalized_status,
        candidate_filter=candidate_filter,
        checked_at=normalized_checked_at,
        idempotency_key=idempotency_key,
    )
    execution = {
        "artifact_retention_execution_schema_version": (
            AE_ARTIFACT_RETENTION_EXECUTION_SCHEMA_VERSION
        ),
        "execution_id": normalized_execution_id,
        "policy_id": DEFAULT_ARTIFACT_RETENTION_POLICY_ID,
        "service_id": "nex-ae-api",
        "mode": normalized_mode,
        "execution_status": normalized_status,
        "delete_enabled": bool(delete_enabled),
        "storage_mutation_enabled": bool(storage_mutation_enabled),
        "database_row_delete_enabled": bool(database_row_delete_enabled),
        "tenant_id": candidate_filter["tenant_id"],
        "workspace_id": candidate_filter["workspace_id"],
        "owner_user_id": candidate_filter["owner_user_id"],
        "retention_days_after_logical_purge": (
            candidate_filter["retention_days_after_logical_purge"]
        ),
        "as_of": candidate_filter["as_of"],
        "cutoff_at": candidate_filter["cutoff_at"],
        "checked_at": normalized_checked_at,
        "scan_limit": candidate_filter["limit"],
        "max_delete_count": normalized_max_delete_count,
        "candidate_count": _non_negative_artifact_retention_int(
            candidate_count,
            "candidate_count",
        ),
        "selected_count": _non_negative_artifact_retention_int(
            selected_count,
            "selected_count",
        ),
        "deleted_counts": normalized_deleted_counts,
        "requested_by": _normalize_artifact_retention_requested_by(requested_by),
        "idempotency_key": optional_text(idempotency_key),
        "trace_id": optional_text(trace_id),
        "request_id": optional_text(request_id),
        "blocked_reason": optional_text(blocked_reason),
        "error": _normalize_artifact_retention_error(error),
        "audit": {
            "audit_event_type": "ae_artifact.retention.execution",
            "audit_event_id": _artifact_retention_audit_id(normalized_execution_id),
            "emitted": False,
        },
        "metadata": {
            "logical_purge_required_before_physical_delete": True,
            "candidate_scan_metadata_only": True,
            "scheduled_batch_timezone": ARTIFACT_RETENTION_BATCH_TIMEZONE,
            "scheduled_batch_window": {
                "start_local_time": ARTIFACT_RETENTION_BATCH_WINDOW_START,
                "end_local_time": ARTIFACT_RETENTION_BATCH_WINDOW_END,
            },
        },
    }
    return validate_artifact_retention_execution(execution)


def validate_artifact_retention_execution(
    execution: dict[str, Any],
) -> dict[str, Any]:
    if not isinstance(execution, dict):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_invalid",
            detail="Artifact retention execution must be an object.",
        )
    for field_name in (
        "artifact_retention_execution_schema_version",
        "execution_id",
        "policy_id",
        "service_id",
        "mode",
        "execution_status",
        "delete_enabled",
        "storage_mutation_enabled",
        "database_row_delete_enabled",
        "tenant_id",
        "workspace_id",
        "owner_user_id",
        "retention_days_after_logical_purge",
        "as_of",
        "cutoff_at",
        "checked_at",
        "scan_limit",
        "max_delete_count",
        "candidate_count",
        "selected_count",
        "deleted_counts",
        "requested_by",
        "idempotency_key",
        "trace_id",
        "request_id",
        "blocked_reason",
        "error",
        "audit",
        "metadata",
    ):
        if field_name not in execution:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_retention_execution_invalid",
                detail=f"Missing artifact retention execution field: {field_name}.",
            )
    if (
        execution["artifact_retention_execution_schema_version"]
        != AE_ARTIFACT_RETENTION_EXECUTION_SCHEMA_VERSION
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_schema_invalid",
            detail="Artifact retention execution schema version is invalid.",
        )
    if execution["policy_id"] != DEFAULT_ARTIFACT_RETENTION_POLICY_ID:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_policy_invalid",
            detail="Artifact retention execution policy id is invalid.",
        )
    if execution["service_id"] != "nex-ae-api":
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_service_invalid",
            detail="Artifact retention execution service id is invalid.",
        )
    _required_collection_scope_text(execution["tenant_id"], "tenant_id")
    _required_collection_scope_text(execution["workspace_id"], "workspace_id")
    _required_collection_scope_text(execution["owner_user_id"], "owner_user_id")
    mode = _normalize_artifact_retention_execution_mode(execution["mode"])
    status = _normalize_artifact_retention_execution_status(
        execution["execution_status"]
    )
    for field_name in (
        "delete_enabled",
        "storage_mutation_enabled",
        "database_row_delete_enabled",
    ):
        if not isinstance(execution[field_name], bool):
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_retention_execution_flag_invalid",
                detail=f"{field_name} must be a boolean.",
            )
    if mode == "DRY_RUN" and (
        execution["delete_enabled"]
        or execution["storage_mutation_enabled"]
        or execution["database_row_delete_enabled"]
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_dry_run_delete_enabled_invalid",
            detail="Dry-run artifact retention execution cannot enable deletes.",
        )
    if (
        mode == "EXECUTE"
        and status == "SUCCEEDED"
        and not (
            execution["delete_enabled"]
            and execution["storage_mutation_enabled"]
            and execution["database_row_delete_enabled"]
        )
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execute_not_enabled",
            detail="Successful artifact retention execute requires all delete flags.",
        )
    retention_days = normalize_artifact_retention_days(
        execution["retention_days_after_logical_purge"]
    )
    if retention_days != execution["retention_days_after_logical_purge"]:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_days_invalid",
            detail="Artifact retention execution days must be normalized.",
        )
    for field_name in ("as_of", "cutoff_at", "checked_at"):
        parse_artifact_retention_timestamp(
            execution[field_name],
            field_name=field_name,
        )
    if normalize_artifact_collection_limit(execution["scan_limit"]) != execution[
        "scan_limit"
    ]:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_scan_limit_invalid",
            detail="Artifact retention scan limit must be normalized.",
        )
    if (
        normalize_artifact_retention_delete_limit(execution["max_delete_count"])
        != execution["max_delete_count"]
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_delete_limit_invalid",
            detail="Artifact retention delete limit must be normalized.",
        )
    candidate_count = _non_negative_artifact_retention_int(
        execution["candidate_count"],
        "candidate_count",
    )
    selected_count = _non_negative_artifact_retention_int(
        execution["selected_count"],
        "selected_count",
    )
    if selected_count > candidate_count:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_selected_count_invalid",
            detail="Artifact retention selected_count cannot exceed candidate_count.",
        )
    if selected_count > execution["max_delete_count"]:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_selected_count_invalid",
            detail="Artifact retention selected_count cannot exceed max_delete_count.",
        )
    deleted_counts = _normalize_artifact_retention_deleted_counts(
        execution["deleted_counts"]
    )
    if mode == "DRY_RUN" or status in {"PLANNED", "BLOCKED"}:
        if any(count != 0 for count in deleted_counts.values()):
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_retention_deleted_counts_invalid",
                detail="Dry-run, planned, and blocked executions cannot delete rows.",
            )
    if mode == "EXECUTE" and status == "SUCCEEDED":
        if deleted_counts["artifacts"] != selected_count:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_retention_deleted_counts_invalid",
                detail="Deleted artifact count must match selected_count.",
            )
    _normalize_artifact_retention_requested_by(execution["requested_by"])
    _normalize_artifact_retention_error(execution["error"])
    _validate_artifact_retention_audit(execution["audit"], execution["execution_id"])
    metadata = execution["metadata"]
    if (
        not isinstance(metadata, dict)
        or metadata.get("candidate_scan_metadata_only") is not True
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_metadata_invalid",
            detail="Artifact retention execution metadata is invalid.",
        )
    assert_artifact_retention_payload_safe(execution)
    return execution


def build_artifact_retention_execution_history_record(
    execution: dict[str, Any],
    *,
    created_at: str | None = None,
) -> dict[str, Any]:
    validated_execution = validate_artifact_retention_execution(execution)
    execution_payload = json.loads(
        json.dumps(validated_execution, ensure_ascii=False, sort_keys=True)
    )
    record = {
        "retention_execution_id": execution_payload["execution_id"],
        "execution_history_schema_version": (
            AE_ARTIFACT_RETENTION_EXECUTION_HISTORY_SCHEMA_VERSION
        ),
        "artifact_retention_execution_schema_version": (
            execution_payload["artifact_retention_execution_schema_version"]
        ),
        "policy_id": execution_payload["policy_id"],
        "service_id": execution_payload["service_id"],
        "mode": execution_payload["mode"],
        "execution_status": execution_payload["execution_status"],
        "tenant_id": execution_payload["tenant_id"],
        "workspace_id": execution_payload["workspace_id"],
        "owner_user_id": execution_payload["owner_user_id"],
        "retention_days_after_logical_purge": (
            execution_payload["retention_days_after_logical_purge"]
        ),
        "as_of": execution_payload["as_of"],
        "cutoff_at": execution_payload["cutoff_at"],
        "checked_at": execution_payload["checked_at"],
        "scan_limit": execution_payload["scan_limit"],
        "max_delete_count": execution_payload["max_delete_count"],
        "candidate_count": execution_payload["candidate_count"],
        "selected_count": execution_payload["selected_count"],
        "delete_enabled": execution_payload["delete_enabled"],
        "storage_mutation_enabled": execution_payload["storage_mutation_enabled"],
        "database_row_delete_enabled": execution_payload["database_row_delete_enabled"],
        "deleted_counts": dict(execution_payload["deleted_counts"]),
        "requested_by": dict(execution_payload["requested_by"]),
        "idempotency_key": execution_payload["idempotency_key"],
        "trace_id": execution_payload["trace_id"],
        "request_id": execution_payload["request_id"],
        "blocked_reason": execution_payload["blocked_reason"],
        "error": execution_payload["error"],
        "audit": dict(execution_payload["audit"]),
        "metadata": dict(execution_payload["metadata"]),
        "execution": execution_payload,
        "execution_payload_hash": sha256_json(execution_payload),
        "created_at": _artifact_retention_history_created_at(
            created_at,
            execution_payload["checked_at"],
        ),
    }
    return validate_artifact_retention_execution_history_record(record)


def validate_artifact_retention_execution_history_record(
    record: dict[str, Any],
) -> dict[str, Any]:
    if not isinstance(record, dict):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_history_invalid",
            detail="Artifact retention execution history record must be an object.",
        )
    required_fields = (
        "retention_execution_id",
        "execution_history_schema_version",
        "artifact_retention_execution_schema_version",
        "policy_id",
        "service_id",
        "mode",
        "execution_status",
        "tenant_id",
        "workspace_id",
        "owner_user_id",
        "retention_days_after_logical_purge",
        "as_of",
        "cutoff_at",
        "checked_at",
        "scan_limit",
        "max_delete_count",
        "candidate_count",
        "selected_count",
        "delete_enabled",
        "storage_mutation_enabled",
        "database_row_delete_enabled",
        "deleted_counts",
        "requested_by",
        "idempotency_key",
        "trace_id",
        "request_id",
        "blocked_reason",
        "error",
        "audit",
        "metadata",
        "execution",
        "execution_payload_hash",
        "created_at",
    )
    for field_name in required_fields:
        if field_name not in record:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_retention_history_invalid",
                detail=f"Missing artifact retention history field: {field_name}.",
            )
    if (
        record["execution_history_schema_version"]
        != AE_ARTIFACT_RETENTION_EXECUTION_HISTORY_SCHEMA_VERSION
    ):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_history_schema_invalid",
            detail="Artifact retention execution history schema version is invalid.",
        )
    execution = validate_artifact_retention_execution(record["execution"])
    for field_name in (
        "artifact_retention_execution_schema_version",
        "policy_id",
        "service_id",
        "mode",
        "execution_status",
        "tenant_id",
        "workspace_id",
        "owner_user_id",
        "retention_days_after_logical_purge",
        "as_of",
        "cutoff_at",
        "checked_at",
        "scan_limit",
        "max_delete_count",
        "candidate_count",
        "selected_count",
        "delete_enabled",
        "storage_mutation_enabled",
        "database_row_delete_enabled",
        "deleted_counts",
        "requested_by",
        "idempotency_key",
        "trace_id",
        "request_id",
        "blocked_reason",
        "error",
        "audit",
        "metadata",
    ):
        if record[field_name] != execution[field_name]:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_retention_history_mismatch",
                detail=f"Artifact retention history field mismatch: {field_name}.",
            )
    if record["retention_execution_id"] != execution["execution_id"]:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_history_mismatch",
            detail="Artifact retention history id must match execution id.",
        )
    if record["execution_payload_hash"] != sha256_json(execution):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_history_hash_invalid",
            detail="Artifact retention history payload hash is invalid.",
        )
    parse_artifact_retention_timestamp(record["created_at"], field_name="created_at")
    assert_artifact_retention_payload_safe(record)
    return record


def _artifact_retention_history_created_at(
    created_at: str | None,
    checked_at: str,
) -> str:
    return format_artifact_retention_timestamp(
        parse_artifact_retention_timestamp(
            created_at or checked_at,
            field_name="created_at",
        )
    )


def normalize_artifact_retention_delete_limit(value: int | str | None) -> int:
    if value is None:
        return DEFAULT_ARTIFACT_RETENTION_MAX_DELETE_COUNT
    if isinstance(value, bool):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_delete_limit_invalid",
            detail="Artifact retention delete limit must be an integer.",
        )
    try:
        normalized = int(value)
    except (TypeError, ValueError) as exc:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_delete_limit_invalid",
            detail="Artifact retention delete limit must be an integer.",
        ) from exc
    if normalized < 1 or normalized > MAX_ARTIFACT_RETENTION_MAX_DELETE_COUNT:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_delete_limit_invalid",
            detail=(
                "Artifact retention delete limit must be between 1 and "
                f"{MAX_ARTIFACT_RETENTION_MAX_DELETE_COUNT}."
            ),
        )
    return normalized


def _normalize_artifact_retention_execution_mode(value: Any) -> str:
    normalized = optional_text(value)
    if normalized is None:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_mode_invalid",
            detail="Artifact retention execution mode is required.",
        )
    normalized = normalized.replace("-", "_").upper()
    if normalized not in ARTIFACT_RETENTION_EXECUTION_MODES:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_mode_invalid",
            detail="Artifact retention execution mode must be DRY_RUN or EXECUTE.",
        )
    return normalized


def _normalize_artifact_retention_execution_status(value: Any) -> str:
    normalized = optional_text(value)
    if normalized is None:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_status_invalid",
            detail="Artifact retention execution status is required.",
        )
    normalized = normalized.upper()
    if normalized not in ARTIFACT_RETENTION_EXECUTION_STATUSES:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_execution_status_invalid",
            detail="Artifact retention execution status is not supported.",
        )
    return normalized


def _normalize_artifact_retention_deleted_counts(
    counts: dict[str, int] | None,
) -> dict[str, int]:
    source = counts or {}
    return {
        "artifacts": _non_negative_artifact_retention_int(
            source.get("artifacts", 0),
            "deleted_counts.artifacts",
        ),
        "source_refs": _non_negative_artifact_retention_int(
            source.get("source_refs", 0),
            "deleted_counts.source_refs",
        ),
        "versions": _non_negative_artifact_retention_int(
            source.get("versions", 0),
            "deleted_counts.versions",
        ),
        "render_jobs": _non_negative_artifact_retention_int(
            source.get("render_jobs", 0),
            "deleted_counts.render_jobs",
        ),
        "files": _non_negative_artifact_retention_int(
            source.get("files", 0),
            "deleted_counts.files",
        ),
        "links": _non_negative_artifact_retention_int(
            source.get("links", 0),
            "deleted_counts.links",
        ),
        "storage_files": _non_negative_artifact_retention_int(
            source.get("storage_files", 0),
            "deleted_counts.storage_files",
        ),
    }


def _empty_artifact_retention_deleted_counts() -> dict[str, int]:
    return _normalize_artifact_retention_deleted_counts(None)


def _non_negative_artifact_retention_int(value: Any, field_name: str) -> int:
    if isinstance(value, bool):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_count_invalid",
            detail=f"{field_name} must be a non-negative integer.",
        )
    try:
        normalized = int(value)
    except (TypeError, ValueError) as exc:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_count_invalid",
            detail=f"{field_name} must be a non-negative integer.",
        ) from exc
    if normalized < 0:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_count_invalid",
            detail=f"{field_name} must be a non-negative integer.",
        )
    return normalized


def _normalize_artifact_retention_requested_by(
    requested_by: dict[str, Any] | None,
) -> dict[str, str]:
    source = requested_by or {}
    actor_type = optional_text(source.get("actor_type")) or "service"
    actor_id = optional_text(source.get("actor_id")) or "nex-ae-api"
    service_id = optional_text(source.get("service_id")) or "nex-ae-api"
    normalized = {
        "actor_type": actor_type,
        "actor_id": actor_id,
        "service_id": service_id,
    }
    assert_artifact_retention_payload_safe(normalized)
    return normalized


def _normalize_artifact_retention_error(
    error: dict[str, Any] | None,
) -> dict[str, str] | None:
    if error is None:
        return None
    if not isinstance(error, dict):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_error_invalid",
            detail="Artifact retention error must be an object.",
        )
    return {
        "error_code": required_string(
            error,
            "error_code",
            "ae.artifact_retention_error_invalid",
        ),
        "detail": required_string(
            error,
            "detail",
            "ae.artifact_retention_error_invalid",
        ),
    }


def _validate_artifact_retention_audit(
    audit: dict[str, Any],
    execution_id: str,
) -> None:
    if not isinstance(audit, dict):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_audit_invalid",
            detail="Artifact retention audit must be an object.",
        )
    if audit.get("audit_event_type") != "ae_artifact.retention.execution":
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_audit_invalid",
            detail="Artifact retention audit event type is invalid.",
        )
    if audit.get("audit_event_id") != _artifact_retention_audit_id(execution_id):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_audit_invalid",
            detail="Artifact retention audit event id is invalid.",
        )
    if not isinstance(audit.get("emitted"), bool):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_retention_audit_invalid",
            detail="Artifact retention audit emitted flag is invalid.",
        )


def _artifact_retention_execution_id(
    *,
    mode: str,
    execution_status: str,
    candidate_filter: dict[str, Any],
    checked_at: str,
    idempotency_key: str | None,
) -> str:
    basis = {
        "mode": mode,
        "execution_status": execution_status,
        "tenant_id": candidate_filter["tenant_id"],
        "workspace_id": candidate_filter["workspace_id"],
        "owner_user_id": candidate_filter["owner_user_id"],
        "cutoff_at": candidate_filter["cutoff_at"],
        "checked_at": checked_at,
        "idempotency_key": optional_text(idempotency_key),
    }
    return str(
        uuid5(
            NAMESPACE_URL,
            f"ae-artifact-retention-execution:{sha256_json(basis)}",
        )
    )


def _artifact_retention_audit_id(execution_id: str) -> str:
    return str(uuid5(NAMESPACE_URL, f"ae-artifact-retention-audit:{execution_id}"))


def lifecycle_actor_ref_from_payload(
    payload: dict[str, Any],
    artifact_record: dict[str, Any],
) -> dict[str, Any]:
    raw_actor = payload.get("actor_ref")
    owner_ref = artifact_record.get("owner_actor_ref", {})
    owner = owner_ref if isinstance(owner_ref, dict) else {}
    if isinstance(raw_actor, dict):
        actor_ref = {
            "actor_type": optional_text(raw_actor.get("actor_type")) or "user",
            "actor_id": optional_text(raw_actor.get("actor_id")) or "unknown",
            "tenant_id": optional_text(raw_actor.get("tenant_id"))
            or optional_text(owner.get("tenant_id"))
            or DEFAULT_TENANT_ID,
        }
    else:
        actor_ref = {
            "actor_type": optional_text(owner.get("actor_type")) or "user",
            "actor_id": optional_text(owner.get("actor_id")) or DEFAULT_OWNER_USER_ID,
            "tenant_id": optional_text(owner.get("tenant_id")) or DEFAULT_TENANT_ID,
        }
    assert_artifact_lifecycle_payload_safe(actor_ref)
    return actor_ref


def assert_artifact_lifecycle_payload_safe(payload: dict[str, Any]) -> None:
    serialized = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    blocked_tokens = (
        "raw_prompt",
        "raw_text",
        "raw_source",
        "rendered_payloads",
        "rendered_markdown",
        "comment_text",
        "storage_ref",
        "storage_root",
        "/data/nex-platform",
        "postgresql://",
        "postgresql+psycopg://",
        "nuri1004",
        "ed6@c496em",
    )
    for token in blocked_tokens:
        if token in serialized:
            raise ArtifactHandoffError(
                status_code=500,
                error_code="ae.artifact_lifecycle_payload_unsafe",
                detail="Artifact lifecycle payload contains private material.",
            )


def build_markdown_render_result(
    *,
    artifact_record: dict[str, Any],
    structured_draft: dict[str, Any],
    target_formats: list[str],
    render_request_id: str,
    render_job_id: str,
) -> dict[str, Any]:
    validate_structured_draft_for_markdown_render(artifact_record, structured_draft)
    markdown = render_markdown_from_structured_draft(structured_draft)
    now = _utc_now()
    artifact_content_hash = sha256_text(markdown)
    source_ref = artifact_record["source_refs"][0]
    artifact_version_id = str(
        uuid5(
            NAMESPACE_URL,
            (
                "ae-artifact-version:"
                f"{artifact_record['artifact_id']}:{len(artifact_record['versions']) + 1}:"
                f"{artifact_content_hash}"
            ),
        )
    )
    artifact_version = {
        "artifact_version_id": artifact_version_id,
        "artifact_id": artifact_record["artifact_id"],
        "version_no": len(artifact_record["versions"]) + 1,
        "version_reason": "initial_render"
        if not artifact_record["versions"]
        else "rerender",
        "source_generation_id": source_ref["cx_generation_id"],
        "source_structured_draft_id": source_ref["structured_draft_id"],
        "source_content_hash": source_ref["structured_draft_content_hash"],
        "source_citation_claims_hash": source_ref["citation_claims_hash"],
        "render_policy_hash": sha256_json(
            {
                "renderer_policy_id": MARKDOWN_RENDERER_POLICY_ID,
                "target_formats": target_formats,
                "render_stage_sequence": render_stage_sequence_for_formats(
                    target_formats
                ),
                "template_ref": artifact_record["template_ref"],
            }
        ),
        "artifact_content_hash": artifact_content_hash,
        "rendered_formats": target_formats,
        "validation_snapshot": dict(source_ref["quality_summary"]),
        "created_at": now,
    }
    render_job = {
        "render_job_id": render_job_id,
        "artifact_id": artifact_record["artifact_id"],
        "artifact_version_id": artifact_version_id,
        "job_status": "COMPLETED",
        "current_stage": "FINALIZING",
        "progress_mode": "DETERMINATE",
        "progress_percent": 100,
        "retryable": False,
        "failure_code": None,
        "started_at": now,
        "completed_at": now,
    }
    rendered_payloads = build_rendered_payloads_from_markdown(markdown, target_formats)
    artifact_files = build_rendered_artifact_files_from_payloads(
        artifact_record=artifact_record,
        artifact_version=artifact_version,
        target_formats=target_formats,
        rendered_payloads=rendered_payloads,
    )
    artifact_links = build_artifact_links_for_files(
        artifact_files=artifact_files,
        created_by_actor_ref=artifact_record["owner_actor_ref"],
        created_at=now,
    )
    return {
        "render_request_id": render_request_id,
        "render_job": render_job,
        "artifact_version": artifact_version,
        "artifact_files": artifact_files,
        "artifact_links": artifact_links,
        "markdown": markdown,
        "rendered_payloads": rendered_payloads,
    }


def build_rendered_payloads_from_markdown(
    markdown: str,
    target_formats: list[str],
) -> dict[str, bytes]:
    payloads: dict[str, bytes] = {}
    if "MD" in target_formats:
        payloads["MD"] = markdown.encode("utf-8")
    if "HTML_PREVIEW" in target_formats:
        payloads["HTML_PREVIEW"] = render_html_preview_from_markdown(markdown).encode(
            "utf-8"
        )
    if "DOCX" in target_formats:
        payloads["DOCX"] = render_docx_export_from_markdown(markdown)
    if "PDF" in target_formats:
        payloads["PDF"] = render_pdf_export_from_markdown(markdown)
    return payloads


def build_rendered_artifact_files_from_payloads(
    *,
    artifact_record: dict[str, Any],
    artifact_version: dict[str, Any],
    target_formats: list[str],
    rendered_payloads: dict[str, bytes],
) -> list[dict[str, Any]]:
    artifact_files: list[dict[str, Any]] = []
    for target_format in target_formats:
        payload = rendered_payloads.get(target_format)
        if payload is None:
            raise ArtifactHandoffError(
                status_code=500,
                error_code="ae.render_payload_missing",
                detail=f"Rendered payload is not available for format: {target_format}",
            )
        artifact_files.append(
            build_rendered_artifact_file(
                artifact_record=artifact_record,
                artifact_version=artifact_version,
                target_format=target_format,
                payload=payload,
            )
        )
    return artifact_files


def build_markdown_artifact_files(
    *,
    artifact_record: dict[str, Any],
    artifact_version: dict[str, Any],
    markdown: str,
) -> list[dict[str, Any]]:
    return [
        build_rendered_artifact_file(
            artifact_record=artifact_record,
            artifact_version=artifact_version,
            target_format="MD",
            payload=markdown.encode("utf-8"),
        )
    ]


def build_html_preview_artifact_file(
    *,
    artifact_record: dict[str, Any],
    artifact_version: dict[str, Any],
    html_preview: str,
) -> dict[str, Any]:
    return build_rendered_artifact_file(
        artifact_record=artifact_record,
        artifact_version=artifact_version,
        target_format="HTML_PREVIEW",
        payload=html_preview.encode("utf-8"),
    )


def build_docx_export_artifact_file(
    *,
    artifact_record: dict[str, Any],
    artifact_version: dict[str, Any],
    docx_payload: bytes,
) -> dict[str, Any]:
    return build_rendered_artifact_file(
        artifact_record=artifact_record,
        artifact_version=artifact_version,
        target_format="DOCX",
        payload=docx_payload,
    )


def build_pdf_export_artifact_file(
    *,
    artifact_record: dict[str, Any],
    artifact_version: dict[str, Any],
    pdf_payload: bytes,
) -> dict[str, Any]:
    return build_rendered_artifact_file(
        artifact_record=artifact_record,
        artifact_version=artifact_version,
        target_format="PDF",
        payload=pdf_payload,
    )


def build_rendered_artifact_file(
    *,
    artifact_record: dict[str, Any],
    artifact_version: dict[str, Any],
    target_format: str,
    payload: bytes,
) -> dict[str, Any]:
    spec = artifact_format_spec(target_format)
    artifact_file_id = str(
        uuid5(
            NAMESPACE_URL,
            (
                "ae-artifact-file:"
                f"{artifact_version['artifact_version_id']}:{spec['format']}"
            ),
        )
    )
    file_name = artifact_file_name_for_format(
        artifact_record["display_title"],
        spec["format"],
    )
    return {
        "artifact_file_id": artifact_file_id,
        "artifact_version_id": artifact_version["artifact_version_id"],
        "format": spec["format"],
        "mime_type": spec["mime_type"],
        "file_name": file_name,
        "storage_ref": (
            "ae://artifacts/"
            f"{artifact_record['artifact_id']}/versions/"
            f"{artifact_version['artifact_version_id']}/{file_name}"
        ),
        "file_size_bytes": len(payload),
        "file_hash": sha256_bytes(payload),
        "source_version_hash": artifact_version["artifact_content_hash"],
        "created_at": artifact_version["created_at"],
    }


def artifact_format_spec(target_format: str) -> dict[str, Any]:
    spec = ARTIFACT_TRANSFORMER_CATALOG.get(target_format)
    if spec is None:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.render_format_unsupported",
            detail=f"Unsupported render format: {target_format}",
        )
    return dict(spec)


def artifact_mime_type(target_format: str) -> str:
    return str(artifact_format_spec(target_format)["mime_type"])


def artifact_file_extension(target_format: str) -> str:
    return str(artifact_format_spec(target_format)["extension"])


def artifact_content_kind(target_format: str) -> str:
    return str(artifact_format_spec(target_format)["content_kind"])


def render_stage_sequence_for_formats(target_formats: list[str]) -> list[str]:
    sequence = ["HANDOFF_VALIDATING", "MARKDOWN_RENDERING"]
    for stage in (
        artifact_format_spec(target_format)["render_stage"]
        for target_format in target_formats
    ):
        if stage not in sequence:
            sequence.append(str(stage))
    for stage in ("LINK_CREATING", "FINALIZING"):
        if stage not in sequence:
            sequence.append(stage)
    return [
        stage
        for stage in MULTI_FORMAT_RENDER_STAGE_ORDER
        if stage in set(sequence)
    ]


def artifact_file_name_for_format(title: str, target_format: str) -> str:
    return f"{safe_file_stem(title)}.{artifact_file_extension(target_format)}"


def build_artifact_links(
    *,
    artifact_file: dict[str, Any],
    created_by_actor_ref: dict[str, str],
    created_at: str,
) -> list[dict[str, Any]]:
    return [
        build_artifact_link(
            artifact_file=artifact_file,
            link_type="preview",
            created_by_actor_ref=created_by_actor_ref,
            created_at=created_at,
        ),
        build_artifact_link(
            artifact_file=artifact_file,
            link_type="download",
            created_by_actor_ref=created_by_actor_ref,
            created_at=created_at,
        ),
    ]


def build_artifact_links_for_files(
    *,
    artifact_files: list[dict[str, Any]],
    created_by_actor_ref: dict[str, str],
    created_at: str,
) -> list[dict[str, Any]]:
    links: list[dict[str, Any]] = []
    for artifact_file in artifact_files:
        links.extend(
            build_artifact_links(
                artifact_file=artifact_file,
                created_by_actor_ref=created_by_actor_ref,
                created_at=created_at,
            )
        )
    return links


def build_artifact_link(
    *,
    artifact_file: dict[str, Any],
    link_type: str,
    created_by_actor_ref: dict[str, str],
    created_at: str,
) -> dict[str, Any]:
    artifact_file_id = artifact_file["artifact_file_id"]
    return {
        "artifact_link_id": str(
            uuid5(NAMESPACE_URL, f"ae-artifact-link:{artifact_file_id}:{link_type}")
        ),
        "artifact_file_id": artifact_file_id,
        "link_type": link_type,
        "access_policy": "owner_only",
        "link_route": f"/api/v1/artifact-files/{artifact_file_id}/{link_type}",
        "expires_at": None,
        "created_by_actor_ref": dict(created_by_actor_ref),
        "download_count": 0,
        "revoked_at": None,
    }


def resolve_artifact_file_payload(
    store: ArtifactRecordStore,
    *,
    artifact_file_id: str,
    link_type: str,
) -> tuple[dict[str, Any], dict[str, Any], str]:
    artifact_file, artifact_link, payload = resolve_rendered_artifact_file_payload(
        store,
        artifact_file_id=artifact_file_id,
        link_type=link_type,
    )
    rendered_text = rendered_text_from_payload(artifact_file, payload)
    return artifact_file, artifact_link, rendered_text


def resolve_rendered_artifact_file_payload(
    store: ArtifactRecordStore,
    *,
    artifact_file_id: str,
    link_type: str,
) -> tuple[dict[str, Any], dict[str, Any], bytes]:
    artifact_file = store.get_file(artifact_file_id)
    if artifact_file is None:
        raise ArtifactHandoffError(
            status_code=404,
            error_code="ae.artifact_file_not_found",
            detail=f"Artifact file was not found: {artifact_file_id}",
        )
    artifact_link = store.get_file_link(artifact_file_id, link_type)
    if artifact_link is None:
        raise ArtifactHandoffError(
            status_code=404,
            error_code="ae.artifact_link_not_found",
            detail=f"Artifact {link_type} link was not found: {artifact_file_id}",
        )
    payload = store.get_rendered_artifact_file(artifact_file)
    if payload is None:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_file_not_ready",
            detail="Artifact file content is not ready.",
        )
    return artifact_file, artifact_link, payload


def rendered_payload_for_file(
    artifact_file: dict[str, Any],
    rendered_payloads: dict[str, bytes],
) -> bytes | None:
    return (
        rendered_payloads.get(artifact_file["artifact_file_id"])
        or rendered_payloads.get(artifact_file["format"])
    )


def rendered_text_from_payload(
    artifact_file: dict[str, Any],
    payload: bytes,
) -> str:
    if artifact_content_kind(artifact_file["format"]) != "text":
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_file_preview_unavailable",
            detail="Artifact file preview is available only for text formats.",
        )
    try:
        return payload.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_file_not_ready",
            detail="Artifact file text payload is not valid UTF-8.",
        ) from exc


def rendered_download_fields_from_payload(
    artifact_file: dict[str, Any],
    payload: bytes,
) -> dict[str, str]:
    if artifact_content_kind(artifact_file["format"]) == "text":
        return {"content": rendered_text_from_payload(artifact_file, payload)}
    return {
        "content_encoding": "base64",
        "content_base64": base64.b64encode(payload).decode("ascii"),
    }


def safe_file_stem(value: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip().lower())
    return normalized.strip(".-_")[:64] or "artifact"


def render_markdown_from_structured_draft(structured_draft: dict[str, Any]) -> str:
    lines: list[str] = [f"# {structured_draft['title'].strip()}"]
    summary = optional_text(structured_draft.get("summary"))
    if summary:
        lines.extend(["", summary])

    for section in sorted(
        structured_draft.get("sections", []),
        key=lambda item: int(item.get("ordinal") or 0),
    ):
        heading = optional_text(section.get("heading"))
        if heading:
            lines.extend(["", f"## {heading}"])
        for block in section.get("blocks", []):
            if block.get("block_type") != "paragraph":
                continue
            text_preview = optional_text(block.get("text_preview"))
            if text_preview:
                lines.extend(["", text_preview])

    valid_citations = [
        citation
        for citation in structured_draft.get("citations", [])
        if citation.get("valid") and optional_text(citation.get("citation_label"))
    ]
    if valid_citations:
        lines.extend(["", "## Citations"])
        for citation in valid_citations:
            evidence_id = optional_text(citation.get("evidence_id")) or "unknown"
            retrieval_package_id = (
                optional_text(citation.get("retrieval_package_id")) or "unknown"
            )
            lines.append(
                "- "
                f"{citation['citation_label']} evidence `{evidence_id}` "
                f"from retrieval package `{retrieval_package_id}`."
            )

    return "\n".join(lines).strip() + "\n"


def render_html_preview_from_markdown(markdown: str) -> str:
    body_lines: list[str] = []
    list_open = False
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            if list_open:
                body_lines.append("</ul>")
                list_open = False
            continue
        if line.startswith("# "):
            if list_open:
                body_lines.append("</ul>")
                list_open = False
            body_lines.append(f"<h1>{html.escape(line[2:].strip())}</h1>")
        elif line.startswith("## "):
            if list_open:
                body_lines.append("</ul>")
                list_open = False
            body_lines.append(f"<h2>{html.escape(line[3:].strip())}</h2>")
        elif line.startswith("- "):
            if not list_open:
                body_lines.append("<ul>")
                list_open = True
            body_lines.append(f"<li>{html.escape(line[2:].strip())}</li>")
        else:
            if list_open:
                body_lines.append("</ul>")
                list_open = False
            body_lines.append(f"<p>{html.escape(line)}</p>")
    if list_open:
        body_lines.append("</ul>")
    return (
        "<!doctype html>\n"
        '<html lang="ko">\n'
        "<head>\n"
        '<meta charset="utf-8" />\n'
        "<title>AE Artifact Preview</title>\n"
        "</head>\n"
        "<body>\n"
        '<article class="ae-artifact-preview">\n'
        f"{chr(10).join(body_lines)}\n"
        "</article>\n"
        "</body>\n"
        "</html>\n"
    )


def render_docx_export_from_markdown(markdown: str) -> bytes:
    document = Document()
    emitted = False
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        emitted = True
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)
    if not emitted:
        document.add_paragraph("")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf_export_from_markdown(markdown: str) -> bytes:
    lines = pdf_lines_from_markdown(markdown)
    pages = [
        lines[index : index + 44]
        for index in range(0, len(lines), 44)
    ] or [[""]]
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    page_object_ids: list[int] = []
    for page_lines in pages:
        page_object_id = len(objects) + 1
        content_object_id = page_object_id + 1
        page_object_ids.append(page_object_id)
        stream = pdf_content_stream(page_lines)
        objects.append(
            (
                "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
                "/Resources << /Font << /F1 3 0 R >> >> "
                f"/Contents {content_object_id} 0 R >>"
            ).encode("ascii")
        )
        objects.append(
            (
                f"<< /Length {len(stream)} >>\n"
            ).encode("ascii")
            + b"stream\n"
            + stream
            + b"\nendstream"
        )
    kids = " ".join(f"{object_id} 0 R" for object_id in page_object_ids)
    objects[1] = (
        f"<< /Type /Pages /Kids [{kids}] /Count {len(page_object_ids)} >>"
    ).encode("ascii")
    return build_pdf_document(objects)


def pdf_lines_from_markdown(markdown: str) -> list[str]:
    lines: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            normalized = line[2:].strip()
        elif line.startswith("## "):
            normalized = line[3:].strip()
        elif line.startswith("- "):
            normalized = f"- {line[2:].strip()}"
        else:
            normalized = line
        wrapped = textwrap.wrap(normalized, width=88) or [""]
        lines.extend(wrapped)
    return lines or [""]


def pdf_content_stream(lines: list[str]) -> bytes:
    commands = ["BT", "/F1 11 Tf", "50 800 Td"]
    first = True
    for line in lines:
        if first:
            first = False
        else:
            commands.append("0 -16 Td")
        commands.append(f"({pdf_literal_text(line)}) Tj")
    commands.append("ET")
    return "\n".join(commands).encode("cp1252")


def pdf_literal_text(value: str) -> str:
    safe_text = value.encode("cp1252", errors="replace").decode("cp1252")
    return (
        safe_text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\r", " ")
        .replace("\n", " ")
    )


def build_pdf_document(objects: list[bytes]) -> bytes:
    pdf = b"%PDF-1.4\n"
    offsets: list[int] = []
    for object_number, body in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += (
            f"{object_number} 0 obj\n".encode("ascii")
            + body
            + b"\nendobj\n"
        )
    startxref = len(pdf)
    xref_entries = b"".join(
        f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets
    )
    return (
        pdf
        + f"xref\n0 {len(objects) + 1}\n".encode("ascii")
        + b"0000000000 65535 f \n"
        + xref_entries
        + f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\n".encode(
            "ascii"
        )
        + b"startxref\n"
        + str(startxref).encode("ascii")
        + b"\n%%EOF\n"
    )


def validate_structured_draft_for_markdown_render(
    artifact_record: dict[str, Any],
    structured_draft: dict[str, Any],
) -> None:
    if artifact_record["artifact_status"] in {"ARCHIVED", "DELETED"}:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_not_renderable",
            detail="Archived or deleted artifacts cannot be rendered.",
        )
    if structured_draft.get("status") != "VALIDATED":
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.citation_validation_required",
            detail="Structured draft must be validated before rendering.",
        )
    if structured_draft.get("validation", {}).get("citation_status") != "VALIDATED":
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.citation_validation_required",
            detail="Citation validation must pass before rendering.",
        )
    source_ref = artifact_record["source_refs"][0]
    if structured_draft.get("structured_draft_id") != source_ref["structured_draft_id"]:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.source_draft_hash_mismatch",
            detail="Structured draft ID does not match the artifact source ref.",
        )
    if structured_draft.get("content_hash") != source_ref["structured_draft_content_hash"]:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.source_draft_hash_mismatch",
            detail="Structured draft content hash does not match the artifact source ref.",
        )


def markdown_target_formats_from_payload(
    payload: dict[str, Any],
    artifact_record: dict[str, Any],
) -> list[str]:
    requested_formats = (
        payload["target_formats"] if "target_formats" in payload else ["MD"]
    )
    if not isinstance(requested_formats, list) or not requested_formats:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.target_formats_invalid",
            detail="target_formats must be a non-empty list.",
        )
    normalized: list[str] = []
    for value in requested_formats:
        if value != "MD":
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.render_format_unsupported",
                detail="Slice 0042 supports Markdown rendering only.",
            )
        if value not in normalized:
            normalized.append(value)
    if "MD" not in artifact_record["target_formats"]:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.render_format_not_requested",
            detail="The artifact handoff did not request Markdown output.",
        )
    return normalized


def render_target_formats_from_payload(
    payload: dict[str, Any],
    artifact_record: dict[str, Any],
) -> list[str]:
    requested_formats = (
        payload["target_formats"] if "target_formats" in payload else ["MD"]
    )
    if not isinstance(requested_formats, list) or not requested_formats:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.target_formats_invalid",
            detail="target_formats must be a non-empty list.",
        )
    normalized: list[str] = []
    for value in requested_formats:
        if not isinstance(value, str) or value not in SUPPORTED_TARGET_FORMATS:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.render_format_unsupported",
                detail=f"Unsupported render format: {value}",
            )
        if value not in IMPLEMENTED_RENDER_FORMATS:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.render_format_unsupported",
                detail=f"Render format is not implemented yet: {value}",
            )
        if value not in normalized:
            normalized.append(value)
    missing_from_handoff = [
        value for value in normalized if value not in artifact_record["target_formats"]
    ]
    if missing_from_handoff:
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.render_format_not_requested",
            detail="The artifact handoff did not request this output format.",
        )
    return normalized


def deterministic_render_job_id(artifact_id: str, render_request_id: str) -> str:
    return str(uuid5(NAMESPACE_URL, f"ae-render-job:{artifact_id}:{render_request_id}"))


def validate_artifact_handoff_record(handoff_record: dict[str, Any]) -> None:
    for field_name in (
        "handoff_schema_version",
        "artifact_handoff_id",
        "artifact_request_id",
        "handoff_status",
        "chat_document_id",
        "interaction_id",
        "actor_claims_ref",
        "workspace_ref",
        "cx_generation_id",
        "structured_draft_id",
        "draft_schema_version",
        "structured_draft_content_hash",
        "citation_claims_hash",
        "validation_result_hash",
        "artifact_intent",
        "target_formats",
        "artifact_title",
        "language",
        "retention_policy_ref",
        "quality_summary",
    ):
        if field_name not in handoff_record:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.artifact_handoff_invalid",
                detail=f"Artifact handoff is missing {field_name}.",
            )
    if handoff_record["handoff_status"] != "READY_FOR_RENDERING":
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.artifact_handoff_not_ready",
            detail="Artifact handoff must be READY_FOR_RENDERING.",
        )


def validate_source_generation(
    generation_record: dict[str, Any],
    structured_draft: dict[str, Any],
) -> None:
    if generation_record.get("status") != "COMPLETED":
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.source_generation_not_ready",
            detail="CX generation must be COMPLETED before artifact handoff.",
        )
    if structured_draft.get("status") != "VALIDATED":
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.citation_validation_required",
            detail="Structured draft must be validated before artifact handoff.",
        )
    if structured_draft.get("validation", {}).get("citation_status") != "VALIDATED":
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.citation_validation_required",
            detail="Citation validation must pass before artifact handoff.",
        )
    if structured_draft.get("cx_generation_id") != generation_record.get("cx_generation_id"):
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.source_draft_hash_mismatch",
            detail="Structured draft does not belong to the CX generation record.",
        )
    expected_draft_id = generation_record.get("request_metadata", {}).get(
        "structured_draft_id"
    )
    if expected_draft_id and expected_draft_id != structured_draft.get("structured_draft_id"):
        raise ArtifactHandoffError(
            status_code=409,
            error_code="ae.source_draft_hash_mismatch",
            detail="Structured draft ID does not match the CX generation record.",
        )


def actor_claims_ref_from_payload(payload: dict[str, Any]) -> dict[str, str]:
    actor_ref = payload.get("actor_claims_ref")
    if actor_ref is not None and not isinstance(actor_ref, dict):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.actor_claims_ref_invalid",
            detail="actor_claims_ref must be an object when supplied.",
        )
    source = actor_ref or payload
    return {
        "actor_type": optional_text(source.get("actor_type")) or "user",
        "actor_id": optional_text(source.get("actor_id"))
        or optional_text(source.get("owner_user_id"))
        or DEFAULT_OWNER_USER_ID,
        "tenant_id": optional_text(source.get("tenant_id")) or DEFAULT_TENANT_ID,
    }


def workspace_ref_from_payload(payload: dict[str, Any]) -> dict[str, str]:
    tenant_id = optional_text(payload.get("tenant_id")) or DEFAULT_TENANT_ID
    owner_user_id = optional_text(payload.get("owner_user_id")) or DEFAULT_OWNER_USER_ID
    workspace_id = optional_text(payload.get("workspace_id")) or str(
        uuid5(NAMESPACE_URL, f"ae-workspace:{tenant_id}:{owner_user_id}:default")
    )
    return {
        "workspace_id": workspace_id,
        "tenant_id": tenant_id,
    }


def artifact_intent_from_payload(payload: dict[str, Any]) -> str:
    artifact_intent = optional_text(payload.get("artifact_intent")) or "create_artifact"
    if artifact_intent not in SUPPORTED_ARTIFACT_INTENTS:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_intent_invalid",
            detail=f"Unsupported artifact intent: {artifact_intent}",
        )
    return artifact_intent


def artifact_type_from_payload(payload: dict[str, Any]) -> str:
    artifact_type = optional_text(payload.get("artifact_type")) or DEFAULT_ARTIFACT_TYPE
    if artifact_type not in SUPPORTED_ARTIFACT_TYPES:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_type_invalid",
            detail=f"Unsupported artifact type: {artifact_type}",
        )
    return artifact_type


def target_formats_from_payload(payload: dict[str, Any]) -> list[str]:
    formats = payload.get("target_formats", DEFAULT_TARGET_FORMATS)
    if not isinstance(formats, list) or not formats:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.target_formats_invalid",
            detail="target_formats must be a non-empty list.",
        )
    normalized: list[str] = []
    for value in formats:
        if not isinstance(value, str) or value not in SUPPORTED_TARGET_FORMATS:
            raise ArtifactHandoffError(
                status_code=422,
                error_code="ae.render_format_unsupported",
                detail=f"Unsupported render format: {value}",
            )
        if value not in normalized:
            normalized.append(value)
    return normalized


def artifact_title_from_payload(
    payload: dict[str, Any],
    structured_draft: dict[str, Any],
) -> str:
    return optional_text(payload.get("artifact_title")) or structured_draft["title"]


def language_from_payload(payload: dict[str, Any]) -> str:
    language = optional_text(payload.get("language")) or "ko"
    if language not in {"ko", "en"}:
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.language_invalid",
            detail="language must be ko or en.",
        )
    return language


def required_string(payload: dict[str, Any], field_name: str, error_code: str) -> str:
    value = payload.get(field_name)
    if not isinstance(value, str) or not value.strip():
        raise ArtifactHandoffError(
            status_code=422,
            error_code=error_code,
            detail=f"{field_name} is required.",
        )
    return value.strip()


def optional_text(value: Any) -> str | None:
    if isinstance(value, str) and value.strip():
        return value.strip()
    return None


def _boolean_from_payload(
    payload: dict[str, Any],
    field_name: str,
    *,
    default: bool,
) -> bool:
    if field_name not in payload or payload[field_name] is None:
        return default
    value = payload[field_name]
    if not isinstance(value, bool):
        raise ArtifactHandoffError(
            status_code=422,
            error_code=f"ae.artifact_retention_{field_name}_invalid",
            detail=f"{field_name} must be a boolean.",
        )
    return value


def sha256_json(value: dict[str, Any]) -> str:
    return hashlib.sha256(
        json.dumps(value, sort_keys=True, separators=(",", ":")).encode("utf-8")
    ).hexdigest()


def sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _authorize_ae_request(
    request: Request,
    authorization: str | None,
) -> JSONResponse | None:
    result = validate_authorization_header(
        authorization,
        expected_audience="nex-ae-api",
        required_scopes=[DEFAULT_SERVICE_SCOPE],
    )
    if result.ok:
        return None

    return problem_response(
        request,
        status_code=401,
        error_code=result.error_code or "SERVICE_CLAIM_INVALID",
        title="Authentication failed",
        detail=result.detail or "AE API requires a valid service claim.",
        type_uri="https://nex-platform.local/problems/authentication-failed",
    )


def _artifact_problem_response(
    request: Request,
    exc: ArtifactHandoffError,
) -> JSONResponse:
    return problem_response(
        request,
        status_code=exc.status_code,
        error_code=exc.error_code,
        title="Artifact handoff failed",
        detail=exc.detail,
        retryable=exc.retryable,
        type_uri="https://nex-platform.local/problems/artifact-handoff-failed",
    )


def _safe_response_json(response: httpx.Response) -> dict[str, Any]:
    try:
        payload = response.json()
    except json.JSONDecodeError:
        return {}
    if isinstance(payload, dict):
        return payload
    return {}


def _artifact_retention_execution_history_upsert_sql(dialect_name: str) -> str:
    json_exprs = _json_param_exprs(
        ARTIFACT_RETENTION_EXECUTION_HISTORY_JSON_FIELDS,
        dialect_name,
    )
    return f"""
        INSERT INTO ae_artifact_retention_executions (
            retention_execution_id,
            execution_history_schema_version,
            artifact_retention_execution_schema_version,
            policy_id,
            service_id,
            mode,
            execution_status,
            tenant_id,
            workspace_id,
            owner_user_id,
            retention_days_after_logical_purge,
            as_of,
            cutoff_at,
            checked_at,
            scan_limit,
            max_delete_count,
            candidate_count,
            selected_count,
            delete_enabled,
            storage_mutation_enabled,
            database_row_delete_enabled,
            deleted_counts,
            requested_by,
            idempotency_key,
            trace_id,
            request_id,
            blocked_reason,
            error,
            audit,
            metadata,
            execution,
            execution_payload_hash,
            created_at
        )
        VALUES (
            :retention_execution_id,
            :execution_history_schema_version,
            :artifact_retention_execution_schema_version,
            :policy_id,
            :service_id,
            :mode,
            :execution_status,
            :tenant_id,
            :workspace_id,
            :owner_user_id,
            :retention_days_after_logical_purge,
            :as_of,
            :cutoff_at,
            :checked_at,
            :scan_limit,
            :max_delete_count,
            :candidate_count,
            :selected_count,
            :delete_enabled,
            :storage_mutation_enabled,
            :database_row_delete_enabled,
            {json_exprs["deleted_counts"]},
            {json_exprs["requested_by"]},
            :idempotency_key,
            :trace_id,
            :request_id,
            :blocked_reason,
            {json_exprs["error"]},
            {json_exprs["audit"]},
            {json_exprs["metadata"]},
            {json_exprs["execution"]},
            :execution_payload_hash,
            :created_at
        )
        ON CONFLICT (retention_execution_id) DO UPDATE SET
            execution_status = excluded.execution_status,
            checked_at = excluded.checked_at,
            candidate_count = excluded.candidate_count,
            selected_count = excluded.selected_count,
            deleted_counts = excluded.deleted_counts,
            requested_by = excluded.requested_by,
            trace_id = excluded.trace_id,
            request_id = excluded.request_id,
            blocked_reason = excluded.blocked_reason,
            error = excluded.error,
            audit = excluded.audit,
            metadata = excluded.metadata,
            execution = excluded.execution,
            execution_payload_hash = excluded.execution_payload_hash,
            created_at = excluded.created_at
    """


def _artifact_retention_execution_history_select_sql(where_clause: str) -> str:
    return f"""
        SELECT
            retention_execution_id,
            execution_history_schema_version,
            artifact_retention_execution_schema_version,
            policy_id,
            service_id,
            mode,
            execution_status,
            tenant_id,
            workspace_id,
            owner_user_id,
            retention_days_after_logical_purge,
            as_of,
            cutoff_at,
            checked_at,
            scan_limit,
            max_delete_count,
            candidate_count,
            selected_count,
            delete_enabled,
            storage_mutation_enabled,
            database_row_delete_enabled,
            deleted_counts,
            requested_by,
            idempotency_key,
            trace_id,
            request_id,
            blocked_reason,
            error,
            audit,
            metadata,
            execution,
            execution_payload_hash,
            created_at
        FROM ae_artifact_retention_executions
        WHERE {where_clause}
    """


def _artifact_retention_execution_history_list_sql(
    history_filter: dict[str, Any],
) -> str:
    where_clauses = [
        "tenant_id = :tenant_id",
        "workspace_id = :workspace_id",
        "owner_user_id = :owner_user_id",
    ]
    if history_filter.get("mode") is not None:
        where_clauses.append("mode = :mode")
    if history_filter.get("execution_status") is not None:
        where_clauses.append("execution_status = :execution_status")
    return (
        _artifact_retention_execution_history_select_sql(
            " AND ".join(where_clauses)
        )
        + """
        ORDER BY checked_at DESC, retention_execution_id DESC
        LIMIT :limit
        """
    )


def _artifact_retention_execution_history_params(
    record: dict[str, Any],
) -> dict[str, Any]:
    validate_artifact_retention_execution_history_record(record)
    params = dict(record)
    for field_name in ARTIFACT_RETENTION_EXECUTION_HISTORY_JSON_FIELDS:
        value = params[field_name]
        params[field_name] = None if value is None else json.dumps(value)
    return params


def _artifact_retention_execution_history_from_row(row: Any) -> dict[str, Any]:
    data = dict(row)
    record = {
        "retention_execution_id": data["retention_execution_id"],
        "execution_history_schema_version": data["execution_history_schema_version"],
        "artifact_retention_execution_schema_version": data[
            "artifact_retention_execution_schema_version"
        ],
        "policy_id": data["policy_id"],
        "service_id": data["service_id"],
        "mode": data["mode"],
        "execution_status": data["execution_status"],
        "tenant_id": data["tenant_id"],
        "workspace_id": data["workspace_id"],
        "owner_user_id": data["owner_user_id"],
        "retention_days_after_logical_purge": data[
            "retention_days_after_logical_purge"
        ],
        "as_of": _datetime_value(data["as_of"]),
        "cutoff_at": _datetime_value(data["cutoff_at"]),
        "checked_at": _datetime_value(data["checked_at"]),
        "scan_limit": data["scan_limit"],
        "max_delete_count": data["max_delete_count"],
        "candidate_count": data["candidate_count"],
        "selected_count": data["selected_count"],
        "delete_enabled": bool(data["delete_enabled"]),
        "storage_mutation_enabled": bool(data["storage_mutation_enabled"]),
        "database_row_delete_enabled": bool(data["database_row_delete_enabled"]),
        "deleted_counts": _json_value(data["deleted_counts"], {}),
        "requested_by": _json_value(data["requested_by"], {}),
        "idempotency_key": data["idempotency_key"],
        "trace_id": data["trace_id"],
        "request_id": data["request_id"],
        "blocked_reason": data["blocked_reason"],
        "error": _json_value(data["error"], None),
        "audit": _json_value(data["audit"], {}),
        "metadata": _json_value(data["metadata"], {}),
        "execution": _json_value(data["execution"], {}),
        "execution_payload_hash": data["execution_payload_hash"],
        "created_at": _datetime_value(data["created_at"]),
    }
    return validate_artifact_retention_execution_history_record(record)


def _artifact_retention_execution_history_filter(
    *,
    tenant_id: str,
    workspace_id: str,
    owner_user_id: str,
    mode: str | None,
    execution_status: str | None,
    limit: int | str | None,
) -> dict[str, Any]:
    return {
        "tenant_id": _required_collection_scope_text(tenant_id, "tenant_id"),
        "workspace_id": _required_collection_scope_text(
            workspace_id,
            "workspace_id",
        ),
        "owner_user_id": _required_collection_scope_text(
            owner_user_id,
            "owner_user_id",
        ),
        "mode": _normalize_artifact_retention_execution_mode(mode)
        if mode is not None
        else None,
        "execution_status": _normalize_artifact_retention_execution_status(
            execution_status
        )
        if execution_status is not None
        else None,
        "limit": normalize_artifact_collection_limit(limit),
    }


def _artifact_retention_execution_history_filter_params(
    history_filter: dict[str, Any],
) -> dict[str, Any]:
    params = {
        "tenant_id": history_filter["tenant_id"],
        "workspace_id": history_filter["workspace_id"],
        "owner_user_id": history_filter["owner_user_id"],
        "limit": history_filter["limit"],
    }
    if history_filter.get("mode") is not None:
        params["mode"] = history_filter["mode"]
    if history_filter.get("execution_status") is not None:
        params["execution_status"] = history_filter["execution_status"]
    return params


def _artifact_retention_execution_history_matches(
    record: dict[str, Any],
    history_filter: dict[str, Any],
) -> bool:
    return (
        record.get("tenant_id") == history_filter["tenant_id"]
        and record.get("workspace_id") == history_filter["workspace_id"]
        and record.get("owner_user_id") == history_filter["owner_user_id"]
        and (
            history_filter["mode"] is None
            or record.get("mode") == history_filter["mode"]
        )
        and (
            history_filter["execution_status"] is None
            or record.get("execution_status") == history_filter["execution_status"]
        )
    )


def _artifact_handoff_upsert_sql(dialect_name: str) -> str:
    json_exprs = _json_param_exprs(ARTIFACT_HANDOFF_JSON_FIELDS, dialect_name)
    return f"""
        INSERT INTO ae_artifact_handoffs (
            artifact_handoff_id,
            handoff_schema_version,
            artifact_request_id,
            handoff_status,
            trace_id,
            request_id,
            tenant_id,
            workspace_id,
            owner_user_id,
            chat_document_id,
            interaction_id,
            cx_generation_id,
            structured_draft_id,
            draft_schema_version,
            structured_draft_content_hash,
            citation_claims_hash,
            validation_result_hash,
            template_id,
            template_version,
            rendering_template_id,
            artifact_intent,
            target_formats,
            artifact_title,
            language,
            retention_policy_ref,
            actor_claims_ref,
            workspace_ref,
            quality_summary,
            created_at,
            updated_at
        )
        VALUES (
            :artifact_handoff_id,
            :handoff_schema_version,
            :artifact_request_id,
            :handoff_status,
            :trace_id,
            :request_id,
            :tenant_id,
            :workspace_id,
            :owner_user_id,
            :chat_document_id,
            :interaction_id,
            :cx_generation_id,
            :structured_draft_id,
            :draft_schema_version,
            :structured_draft_content_hash,
            :citation_claims_hash,
            :validation_result_hash,
            :template_id,
            :template_version,
            :rendering_template_id,
            :artifact_intent,
            {json_exprs["target_formats"]},
            :artifact_title,
            :language,
            :retention_policy_ref,
            {json_exprs["actor_claims_ref"]},
            {json_exprs["workspace_ref"]},
            {json_exprs["quality_summary"]},
            :created_at,
            :updated_at
        )
        ON CONFLICT (artifact_handoff_id) DO UPDATE SET
            artifact_request_id = excluded.artifact_request_id,
            handoff_status = excluded.handoff_status,
            target_formats = excluded.target_formats,
            artifact_title = excluded.artifact_title,
            retention_policy_ref = excluded.retention_policy_ref,
            actor_claims_ref = excluded.actor_claims_ref,
            workspace_ref = excluded.workspace_ref,
            quality_summary = excluded.quality_summary,
            updated_at = excluded.updated_at
    """


def _artifact_handoff_select_sql() -> str:
    return """
        SELECT
            artifact_handoff_id,
            handoff_schema_version,
            artifact_request_id,
            handoff_status,
            trace_id,
            request_id,
            chat_document_id,
            interaction_id,
            actor_claims_ref,
            workspace_ref,
            cx_generation_id,
            structured_draft_id,
            draft_schema_version,
            structured_draft_content_hash,
            citation_claims_hash,
            validation_result_hash,
            template_id,
            template_version,
            rendering_template_id,
            artifact_intent,
            target_formats,
            artifact_title,
            language,
            retention_policy_ref,
            quality_summary,
            created_at,
            updated_at
        FROM ae_artifact_handoffs
        WHERE artifact_handoff_id = :artifact_handoff_id
    """


def _artifact_handoff_params(record: dict[str, Any]) -> dict[str, Any]:
    actor = dict(record["actor_claims_ref"])
    workspace = dict(record["workspace_ref"])
    params = dict(record)
    params["tenant_id"] = actor["tenant_id"]
    params["workspace_id"] = workspace["workspace_id"]
    params["owner_user_id"] = actor["actor_id"]
    return _json_params(params, ARTIFACT_HANDOFF_JSON_FIELDS)


def _artifact_handoff_from_row(row: Any) -> dict[str, Any]:
    data = dict(row)
    return {
        "handoff_schema_version": data["handoff_schema_version"],
        "artifact_handoff_id": data["artifact_handoff_id"],
        "artifact_request_id": data["artifact_request_id"],
        "handoff_status": data["handoff_status"],
        "trace_id": data["trace_id"],
        "request_id": data["request_id"],
        "chat_document_id": data["chat_document_id"],
        "interaction_id": data["interaction_id"],
        "actor_claims_ref": _json_value(data["actor_claims_ref"], {}),
        "workspace_ref": _json_value(data["workspace_ref"], {}),
        "cx_generation_id": data["cx_generation_id"],
        "structured_draft_id": data["structured_draft_id"],
        "draft_schema_version": data["draft_schema_version"],
        "structured_draft_content_hash": data["structured_draft_content_hash"],
        "citation_claims_hash": data["citation_claims_hash"],
        "validation_result_hash": data["validation_result_hash"],
        "template_id": data["template_id"],
        "template_version": data["template_version"],
        "rendering_template_id": data["rendering_template_id"],
        "artifact_intent": data["artifact_intent"],
        "target_formats": _json_value(data["target_formats"], []),
        "artifact_title": data["artifact_title"],
        "language": data["language"],
        "retention_policy_ref": data["retention_policy_ref"],
        "quality_summary": _json_value(data["quality_summary"], {}),
        "created_at": _datetime_value(data["created_at"]),
        "updated_at": _datetime_value(data["updated_at"]),
    }


def _persist_artifact_record(session: Session, record: dict[str, Any]) -> None:
    dialect_name = _dialect_name(session)
    session.execute(
        text(_artifact_upsert_sql(dialect_name)),
        _artifact_params(record),
    )
    for source_ref in record.get("source_refs", []):
        session.execute(
            text(_artifact_source_ref_upsert_sql(dialect_name)),
            _artifact_source_ref_params(record["artifact_id"], source_ref),
        )
    for version in record.get("versions", []):
        session.execute(
            text(_artifact_version_upsert_sql(dialect_name)),
            _artifact_version_params(version),
        )
    for render_job in record.get("render_jobs", []):
        session.execute(
            text(_artifact_render_job_upsert_sql()),
            _artifact_render_job_params(render_job),
        )
    for artifact_file in record.get("files", []):
        session.execute(
            text(_artifact_file_upsert_sql()),
            _artifact_file_params(record["artifact_id"], artifact_file),
        )
    for artifact_link in record.get("links", []):
        session.execute(
            text(_artifact_link_upsert_sql(dialect_name)),
            _artifact_link_params(artifact_link),
        )


def _artifact_upsert_sql(dialect_name: str) -> str:
    json_exprs = _json_param_exprs(ARTIFACT_RECORD_JSON_FIELDS, dialect_name)
    return f"""
        INSERT INTO ae_artifacts (
            artifact_id,
            artifact_schema_version,
            artifact_type,
            artifact_status,
            current_version_id,
            artifact_handoff_id,
            artifact_request_id,
            tenant_id,
            workspace_id,
            owner_user_id,
            chat_document_id,
            interaction_id,
            trace_id,
            request_id,
            display_title,
            language,
            artifact_intent,
            target_formats,
            retention_policy_ref,
            owner_actor_ref,
            workspace_ref,
            template_ref,
            handoff_ref,
            created_at,
            updated_at
        )
        VALUES (
            :artifact_id,
            :artifact_schema_version,
            :artifact_type,
            :artifact_status,
            :current_version_id,
            :artifact_handoff_id,
            :artifact_request_id,
            :tenant_id,
            :workspace_id,
            :owner_user_id,
            :chat_document_id,
            :interaction_id,
            :trace_id,
            :request_id,
            :display_title,
            :language,
            :artifact_intent,
            {json_exprs["target_formats"]},
            :retention_policy_ref,
            {json_exprs["owner_actor_ref"]},
            {json_exprs["workspace_ref"]},
            {json_exprs["template_ref"]},
            {json_exprs["handoff_ref"]},
            :created_at,
            :updated_at
        )
        ON CONFLICT (artifact_id) DO UPDATE SET
            artifact_type = excluded.artifact_type,
            artifact_status = excluded.artifact_status,
            current_version_id = excluded.current_version_id,
            display_title = excluded.display_title,
            target_formats = excluded.target_formats,
            retention_policy_ref = excluded.retention_policy_ref,
            owner_actor_ref = excluded.owner_actor_ref,
            workspace_ref = excluded.workspace_ref,
            template_ref = excluded.template_ref,
            handoff_ref = excluded.handoff_ref,
            updated_at = excluded.updated_at
    """


def _artifact_params(record: dict[str, Any]) -> dict[str, Any]:
    owner = dict(record["owner_actor_ref"])
    workspace = dict(record["workspace_ref"])
    handoff_ref = dict(record["handoff_ref"])
    params = dict(record)
    params["artifact_handoff_id"] = handoff_ref["artifact_handoff_id"]
    params["artifact_request_id"] = handoff_ref["artifact_request_id"]
    params["tenant_id"] = owner["tenant_id"]
    params["workspace_id"] = workspace["workspace_id"]
    params["owner_user_id"] = owner["actor_id"]
    return _json_params(params, ARTIFACT_RECORD_JSON_FIELDS)


def _load_artifact_record(session: Session, artifact_id: str) -> dict[str, Any] | None:
    artifact_row = (
        session.execute(
            text(_artifact_select_sql("artifact_id = :artifact_id")),
            {"artifact_id": artifact_id},
        )
        .mappings()
        .first()
    )
    if artifact_row is None:
        return None
    return _artifact_record_from_row(
        artifact_row,
        source_refs=[
            _artifact_source_ref_from_row(row)
            for row in session.execute(
                text(
                    _artifact_source_ref_select_sql(
                        "artifact_id = :artifact_id ORDER BY source_ref_id ASC"
                    )
                ),
                {"artifact_id": artifact_id},
            )
            .mappings()
            .all()
        ],
        versions=[
            _artifact_version_from_row(row)
            for row in session.execute(
                text(
                    _artifact_version_select_sql(
                        "artifact_id = :artifact_id ORDER BY version_no ASC"
                    )
                ),
                {"artifact_id": artifact_id},
            )
            .mappings()
            .all()
        ],
        render_jobs=[
            _artifact_render_job_from_row(row)
            for row in session.execute(
                text(
                    _artifact_render_job_select_sql(
                        "artifact_id = :artifact_id "
                        "ORDER BY created_at ASC, render_job_id ASC"
                    )
                ),
                {"artifact_id": artifact_id},
            )
            .mappings()
            .all()
        ],
        files=[
            _artifact_file_from_row(row)
            for row in session.execute(
                text(
                    _artifact_file_select_sql(
                        "artifact_id = :artifact_id "
                        "ORDER BY created_at ASC, artifact_file_id ASC"
                    )
                ),
                {"artifact_id": artifact_id},
            )
            .mappings()
            .all()
        ],
        links=[
            _artifact_link_from_row(row)
            for row in session.execute(
                text(
                    _artifact_link_select_sql(
                        "artifact_file_id IN ("
                        "SELECT artifact_file_id FROM ae_artifact_files "
                        "WHERE artifact_id = :artifact_id"
                        ") ORDER BY CASE link_type "
                        "WHEN 'preview' THEN 0 ELSE 1 END, artifact_link_id ASC"
                    )
                ),
                {"artifact_id": artifact_id},
            )
            .mappings()
            .all()
        ],
    )


def _load_artifact_records_for_collection(
    session: Session,
    collection_filter: dict[str, Any],
) -> list[dict[str, Any]]:
    rows = (
        session.execute(
            text(_artifact_collection_select_ids_sql(collection_filter)),
            _artifact_collection_params(collection_filter),
        )
        .mappings()
        .all()
    )
    records = []
    for row in rows:
        record = _load_artifact_record(session, row["artifact_id"])
        if record is not None:
            records.append(record)
    return records


def _load_artifact_records_for_retention_candidates(
    session: Session,
    candidate_filter: dict[str, Any],
) -> list[dict[str, Any]]:
    rows = (
        session.execute(
            text(_artifact_retention_candidate_select_ids_sql()),
            _artifact_retention_candidate_params(candidate_filter),
        )
        .mappings()
        .all()
    )
    records = []
    for row in rows:
        record = _load_artifact_record(session, row["artifact_id"])
        if record is not None:
            records.append(record)
    return records


def _delete_retention_selected_artifact_records(
    session: Session,
    records: list[dict[str, Any]],
    *,
    rendered_storage: RenderedArtifactStorage,
) -> dict[str, int]:
    counts = _empty_artifact_retention_deleted_counts()
    for record in records:
        artifact_id = record["artifact_id"]
        counts["storage_files"] += _delete_rendered_files_for_retention(
            rendered_storage,
            record,
        )
        counts["links"] += _delete_artifact_rows(
            session,
            """
            DELETE FROM ae_artifact_links
            WHERE artifact_file_id IN (
                SELECT artifact_file_id
                FROM ae_artifact_files
                WHERE artifact_id = :artifact_id
            )
            """,
            artifact_id,
        )
        counts["files"] += _delete_artifact_rows(
            session,
            "DELETE FROM ae_artifact_files WHERE artifact_id = :artifact_id",
            artifact_id,
        )
        counts["render_jobs"] += _delete_artifact_rows(
            session,
            "DELETE FROM ae_artifact_render_jobs WHERE artifact_id = :artifact_id",
            artifact_id,
        )
        counts["versions"] += _delete_artifact_rows(
            session,
            "DELETE FROM ae_artifact_versions WHERE artifact_id = :artifact_id",
            artifact_id,
        )
        counts["source_refs"] += _delete_artifact_rows(
            session,
            "DELETE FROM ae_artifact_source_refs WHERE artifact_id = :artifact_id",
            artifact_id,
        )
        counts["artifacts"] += _delete_artifact_rows(
            session,
            "DELETE FROM ae_artifacts WHERE artifact_id = :artifact_id",
            artifact_id,
        )
    return counts


def _delete_rendered_files_for_retention(
    rendered_storage: RenderedArtifactStorage,
    record: dict[str, Any],
) -> int:
    deleted = 0
    for artifact_file in _list_of_mappings(record.get("files")):
        if rendered_storage.delete_rendered_artifact_file(artifact_file):
            deleted += 1
    return deleted


def _delete_artifact_rows(session: Session, sql: str, artifact_id: str) -> int:
    result = session.execute(text(sql), {"artifact_id": artifact_id})
    return int(result.rowcount or 0)


def _artifact_collection_select_ids_sql(collection_filter: dict[str, Any]) -> str:
    where_clauses = [
        "tenant_id = :tenant_id",
        "workspace_id = :workspace_id",
        "owner_user_id = :owner_user_id",
    ]
    if collection_filter.get("status") is not None:
        where_clauses.append("artifact_status = :status")
    where_sql = " AND ".join(where_clauses)
    return f"""
        SELECT artifact_id
        FROM ae_artifacts
        WHERE {where_sql}
        ORDER BY updated_at DESC, artifact_id ASC
        LIMIT :limit
    """


def _artifact_collection_params(collection_filter: dict[str, Any]) -> dict[str, Any]:
    params = {
        "tenant_id": collection_filter["tenant_id"],
        "workspace_id": collection_filter["workspace_id"],
        "owner_user_id": collection_filter["owner_user_id"],
        "limit": collection_filter["limit"],
    }
    if collection_filter.get("status") is not None:
        params["status"] = collection_filter["status"]
    return params


def _artifact_retention_candidate_select_ids_sql() -> str:
    return """
        SELECT artifact_id
        FROM ae_artifacts
        WHERE tenant_id = :tenant_id
          AND workspace_id = :workspace_id
          AND owner_user_id = :owner_user_id
          AND artifact_status = :status
          AND updated_at <= :cutoff_at
        ORDER BY updated_at ASC, artifact_id ASC
        LIMIT :limit
    """


def _artifact_retention_candidate_params(
    candidate_filter: dict[str, Any],
) -> dict[str, Any]:
    return {
        "tenant_id": candidate_filter["tenant_id"],
        "workspace_id": candidate_filter["workspace_id"],
        "owner_user_id": candidate_filter["owner_user_id"],
        "status": candidate_filter["status"],
        "cutoff_at": candidate_filter["cutoff_at"],
        "limit": candidate_filter["limit"],
    }


def _artifact_select_sql(where_clause: str) -> str:
    return f"""
        SELECT
            artifact_id,
            artifact_schema_version,
            artifact_type,
            artifact_status,
            current_version_id,
            trace_id,
            request_id,
            chat_document_id,
            interaction_id,
            owner_actor_ref,
            workspace_ref,
            display_title,
            language,
            artifact_intent,
            target_formats,
            retention_policy_ref,
            template_ref,
            handoff_ref,
            created_at,
            updated_at
        FROM ae_artifacts
        WHERE {where_clause}
    """


def _artifact_record_from_row(
    row: Any,
    *,
    source_refs: list[dict[str, Any]],
    versions: list[dict[str, Any]],
    render_jobs: list[dict[str, Any]],
    files: list[dict[str, Any]],
    links: list[dict[str, Any]],
) -> dict[str, Any]:
    data = dict(row)
    return {
        "artifact_schema_version": data["artifact_schema_version"],
        "artifact_id": data["artifact_id"],
        "artifact_type": data["artifact_type"],
        "artifact_status": data["artifact_status"],
        "current_version_id": data["current_version_id"],
        "trace_id": data["trace_id"],
        "request_id": data["request_id"],
        "chat_document_id": data["chat_document_id"],
        "interaction_id": data["interaction_id"],
        "owner_actor_ref": _json_value(data["owner_actor_ref"], {}),
        "workspace_ref": _json_value(data["workspace_ref"], {}),
        "display_title": data["display_title"],
        "language": data["language"],
        "artifact_intent": data["artifact_intent"],
        "target_formats": _json_value(data["target_formats"], []),
        "retention_policy_ref": data["retention_policy_ref"],
        "template_ref": _json_value(data["template_ref"], {}),
        "handoff_ref": _json_value(data["handoff_ref"], {}),
        "source_refs": source_refs,
        "versions": versions,
        "render_jobs": render_jobs,
        "files": files,
        "links": links,
        "created_at": _datetime_value(data["created_at"]),
        "updated_at": _datetime_value(data["updated_at"]),
    }


def _artifact_source_ref_upsert_sql(dialect_name: str) -> str:
    json_exprs = _json_param_exprs(ARTIFACT_SOURCE_REF_JSON_FIELDS, dialect_name)
    return f"""
        INSERT INTO ae_artifact_source_refs (
            source_ref_id,
            artifact_id,
            cx_generation_id,
            structured_draft_id,
            draft_schema_version,
            structured_draft_content_hash,
            citation_claims_hash,
            validation_result_hash,
            retrieval_package_id,
            retrieval_package_hash,
            evidence_ref_count,
            source_anchor_count,
            quality_summary,
            created_at
        )
        VALUES (
            :source_ref_id,
            :artifact_id,
            :cx_generation_id,
            :structured_draft_id,
            :draft_schema_version,
            :structured_draft_content_hash,
            :citation_claims_hash,
            :validation_result_hash,
            :retrieval_package_id,
            :retrieval_package_hash,
            :evidence_ref_count,
            :source_anchor_count,
            {json_exprs["quality_summary"]},
            :created_at
        )
        ON CONFLICT (source_ref_id) DO UPDATE SET
            retrieval_package_id = excluded.retrieval_package_id,
            retrieval_package_hash = excluded.retrieval_package_hash,
            evidence_ref_count = excluded.evidence_ref_count,
            source_anchor_count = excluded.source_anchor_count,
            quality_summary = excluded.quality_summary
    """


def _artifact_source_ref_params(
    artifact_id: str,
    source_ref: dict[str, Any],
) -> dict[str, Any]:
    params = dict(source_ref)
    params["artifact_id"] = artifact_id
    params["created_at"] = params.get("created_at") or _utc_now()
    return _json_params(params, ARTIFACT_SOURCE_REF_JSON_FIELDS)


def _artifact_source_ref_select_sql(where_clause: str) -> str:
    return f"""
        SELECT
            source_ref_id,
            cx_generation_id,
            structured_draft_id,
            draft_schema_version,
            structured_draft_content_hash,
            citation_claims_hash,
            validation_result_hash,
            retrieval_package_id,
            retrieval_package_hash,
            evidence_ref_count,
            source_anchor_count,
            quality_summary
        FROM ae_artifact_source_refs
        WHERE {where_clause}
    """


def _artifact_source_ref_from_row(row: Any) -> dict[str, Any]:
    data = dict(row)
    return {
        "source_ref_id": data["source_ref_id"],
        "cx_generation_id": data["cx_generation_id"],
        "structured_draft_id": data["structured_draft_id"],
        "draft_schema_version": data["draft_schema_version"],
        "structured_draft_content_hash": data["structured_draft_content_hash"],
        "citation_claims_hash": data["citation_claims_hash"],
        "validation_result_hash": data["validation_result_hash"],
        "retrieval_package_id": data["retrieval_package_id"],
        "retrieval_package_hash": data["retrieval_package_hash"],
        "evidence_ref_count": data["evidence_ref_count"],
        "source_anchor_count": data["source_anchor_count"],
        "quality_summary": _json_value(data["quality_summary"], {}),
    }


def _artifact_version_upsert_sql(dialect_name: str) -> str:
    json_exprs = _json_param_exprs(ARTIFACT_VERSION_JSON_FIELDS, dialect_name)
    return f"""
        INSERT INTO ae_artifact_versions (
            artifact_version_id,
            artifact_id,
            version_no,
            version_reason,
            source_generation_id,
            source_structured_draft_id,
            source_content_hash,
            source_citation_claims_hash,
            render_policy_hash,
            artifact_content_hash,
            rendered_formats,
            validation_snapshot,
            created_at
        )
        VALUES (
            :artifact_version_id,
            :artifact_id,
            :version_no,
            :version_reason,
            :source_generation_id,
            :source_structured_draft_id,
            :source_content_hash,
            :source_citation_claims_hash,
            :render_policy_hash,
            :artifact_content_hash,
            {json_exprs["rendered_formats"]},
            {json_exprs["validation_snapshot"]},
            :created_at
        )
        ON CONFLICT (artifact_version_id) DO UPDATE SET
            artifact_content_hash = excluded.artifact_content_hash,
            rendered_formats = excluded.rendered_formats,
            validation_snapshot = excluded.validation_snapshot
    """


def _artifact_version_params(version: dict[str, Any]) -> dict[str, Any]:
    return _json_params(dict(version), ARTIFACT_VERSION_JSON_FIELDS)


def _artifact_version_select_sql(where_clause: str) -> str:
    return f"""
        SELECT
            artifact_version_id,
            artifact_id,
            version_no,
            version_reason,
            source_generation_id,
            source_structured_draft_id,
            source_content_hash,
            source_citation_claims_hash,
            render_policy_hash,
            artifact_content_hash,
            rendered_formats,
            validation_snapshot,
            created_at
        FROM ae_artifact_versions
        WHERE {where_clause}
    """


def _artifact_version_from_row(row: Any) -> dict[str, Any]:
    data = dict(row)
    return {
        "artifact_version_id": data["artifact_version_id"],
        "artifact_id": data["artifact_id"],
        "version_no": data["version_no"],
        "version_reason": data["version_reason"],
        "source_generation_id": data["source_generation_id"],
        "source_structured_draft_id": data["source_structured_draft_id"],
        "source_content_hash": data["source_content_hash"],
        "source_citation_claims_hash": data["source_citation_claims_hash"],
        "render_policy_hash": data["render_policy_hash"],
        "artifact_content_hash": data["artifact_content_hash"],
        "rendered_formats": _json_value(data["rendered_formats"], []),
        "validation_snapshot": _json_value(data["validation_snapshot"], {}),
        "created_at": _datetime_value(data["created_at"]),
    }


def _artifact_render_job_upsert_sql() -> str:
    return """
        INSERT INTO ae_artifact_render_jobs (
            render_job_id,
            artifact_id,
            artifact_version_id,
            job_status,
            current_stage,
            progress_mode,
            progress_percent,
            retryable,
            failure_code,
            started_at,
            completed_at,
            created_at,
            updated_at
        )
        VALUES (
            :render_job_id,
            :artifact_id,
            :artifact_version_id,
            :job_status,
            :current_stage,
            :progress_mode,
            :progress_percent,
            :retryable,
            :failure_code,
            :started_at,
            :completed_at,
            :created_at,
            :updated_at
        )
        ON CONFLICT (render_job_id) DO UPDATE SET
            artifact_version_id = excluded.artifact_version_id,
            job_status = excluded.job_status,
            current_stage = excluded.current_stage,
            progress_percent = excluded.progress_percent,
            retryable = excluded.retryable,
            failure_code = excluded.failure_code,
            completed_at = excluded.completed_at,
            updated_at = excluded.updated_at
    """


def _artifact_render_job_params(render_job: dict[str, Any]) -> dict[str, Any]:
    params = dict(render_job)
    params["created_at"] = params.get("created_at") or params.get("started_at") or _utc_now()
    params["updated_at"] = params.get("updated_at") or params.get("completed_at") or _utc_now()
    return params


def _artifact_render_job_select_sql(where_clause: str) -> str:
    return f"""
        SELECT
            render_job_id,
            artifact_id,
            artifact_version_id,
            job_status,
            current_stage,
            progress_mode,
            progress_percent,
            retryable,
            failure_code,
            started_at,
            completed_at
        FROM ae_artifact_render_jobs
        WHERE {where_clause}
    """


def _artifact_render_job_from_row(row: Any) -> dict[str, Any]:
    data = dict(row)
    return {
        "render_job_id": data["render_job_id"],
        "artifact_id": data["artifact_id"],
        "artifact_version_id": data["artifact_version_id"],
        "job_status": data["job_status"],
        "current_stage": data["current_stage"],
        "progress_mode": data["progress_mode"],
        "progress_percent": data["progress_percent"],
        "retryable": bool(data["retryable"]),
        "failure_code": data["failure_code"],
        "started_at": _datetime_value(data["started_at"]),
        "completed_at": _datetime_value(data["completed_at"]),
    }


def _artifact_file_upsert_sql() -> str:
    return """
        INSERT INTO ae_artifact_files (
            artifact_file_id,
            artifact_version_id,
            artifact_id,
            format,
            mime_type,
            file_name,
            storage_ref,
            file_size_bytes,
            file_hash,
            source_version_hash,
            created_at
        )
        VALUES (
            :artifact_file_id,
            :artifact_version_id,
            :artifact_id,
            :format,
            :mime_type,
            :file_name,
            :storage_ref,
            :file_size_bytes,
            :file_hash,
            :source_version_hash,
            :created_at
        )
        ON CONFLICT (artifact_file_id) DO UPDATE SET
            storage_ref = excluded.storage_ref,
            file_size_bytes = excluded.file_size_bytes,
            file_hash = excluded.file_hash,
            source_version_hash = excluded.source_version_hash
    """


def _artifact_file_params(
    artifact_id: str,
    artifact_file: dict[str, Any],
) -> dict[str, Any]:
    params = dict(artifact_file)
    params["artifact_id"] = artifact_id
    return params


def _artifact_file_select_sql(where_clause: str) -> str:
    return f"""
        SELECT
            artifact_file_id,
            artifact_version_id,
            format,
            mime_type,
            file_name,
            storage_ref,
            file_size_bytes,
            file_hash,
            source_version_hash,
            created_at
        FROM ae_artifact_files
        WHERE {where_clause}
    """


def _artifact_file_from_row(row: Any) -> dict[str, Any]:
    data = dict(row)
    return {
        "artifact_file_id": data["artifact_file_id"],
        "artifact_version_id": data["artifact_version_id"],
        "format": data["format"],
        "mime_type": data["mime_type"],
        "file_name": data["file_name"],
        "storage_ref": data["storage_ref"],
        "file_size_bytes": data["file_size_bytes"],
        "file_hash": data["file_hash"],
        "source_version_hash": data["source_version_hash"],
        "created_at": _datetime_value(data["created_at"]),
    }


def _artifact_link_upsert_sql(dialect_name: str) -> str:
    json_exprs = _json_param_exprs(ARTIFACT_LINK_JSON_FIELDS, dialect_name)
    return f"""
        INSERT INTO ae_artifact_links (
            artifact_link_id,
            artifact_file_id,
            link_type,
            access_policy,
            link_route,
            expires_at,
            created_by_actor_ref,
            download_count,
            revoked_at,
            created_at
        )
        VALUES (
            :artifact_link_id,
            :artifact_file_id,
            :link_type,
            :access_policy,
            :link_route,
            :expires_at,
            {json_exprs["created_by_actor_ref"]},
            :download_count,
            :revoked_at,
            :created_at
        )
        ON CONFLICT (artifact_link_id) DO UPDATE SET
            access_policy = excluded.access_policy,
            link_route = excluded.link_route,
            expires_at = excluded.expires_at,
            created_by_actor_ref = excluded.created_by_actor_ref,
            download_count = excluded.download_count,
            revoked_at = excluded.revoked_at
    """


def _artifact_link_params(artifact_link: dict[str, Any]) -> dict[str, Any]:
    params = dict(artifact_link)
    params["created_at"] = params.get("created_at") or _utc_now()
    return _json_params(params, ARTIFACT_LINK_JSON_FIELDS)


def _artifact_link_select_sql(where_clause: str) -> str:
    return f"""
        SELECT
            artifact_link_id,
            artifact_file_id,
            link_type,
            access_policy,
            link_route,
            expires_at,
            created_by_actor_ref,
            download_count,
            revoked_at
        FROM ae_artifact_links
        WHERE {where_clause}
    """


def _artifact_link_from_row(row: Any) -> dict[str, Any]:
    data = dict(row)
    return {
        "artifact_link_id": data["artifact_link_id"],
        "artifact_file_id": data["artifact_file_id"],
        "link_type": data["link_type"],
        "access_policy": data["access_policy"],
        "link_route": data["link_route"],
        "expires_at": _nullable_datetime_value(data["expires_at"]),
        "created_by_actor_ref": _json_value(data["created_by_actor_ref"], {}),
        "download_count": data["download_count"],
        "revoked_at": _nullable_datetime_value(data["revoked_at"]),
    }


def _json_param_exprs(
    field_names: tuple[str, ...],
    dialect_name: str,
) -> dict[str, str]:
    return {
        field_name: _json_param_expr(field_name, dialect_name)
        for field_name in field_names
    }


def _json_param_expr(name: str, dialect_name: str) -> str:
    if dialect_name == "postgresql":
        return f"CAST(:{name} AS jsonb)"
    return f":{name}"


def _json_params(
    params: dict[str, Any],
    field_names: tuple[str, ...],
) -> dict[str, Any]:
    updated = dict(params)
    for field_name in field_names:
        updated[field_name] = json.dumps(updated[field_name])
    return updated


def _json_value(value: Any, default: Any) -> Any:
    if value is None:
        return default
    if isinstance(value, str):
        return json.loads(value)
    return value


def _storage_ref_relative_path(storage_ref: str) -> str:
    prefix = "ae://artifacts/"
    if not isinstance(storage_ref, str) or not storage_ref.startswith(prefix):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_storage_ref_invalid",
            detail="Artifact storage ref must use the ae://artifacts scheme.",
        )
    relative = storage_ref.removeprefix(prefix)
    parts = relative.split("/")
    if not relative or any(part in {"", ".", ".."} for part in parts):
        raise ArtifactHandoffError(
            status_code=422,
            error_code="ae.artifact_storage_ref_invalid",
            detail="Artifact storage ref contains an unsafe path segment.",
        )
    return relative


def _datetime_value(value: Any) -> str:
    if value is None:
        return _utc_now()
    if hasattr(value, "isoformat"):
        return value.isoformat().replace("+00:00", "Z")
    return str(value)


def _nullable_datetime_value(value: Any) -> str | None:
    return None if value is None else _datetime_value(value)


def _dialect_name(session: Session) -> str:
    return session.get_bind().dialect.name


def _utc_now() -> str:
    return datetime.now(UTC).isoformat().replace("+00:00", "Z")
