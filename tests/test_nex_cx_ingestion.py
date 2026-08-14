from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from nex_cx.ingestion import (
    ContentIngestionStore,
    CX_SOURCE_FILE_MATERIALIZATION_RECEIPT_SCHEMA_VERSION,
    CxStorageConfig,
    IngestionError,
    UPLOAD_OWNER_RESOLVER_DISABLED,
    UPLOAD_OWNER_RESOLVER_VERIFY,
    SOURCE_READER_MATERIALIZED_LOCAL_FILE,
    SOURCE_READER_RUNTIME_MEMORY,
    build_ingestion_job,
    build_source_file_materialization_receipt,
    build_storage_config,
    build_upload_ownership_ref,
    build_upload_registration,
    markdown_from_source_text,
    materialize_local_source_bytes,
    materialize_local_source_file,
    normalize_upload_owner_resolver_mode,
    payload_source_kind,
    read_verified_materialized_source_bytes,
    register_ingestion_routes,
    resolve_upload_ownership,
    run_text_extraction_job,
    sha256_bytes,
    sanitize_filename,
    sha256_text,
    source_content_from_payload,
    source_bytes_for_extraction,
    storage_date_partition,
    storage_paths_for_document,
    stored_extension_for,
    validate_upload_size,
    write_extracted_markdown,
)
import nex_cx.ingestion as cx_ingestion
from nex_cx.repository import CxContentRepositoryError
from nex_runtime import (
    SERVICE_SPECS,
    SubjectRegistryResolverError,
    build_service_app,
    issue_mock_service_token,
)

TRACE_ID = "4bf92f3577b34da6a3ce929d0e0e4736"
REQUEST_ID = "0189f0ff-8f22-4f72-9b47-b481dc21bb21"
OWNER_REF = {
    "ownership_schema_version": "cx_source_ownership_ref.v1",
    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
    "uploaded_by_subject_ref": {"type": "oa.user", "id": "uploader-a"},
    "legacy": {"tenant_id": "tenant-a", "owner_user_id": "user-a"},
    "compatibility_mode": "legacy_owner_fields_mapped_to_oa_subject_refs",
}


def sample_pdf_bytes(text: str = "Slice 0285 PDF ingestion text") -> bytes:
    text_bytes = text.encode("ascii")
    stream = b"BT /F1 18 Tf 36 96 Td (" + text_bytes + b") Tj ET"
    objects = (
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        (
            b"<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        ),
    )
    pdf = b"%PDF-1.4\n"
    offsets: list[int] = []
    for object_number, body in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += (
            f"{object_number} 0 obj\n".encode("ascii")
            + body
            + b"\nendobj\n"
        )
    startxref = len(pdf)
    xref_entries = b"".join(
        f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets
    )
    return (
        pdf
        + b"xref\n0 6\n0000000000 65535 f \n"
        + xref_entries
        + b"trailer\n<< /Root 1 0 R /Size 6 >>\nstartxref\n"
        + str(startxref).encode("ascii")
        + b"\n%%EOF\n"
    )


def sample_docx_bytes(
    *,
    title: str = "Slice 0286 DOCX ingestion title",
    body: str = "Slice 0286 DOCX ingestion body",
) -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(title)
    document.add_paragraph(body)
    document.save(buffer)
    return buffer.getvalue()


class FakeOwnerResolver:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def resolve_ownership_ref(
        self,
        ownership_ref: dict[str, object],
        *,
        request_id: str,
        trace_id: str,
        ensure: bool = False,
    ) -> dict[str, object]:
        self.calls.append(
            {
                "ownership_ref": ownership_ref,
                "request_id": request_id,
                "trace_id": trace_id,
                "ensure": ensure,
            }
        )
        return {
            "resolver_schema_version": "oa_subject_registry_resolver.v1",
            "resolution_status": "RESOLVED",
            "ensure": ensure,
        }


class FailingOwnerResolver:
    def resolve_ownership_ref(
        self,
        ownership_ref: dict[str, object],
        *,
        request_id: str,
        trace_id: str,
        ensure: bool = False,
    ) -> dict[str, object]:
        raise SubjectRegistryResolverError(
            status_code=404,
            error_code="oa.subject_not_found",
            detail="Subject was not found.",
            retryable=False,
        )


class FailingDocumentDetailRepository:
    def get_content_object(self, content_object_id: str) -> dict[str, object]:
        raise CxContentRepositoryError(
            error_code="cx.content_repository_unavailable",
            detail=f"repository offline for {content_object_id}",
            status_code=503,
        )


def auth_headers() -> dict[str, str]:
    issued = issue_mock_service_token(service_id="nex-ae-api", audience="nex-cx")
    return {
        "Authorization": f"Bearer {issued.access_token}",
        "X-Request-ID": REQUEST_ID,
        "traceparent": f"00-{TRACE_ID}-00f067aa0ba902b7-01",
    }


def storage_config(tmp_path: Path) -> CxStorageConfig:
    return CxStorageConfig(
        data_root=tmp_path,
        source_root=tmp_path / "cx" / "source-files",
        extracted_markdown_root=tmp_path / "cx" / "extracted-markdown",
        extraction_temp_root=tmp_path / "cx" / "extraction-temp",
        chunk_policy="chunk_1000_100",
        chunk_size=1000,
        chunk_overlap=100,
        bm25_tokenizer="mecab_ko",
        bm25_tokenizer_fallback="korean_mixed_v1",
    )


def build_test_client(
    tmp_path: Path,
    *,
    owner_resolver: FakeOwnerResolver | FailingOwnerResolver | None = None,
    owner_resolver_mode: str | None = None,
) -> tuple[TestClient, ContentIngestionStore]:
    app = build_service_app(SERVICE_SPECS["nex-cx"])
    store = ContentIngestionStore()
    register_ingestion_routes(
        app,
        store=store,
        storage_config=storage_config(tmp_path),
        owner_resolver=owner_resolver,
        owner_resolver_mode=owner_resolver_mode,
    )
    return TestClient(app), store


def test_build_storage_config_uses_data_root_defaults() -> None:
    config = build_storage_config({"NEX_DATA_ROOT": "/data/nex-platform"})

    assert str(config.data_root) == "/data/nex-platform"
    assert str(config.source_root) == "/data/nex-platform/cx/source-files"
    assert str(config.extracted_markdown_root) == "/data/nex-platform/cx/extracted-markdown"
    assert str(config.extraction_temp_root) == "/data/nex-platform/cx/extraction-temp"
    assert config.chunk_size == 1000
    assert config.chunk_overlap == 100
    assert config.bm25_tokenizer == "mecab_ko"
    assert config.bm25_tokenizer_fallback == "korean_mixed_v1"
    assert config.max_upload_size_bytes == 50 * 1024 * 1024


def test_build_storage_config_accepts_explicit_overrides() -> None:
    config = build_storage_config(
        {
            "NEX_DATA_ROOT": "/nex-data",
            "NEX_CX_SOURCE_STORAGE_ROOT": "/source",
            "NEX_CX_EXTRACTED_MARKDOWN_ROOT": "/markdown",
            "NEX_CX_EXTRACTION_TEMP_ROOT": "/temp",
            "NEX_CX_DEFAULT_CHUNK_POLICY": "custom",
            "NEX_CX_CHUNK_SIZE": "1200",
            "NEX_CX_CHUNK_OVERLAP": "80",
            "NEX_CX_BM25_TOKENIZER": "mecab_ko",
            "NEX_CX_BM25_TOKENIZER_FALLBACK": "korean_mixed_v1",
            "NEX_CX_MAX_UPLOAD_SIZE_BYTES": "4096",
        }
    )

    assert str(config.source_root) == "/source"
    assert str(config.extracted_markdown_root) == "/markdown"
    assert str(config.extraction_temp_root) == "/temp"
    assert config.chunk_policy == "custom"
    assert config.chunk_size == 1200
    assert config.chunk_overlap == 80
    assert config.max_upload_size_bytes == 4096


@pytest.mark.parametrize(
    "env",
    [
        {"NEX_CX_CHUNK_SIZE": "0"},
        {"NEX_CX_CHUNK_SIZE": "not-an-int"},
        {"NEX_CX_CHUNK_OVERLAP": "-1"},
        {"NEX_CX_CHUNK_OVERLAP": "not-an-int"},
        {"NEX_CX_MAX_UPLOAD_SIZE_BYTES": "0"},
        {"NEX_CX_MAX_UPLOAD_SIZE_BYTES": "not-an-int"},
    ],
)
def test_build_storage_config_rejects_bad_numeric_values(env: dict[str, str]) -> None:
    with pytest.raises(ValueError):
        build_storage_config(env)


def test_sanitize_filename_trims_basename() -> None:
    assert sanitize_filename("  report.pdf  ") == "report.pdf"


@pytest.mark.parametrize("filename", ["", ".", "..", "../report.pdf", "a/b.pdf", "bad\x00.pdf"])
def test_sanitize_filename_rejects_unsafe_names(filename: str) -> None:
    with pytest.raises(IngestionError):
        sanitize_filename(filename)


def test_sanitize_filename_rejects_long_name() -> None:
    with pytest.raises(IngestionError):
        sanitize_filename("a" * 256)


def test_storage_paths_partition_by_date_hash_and_generated_filename(tmp_path: Path) -> None:
    paths = storage_paths_for_document(
        storage_config=storage_config(tmp_path),
        filename="report.md",
        source_sha256="a" * 64,
        document_id="doc-001",
        created_at="2026-08-02T00:00:00Z",
    )

    assert paths["source_storage_backend"] == "local_filesystem"
    assert paths["source_storage_key"] == "20260802/aa/aa/doc-001.md"
    assert paths["source_storage_path"].endswith("/cx/source-files/20260802/aa/aa/doc-001.md")
    assert paths["stored_filename"] == "doc-001.md"
    assert paths["stored_extension"] == ".md"
    assert paths["extracted_markdown_path"].endswith("/cx/extracted-markdown/aa/doc-001.md")
    assert paths["extraction_temp_path"].endswith("/cx/extraction-temp/doc-001")


def test_stored_extension_for_uses_safe_lowercase_suffix_only() -> None:
    assert stored_extension_for("REPORT.PDF") == ".pdf"
    assert stored_extension_for("archive.tar.gz") == ".gz"
    assert stored_extension_for("source") == ""
    assert stored_extension_for("unsafe.$$$") == ""


def test_storage_date_partition_falls_back_for_invalid_timestamp() -> None:
    partition = storage_date_partition("not-a-timestamp")

    assert len(partition) == 8
    assert partition.isdigit()


def test_build_ingestion_job_links_document() -> None:
    job = build_ingestion_job(
        document_id="doc-001",
        upload_id="upload-001",
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
        created_at="2026-08-02T00:00:00Z",
    )

    assert job["job_schema_version"] == "common_job.v1"
    assert job["status"] == "QUEUED"
    assert job["subject_ref"] == {"type": "cx.document", "id": "doc-001"}
    assert job["links"]["document"] == "/api/v1/documents/doc-001"


def test_build_upload_registration_hashes_content_without_leaking_text(tmp_path: Path) -> None:
    record = build_upload_registration(
        {
            "filename": "mvp-srs.md",
            "content_type": "text/markdown",
            "content_text": "# MVP\nTraceable content",
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    assert record["document_schema_version"] == "cx_upload_registration.v1"
    assert record["source_sha256"] == sha256_text("# MVP\nTraceable content")
    assert record["size_bytes"] == len("# MVP\nTraceable content".encode("utf-8"))
    assert record["retrieval_policy"]["chunk_policy"] == "chunk_1000_100"
    assert record["retrieval_policy"]["bm25_tokenizer_fallback"] == "korean_mixed_v1"
    assert record["ownership_ref"] == {
        "ownership_schema_version": "cx_source_ownership_ref.v1",
        "tenant_ref": {"type": "oa.tenant", "id": "local-tenant"},
        "owner_subject_ref": {"type": "oa.user", "id": "local-user"},
        "uploaded_by_subject_ref": {"type": "oa.user", "id": "local-user"},
        "legacy": {"tenant_id": "local-tenant", "owner_user_id": "local-user"},
        "compatibility_mode": "legacy_owner_fields_mapped_to_oa_subject_refs",
    }
    assert "Traceable content" not in str(record)


def test_build_upload_registration_accepts_base64_source_bytes_without_leak(
    tmp_path: Path,
) -> None:
    source_bytes = b"%PDF-1.7\nbinary-ish\x00content\n"
    content_base64 = base64.b64encode(source_bytes).decode("ascii")

    record = build_upload_registration(
        {
            "filename": "source.pdf",
            "content_type": "application/pdf",
            "content_base64": content_base64,
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    assert record["source_sha256"] == sha256_bytes(source_bytes)
    assert record["size_bytes"] == len(source_bytes)
    assert record["upload_boundary"] == {
        "payload_source": "content_base64",
        "source_content_in_record": False,
        "checksum_algorithm": "sha256",
        "max_size_bytes": 50 * 1024 * 1024,
    }
    assert content_base64 not in str(record)
    assert "binary-ish" not in str(record)


def test_build_upload_registration_accepts_matching_explicit_source_hash(
    tmp_path: Path,
) -> None:
    source_bytes = b"explicit source hash bytes"
    source_sha256 = sha256_bytes(source_bytes)

    text_record = build_upload_registration(
        {
            "filename": "source.txt",
            "content_text": "explicit source hash text",
            "source_sha256": sha256_text("explicit source hash text").upper(),
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    binary_record = build_upload_registration(
        {
            "filename": "source.bin",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
            "source_sha256": source_sha256,
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    assert text_record["source_sha256"] == sha256_text("explicit source hash text")
    assert binary_record["source_sha256"] == source_sha256


def test_source_content_from_payload_rejects_conflicts_and_bad_base64() -> None:
    assert source_content_from_payload({"content_text": "hello"}) == ("hello", None)
    assert source_content_from_payload(
        {"content_base64": base64.b64encode(b"hello").decode("ascii")}
    ) == (None, b"hello")

    with pytest.raises(IngestionError) as conflict:
        source_content_from_payload({"content_text": "hello", "content_base64": "aGVsbG8="})
    with pytest.raises(IngestionError) as bad_type:
        source_content_from_payload({"content_base64": 123})
    with pytest.raises(IngestionError) as bad_value:
        source_content_from_payload({"content_base64": "not base64"})

    assert conflict.value.error_code == "cx.upload_content_source_conflict"
    assert bad_type.value.error_code == "cx.upload_content_base64_invalid"
    assert bad_value.value.error_code == "cx.upload_content_base64_invalid"


def test_upload_size_validation_rejects_mismatch_and_limit(tmp_path: Path) -> None:
    with pytest.raises(IngestionError) as mismatch:
        build_upload_registration(
            {
                "filename": "source.txt",
                "content_text": "12345",
                "size_bytes": 4,
            },
            storage_config=storage_config(tmp_path),
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )
    with pytest.raises(IngestionError) as too_large:
        validate_upload_size(11, max_upload_size_bytes=10)

    assert mismatch.value.error_code == "cx.upload_size_mismatch"
    assert too_large.value.status_code == 413
    assert too_large.value.error_code == "cx.upload_size_exceeds_limit"


def test_store_keeps_source_text_private_for_mock_extraction(tmp_path: Path) -> None:
    store = ContentIngestionStore()
    record = build_upload_registration(
        {"filename": "source.md", "content_text": "private source text"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    store.save_upload_registration(record, source_text="private source text")

    assert store.get_source_text(record["upload_id"]) == "private source text"
    assert store.get_source_bytes(record["upload_id"]) == b"private source text"
    assert "private source text" not in str(store.get_document(record["document_id"]))


def test_store_materializes_source_file_and_marks_checksum_verified(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    record = build_upload_registration(
        {"filename": "source.md", "content_text": "private source text"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    saved = store.save_upload_registration(record, source_text="private source text")
    refs = store.get_content_ref(saved["document_id"])
    source_file = store.content_repository.get_source_file(refs["source_file_id"])

    source_path = Path(saved["storage"]["source_storage_path"])
    assert source_path.read_bytes() == b"private source text"
    assert sha256_bytes(source_path.read_bytes()) == saved["source_sha256"]
    assert source_file["checksum_verified_at"] is not None
    assert saved["original_filename"] not in str(source_path.parent)


def test_source_file_materialization_receipt_hides_local_path_and_source_bytes(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    source_bytes = b"\x00binary source materialization receipt"
    record = build_upload_registration(
        {
            "filename": "source.bin",
            "content_type": "application/octet-stream",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    saved = store.save_upload_registration(record, source_bytes=source_bytes)

    receipt = build_source_file_materialization_receipt(
        store=store,
        document_id=saved["document_id"],
        tenant_id="tenant-a",
        owner_user_id="user-a",
        source_kind="memory",
        database_env="NEX_CX_TEST_DATABASE_URL",
        redacted_database_url="postgresql+psycopg://nex_cx_user:***@127.0.0.1/nex_cx_test",
    )
    wrong_owner = build_source_file_materialization_receipt(
        store=store,
        document_id=saved["document_id"],
        tenant_id="tenant-a",
        owner_user_id="other-user",
    )

    assert receipt is not None
    assert receipt["receipt_schema_version"] == (
        CX_SOURCE_FILE_MATERIALIZATION_RECEIPT_SCHEMA_VERSION
    )
    assert receipt["source"]["database_env"] == "NEX_CX_TEST_DATABASE_URL"
    assert receipt["source_file"]["checksum_verified"] is True
    assert receipt["source_file"]["checksum_verified_at"] is not None
    assert receipt["materialization"] == {
        "status": "VERIFIED",
        "payload_source": "content_base64",
        "source_bytes_captured": True,
        "checksum_algorithm": "sha256",
        "source_content_in_receipt": False,
        "local_storage_path_included": False,
    }
    assert receipt["metadata"]["storage_path_redacted"] is True
    assert "source_storage_path" not in str(receipt)
    assert str(tmp_path) not in str(receipt)
    assert "binary source materialization receipt" not in str(receipt)
    assert wrong_owner is None


def test_source_file_materialization_receipt_reports_metadata_only_pending(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    record = build_upload_registration(
        {
            "filename": "source.pdf",
            "source_sha256": "a" * 64,
            "size_bytes": 4096,
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    saved = store.save_upload_registration(record)

    receipt = build_source_file_materialization_receipt(
        store=store,
        document_id=saved["document_id"],
        tenant_id="tenant-a",
        owner_user_id="user-a",
    )

    assert receipt is not None
    assert receipt["source_file"]["checksum_verified"] is False
    assert receipt["source_file"]["checksum_verified_at"] is None
    assert receipt["materialization"]["status"] == "PENDING"
    assert receipt["materialization"]["payload_source"] == "precomputed_hash"
    assert receipt["materialization"]["source_bytes_captured"] is False


def test_source_file_materialization_receipt_collapses_missing_lineage_or_file(
    tmp_path: Path,
) -> None:
    missing_lineage_store = ContentIngestionStore()
    missing_file_store = ContentIngestionStore()
    record = build_upload_registration(
        {
            "filename": "source.md",
            "content_text": "lineage source",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    missing_lineage = missing_lineage_store.save_upload_registration(
        record,
        source_text="lineage source",
    )
    missing_file = missing_file_store.save_upload_registration(
        record,
        source_text="lineage source",
    )
    lineage_refs = missing_lineage_store.get_content_ref(missing_lineage["document_id"])
    file_refs = missing_file_store.get_content_ref(missing_file["document_id"])
    missing_lineage_store.content_repository.content_objects[
        lineage_refs["content_object_id"]
    ]["source_file_id"] = ""
    missing_file_store.content_repository.content_objects[
        file_refs["content_object_id"]
    ]["source_file_id"] = "missing-source-file"

    assert (
        build_source_file_materialization_receipt(
            store=missing_lineage_store,
            document_id=missing_lineage["document_id"],
            tenant_id="tenant-a",
            owner_user_id="user-a",
        )
        is None
    )
    assert (
        build_source_file_materialization_receipt(
            store=missing_file_store,
            document_id=missing_file["document_id"],
            tenant_id="tenant-a",
            owner_user_id="user-a",
        )
        is None
    )


def test_store_derives_owner_scope_and_uploaded_by_from_upload_record(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    record = build_upload_registration(
        {
            "filename": "source.md",
            "content_text": "canonical source text",
            "ownership_ref": OWNER_REF,
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    saved = store.save_upload_registration(record, source_text="canonical source text")
    refs = store.get_content_ref(saved["document_id"])
    content_object = store.content_repository.get_content_object(
        refs["content_object_id"]
    )

    assert saved["ownership_ref"] == OWNER_REF
    assert content_object["ownership_ref"] == OWNER_REF
    assert content_object["tenant_id"] == "tenant-a"
    assert content_object["owner_user_id"] == "user-a"


def test_store_defaults_owner_scope_for_legacy_record_without_ownership_metadata(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    record = build_upload_registration(
        {"filename": "source.md", "content_text": "legacy fallback text"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    legacy_record = {
        key: value
        for key, value in record.items()
        if key not in {"ownership", "ownership_ref"}
    }

    saved = store.save_upload_registration(legacy_record)
    refs = store.get_content_ref(saved["document_id"])
    content_object = store.content_repository.get_content_object(
        refs["content_object_id"]
    )

    assert content_object["tenant_id"] == "local-tenant"
    assert content_object["owner_user_id"] == "local-user"


def test_materialize_local_source_file_is_idempotent_for_matching_file(
    tmp_path: Path,
) -> None:
    record = build_upload_registration(
        {"filename": "source.md", "content_text": "same bytes"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    first_verified_at = materialize_local_source_file(record, "same bytes")
    second_verified_at = materialize_local_source_file(record, "same bytes")

    assert first_verified_at.endswith("Z")
    assert second_verified_at.endswith("Z")
    assert Path(record["storage"]["source_storage_path"]).read_text(encoding="utf-8") == (
        "same bytes"
    )


def test_materialize_local_source_bytes_writes_binary_source(tmp_path: Path) -> None:
    source_bytes = b"\x00\x01binary upload"
    record = build_upload_registration(
        {
            "filename": "source.bin",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    verified_at = materialize_local_source_bytes(record, source_bytes)

    assert verified_at.endswith("Z")
    assert Path(record["storage"]["source_storage_path"]).read_bytes() == source_bytes


def test_materialize_local_source_file_rejects_bad_checksum(tmp_path: Path) -> None:
    record = build_upload_registration(
        {"filename": "source.md", "content_text": "original"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    with pytest.raises(IngestionError) as exc:
        materialize_local_source_file(record, "changed")

    assert exc.value.error_code == "cx.source_checksum_mismatch"


def test_materialize_local_source_file_rejects_unsafe_storage_metadata(
    tmp_path: Path,
) -> None:
    record = build_upload_registration(
        {"filename": "source.md", "content_text": "source"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    with pytest.raises(IngestionError) as backend_exc:
        materialize_local_source_file(
            {
                **record,
                "storage": {**record["storage"], "source_storage_backend": "s3"},
            },
            "source",
        )
    with pytest.raises(IngestionError) as key_exc:
        materialize_local_source_file(
            {
                **record,
                "storage": {**record["storage"], "source_storage_key": "../source.md"},
            },
            "source",
        )
    with pytest.raises(IngestionError) as path_exc:
        materialize_local_source_file(
            {
                **record,
                "storage": {**record["storage"], "source_storage_path": "relative/source.md"},
            },
            "source",
        )

    assert backend_exc.value.error_code == "cx.source_storage_backend_unsupported"
    assert key_exc.value.error_code == "cx.source_storage_key_invalid"
    assert path_exc.value.error_code == "cx.source_storage_path_invalid"


def test_materialize_local_source_file_rejects_existing_file_collision(
    tmp_path: Path,
) -> None:
    record = build_upload_registration(
        {"filename": "source.md", "content_text": "expected"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    source_path = Path(record["storage"]["source_storage_path"])
    source_path.parent.mkdir(parents=True)
    source_path.write_text("different", encoding="utf-8")

    with pytest.raises(IngestionError) as exc:
        materialize_local_source_file(record, "expected")

    assert exc.value.error_code == "cx.source_file_collision"


def test_build_upload_registration_accepts_precomputed_hash(tmp_path: Path) -> None:
    record = build_upload_registration(
        {
            "filename": "source.pdf",
            "content_type": "application/pdf",
            "source_sha256": "A" * 64,
            "size_bytes": 2048,
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    assert record["source_sha256"] == "a" * 64
    assert record["size_bytes"] == 2048
    assert record["upload_boundary"]["payload_source"] == "precomputed_hash"
    assert not record["storage"]["source_storage_key"].endswith(
        f"/{record['document_id']}.pdf"
    )
    assert record["storage"]["source_storage_key"].endswith(".pdf")
    assert record["storage"]["stored_extension"] == ".pdf"
    assert "/source.pdf" not in record["storage"]["source_storage_path"]


def test_build_upload_registration_scopes_document_id_to_owner(tmp_path: Path) -> None:
    base_payload = {
        "filename": "source.md",
        "content_type": "text/markdown",
        "content_text": "same source bytes",
        "tenant_id": "tenant-a",
    }

    user_a = build_upload_registration(
        {**base_payload, "owner_user_id": "user-a"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    user_b = build_upload_registration(
        {**base_payload, "owner_user_id": "user-b"},
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    assert user_a["document_id"] != user_b["document_id"]
    assert user_a["source_sha256"] == user_b["source_sha256"]
    assert user_a["storage"]["source_storage_key"] == user_b["storage"]["source_storage_key"]
    assert user_a["ownership"] == {"tenant_id": "tenant-a", "owner_user_id": "user-a"}
    assert user_a["ownership_ref"]["tenant_ref"]["id"] == "tenant-a"
    assert user_a["ownership_ref"]["owner_subject_ref"]["id"] == "user-a"


def test_build_upload_registration_consumes_canonical_ownership_ref(
    tmp_path: Path,
) -> None:
    record = build_upload_registration(
        {
            "filename": "source.md",
            "content_text": "canonical owner source bytes",
            "ownership_ref": OWNER_REF,
        },
        storage_config=storage_config(tmp_path),
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    legacy_record = build_upload_registration(
        {
            "filename": "renamed.md",
            "content_text": "canonical owner source bytes",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id="other-request",
        trace_id=TRACE_ID,
    )

    assert record["ownership"] == {"tenant_id": "tenant-a", "owner_user_id": "user-a"}
    assert record["ownership_ref"] == OWNER_REF
    assert record["document_id"] == legacy_record["document_id"]


def test_build_upload_ownership_ref_accepts_direct_subject_refs() -> None:
    ownership_ref = build_upload_ownership_ref(
        {
            "tenant_ref": {"type": "oa.tenant", "id": "tenant-b"},
            "owner_subject_ref": {"type": "oa.user", "id": "user-b"},
            "uploaded_by_subject_ref": {"type": "oa.user", "id": "uploader-b"},
        }
    )

    assert ownership_ref["legacy"] == {
        "tenant_id": "tenant-b",
        "owner_user_id": "user-b",
    }
    assert ownership_ref["uploaded_by_subject_ref"] == {
        "type": "oa.user",
        "id": "uploader-b",
    }


def test_build_upload_ownership_ref_accepts_matching_direct_and_partial_legacy() -> None:
    ownership_ref = build_upload_ownership_ref(
        {
            "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
            "ownership_ref": {
                "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                "legacy": {"tenant_id": "tenant-a"},
            },
        }
    )

    assert ownership_ref == {
        "ownership_schema_version": "cx_source_ownership_ref.v1",
        "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
        "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
        "uploaded_by_subject_ref": {"type": "oa.user", "id": "user-a"},
        "legacy": {"tenant_id": "tenant-a", "owner_user_id": "user-a"},
        "compatibility_mode": "legacy_owner_fields_mapped_to_oa_subject_refs",
    }


def test_store_returns_existing_registration_for_same_owner_duplicate(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    first = build_upload_registration(
        {
            "filename": "source.md",
            "content_text": "same source bytes",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id="request-a",
        trace_id=TRACE_ID,
    )
    second = build_upload_registration(
        {
            "filename": "renamed.md",
            "content_text": "same source bytes",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id="request-b",
        trace_id=TRACE_ID,
    )

    created = store.save_upload_registration(
        first,
        source_text="same source bytes",
        tenant_id="tenant-a",
        owner_user_id="user-a",
    )
    duplicate = store.save_upload_registration(
        second,
        source_text="same source bytes",
        tenant_id="tenant-a",
        owner_user_id="user-a",
    )

    assert created["dedupe"]["status"] == "CREATED"
    assert duplicate["document_id"] == created["document_id"]
    assert duplicate["dedupe"]["status"] == "ALREADY_EXISTS"
    assert duplicate["dedupe"]["existing_document_id"] == created["document_id"]
    assert store.get_source_text(created["upload_id"]) == "same source bytes"
    assert store.get_source_text(second["upload_id"]) is None


def test_duplicate_upload_without_source_content_returns_existing_record(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    payload = {
        "filename": "source.pdf",
        "source_sha256": "a" * 64,
        "size_bytes": 10,
        "tenant_id": "tenant-a",
        "owner_user_id": "user-a",
    }
    first = build_upload_registration(
        payload,
        storage_config=storage_config(tmp_path),
        request_id="request-a",
        trace_id=TRACE_ID,
    )
    second = build_upload_registration(
        {**payload, "filename": "renamed.pdf"},
        storage_config=storage_config(tmp_path),
        request_id="request-b",
        trace_id=TRACE_ID,
    )

    created = store.save_upload_registration(first)
    duplicate = store.save_upload_registration(second)

    assert duplicate["dedupe"]["status"] == "ALREADY_EXISTS"
    assert duplicate["document_id"] == created["document_id"]
    assert store.get_source_bytes(created["upload_id"]) is None


def test_store_allows_same_hash_for_different_owner_without_leaking_duplicate(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    payload = {
        "filename": "source.md",
        "content_text": "same source bytes",
        "tenant_id": "tenant-a",
    }
    user_a = build_upload_registration(
        {**payload, "owner_user_id": "user-a"},
        storage_config=storage_config(tmp_path),
        request_id="request-a",
        trace_id=TRACE_ID,
    )
    user_b = build_upload_registration(
        {**payload, "owner_user_id": "user-b"},
        storage_config=storage_config(tmp_path),
        request_id="request-b",
        trace_id=TRACE_ID,
    )

    store.save_upload_registration(user_a, tenant_id="tenant-a", owner_user_id="user-a")
    created_for_b = store.save_upload_registration(
        user_b,
        tenant_id="tenant-a",
        owner_user_id="user-b",
    )

    assert created_for_b["document_id"] == user_b["document_id"]
    assert created_for_b["dedupe"]["status"] == "CREATED"
    assert created_for_b["dedupe"]["existing_document_id"] is None
    assert len(store.content_repository.source_files) == 1


def test_duplicate_upload_with_bytes_materializes_existing_metadata_only_source(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    source_bytes = b"same source bytes"
    source_sha256 = sha256_bytes(source_bytes)
    first = build_upload_registration(
        {
            "filename": "source.pdf",
            "source_sha256": source_sha256,
            "size_bytes": len(source_bytes),
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id="request-a",
        trace_id=TRACE_ID,
    )
    second = build_upload_registration(
        {
            "filename": "source.pdf",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        storage_config=storage_config(tmp_path),
        request_id="request-b",
        trace_id=TRACE_ID,
    )

    created = store.save_upload_registration(first, tenant_id="tenant-a", owner_user_id="user-a")
    duplicate = store.save_upload_registration(
        second,
        source_bytes=source_bytes,
        tenant_id="tenant-a",
        owner_user_id="user-a",
    )
    refs = store.get_content_ref(created["document_id"])
    source_file = store.content_repository.get_source_file(refs["source_file_id"])

    assert duplicate["dedupe"]["status"] == "ALREADY_EXISTS"
    assert Path(created["storage"]["source_storage_path"]).read_bytes() == source_bytes
    assert store.get_source_bytes(created["upload_id"]) == source_bytes
    assert source_file["checksum_verified_at"] is not None


@pytest.mark.parametrize(
    ("payload", "error_code"),
    [
        ({"content_text": "hello"}, "cx.filename_required"),
        ({"filename": "report.pdf", "size_bytes": 10}, "cx.source_sha256_required"),
        (
            {"filename": "report.pdf", "source_sha256": "bad", "size_bytes": 10},
            "cx.upload_hash_invalid",
        ),
        (
            {"filename": "report.pdf", "content_text": "hello", "source_sha256": True},
            "cx.upload_hash_invalid",
        ),
        (
            {"filename": "report.pdf", "source_sha256": "a" * 64},
            "cx.upload_size_required",
        ),
        (
            {"filename": "report.pdf", "source_sha256": "a" * 64, "size_bytes": -1},
            "cx.upload_size_invalid",
        ),
        (
            {"filename": "report.pdf", "content_type": "", "content_text": "hello"},
            "cx.content_type_invalid",
        ),
        (
            {"filename": "report.pdf", "content_text": ["hello"]},
            "cx.upload_content_text_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "source_sha256": "a" * 64,
            },
            "cx.upload_hash_mismatch",
        ),
        (
            {
                "filename": "report.pdf",
                "content_base64": base64.b64encode(b"hello").decode("ascii"),
                "source_sha256": "a" * 64,
            },
            "cx.upload_hash_mismatch",
        ),
        (
            {"filename": "report.pdf", "content_text": "hello", "content_base64": "aGVsbG8="},
            "cx.upload_content_source_conflict",
        ),
        (
            {"filename": "report.pdf", "content_text": "hello", "ownership_ref": "owner-a"},
            "cx.upload_owner_invalid",
        ),
        (
            {"filename": "report.pdf", "content_text": "hello", "tenant_id": ""},
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "tenant_id": "tenant-a",
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-b"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "tenant_ref": {"type": "cx.tenant", "id": "tenant-a"},
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "tenant_ref": {"type": "oa.tenant", "id": "tenant-b"},
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                    "legacy": "tenant-a/user-a",
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                    "legacy": {
                        "tenant_id": "tenant-a",
                        "owner_user_id": "user-a",
                        "raw_profile": {},
                    },
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                    "email": "private@example.com",
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "ownership_schema_version": "cx_source_ownership_ref.v0",
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "compatibility_mode": "raw_identity",
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "tenant_ref": {
                        "type": "oa.tenant",
                        "id": "tenant-a",
                        "email": "private@example.com",
                    },
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                    "legacy": {"tenant_id": "tenant-b"},
                },
            },
            "cx.upload_owner_invalid",
        ),
        (
            {
                "filename": "report.pdf",
                "content_text": "hello",
                "ownership_ref": {
                    "tenant_ref": {"type": "oa.tenant", "id": "tenant-a"},
                    "owner_subject_ref": {"type": "oa.user", "id": "user-a"},
                    "legacy": {"tenant_id": "tenant-a", "owner_user_id": "user-b"},
                },
            },
            "cx.upload_owner_invalid",
        ),
    ],
)
def test_build_upload_registration_reports_validation_errors(
    tmp_path: Path,
    payload: dict[str, object],
    error_code: str,
) -> None:
    with pytest.raises(IngestionError) as exc:
        build_upload_registration(
            payload,
            storage_config=storage_config(tmp_path),
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )

    assert exc.value.error_code == error_code


def test_upload_owner_resolver_mode_normalization_and_disabled_skip() -> None:
    assert normalize_upload_owner_resolver_mode(None) == UPLOAD_OWNER_RESOLVER_DISABLED
    assert normalize_upload_owner_resolver_mode(" VERIFY ") == UPLOAD_OWNER_RESOLVER_VERIFY
    assert (
        resolve_upload_ownership(
            OWNER_REF,
            owner_resolver=None,
            owner_resolver_mode=UPLOAD_OWNER_RESOLVER_DISABLED,
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )
        is None
    )

    with pytest.raises(IngestionError) as bad_mode:
        normalize_upload_owner_resolver_mode("ensure")

    assert bad_mode.value.error_code == "cx.upload_owner_resolver_mode_invalid"


def test_resolve_upload_ownership_verify_mode_and_error_mapping() -> None:
    resolver = FakeOwnerResolver()
    resolved = resolve_upload_ownership(
        OWNER_REF,
        owner_resolver=resolver,
        owner_resolver_mode=UPLOAD_OWNER_RESOLVER_VERIFY,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    assert resolved["resolution_status"] == "RESOLVED"
    assert resolver.calls == [
        {
            "ownership_ref": OWNER_REF,
            "request_id": REQUEST_ID,
            "trace_id": TRACE_ID,
            "ensure": False,
        }
    ]

    with pytest.raises(IngestionError) as missing_resolver:
        resolve_upload_ownership(
            OWNER_REF,
            owner_resolver=None,
            owner_resolver_mode=UPLOAD_OWNER_RESOLVER_VERIFY,
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )

    assert missing_resolver.value.status_code == 503
    assert missing_resolver.value.error_code == "cx.upload_owner_resolver_unavailable"
    assert missing_resolver.value.retryable is True

    with pytest.raises(IngestionError) as unresolved:
        resolve_upload_ownership(
            OWNER_REF,
            owner_resolver=FailingOwnerResolver(),
            owner_resolver_mode=UPLOAD_OWNER_RESOLVER_VERIFY,
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )

    assert unresolved.value.status_code == 404
    assert unresolved.value.error_code == "cx.upload_owner_unresolved"
    assert unresolved.value.detail == "Subject was not found."


def test_upload_registration_endpoint_requires_service_claim(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)

    response = client.post("/api/v1/documents/uploads", json={"filename": "report.pdf"})

    assert response.status_code == 401
    assert response.json()["error_code"] == "AUTHORIZATION_HEADER_MISSING"


def test_upload_registration_endpoint_creates_document_and_job(tmp_path: Path) -> None:
    client, store = build_test_client(tmp_path)

    response = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.md",
            "content_type": "text/markdown",
            "content_text": "hello from upload",
        },
        headers=auth_headers(),
    )

    assert response.status_code == 202
    payload = response.json()
    assert payload["trace_id"] == TRACE_ID
    assert payload["request_id"] == REQUEST_ID
    assert payload["ownership"] == {
        "tenant_id": "local-tenant",
        "owner_user_id": "local-user",
    }
    assert payload["ownership_ref"]["legacy"] == payload["ownership"]
    assert payload["dedupe"]["status"] == "CREATED"
    assert payload["extraction"]["status"] == "PENDING"
    assert payload["extraction"]["markdown_available"] is False
    assert store.get_document(payload["document_id"]) == payload
    assert store.get_job(payload["extraction"]["job_id"]) == payload["ingestion_job"]
    assert Path(payload["storage"]["source_storage_path"]).read_text(encoding="utf-8") == (
        "hello from upload"
    )


def test_upload_registration_endpoint_accepts_canonical_ownership_ref(
    tmp_path: Path,
) -> None:
    client, store = build_test_client(tmp_path)

    response = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.md",
            "content_type": "text/markdown",
            "content_text": "hello from canonical owner upload",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
            "ownership_ref": OWNER_REF,
        },
        headers=auth_headers(),
    )

    assert response.status_code == 202
    payload = response.json()
    refs = store.get_content_ref(payload["document_id"])
    content_object = store.content_repository.get_content_object(
        refs["content_object_id"]
    )
    assert payload["ownership_ref"] == OWNER_REF
    assert payload["ownership"] == {"tenant_id": "tenant-a", "owner_user_id": "user-a"}
    assert content_object["ownership_ref"]["uploaded_by_subject_ref"]["id"] == "uploader-a"


def test_upload_registration_endpoint_resolves_owner_before_persisting(
    tmp_path: Path,
) -> None:
    owner_resolver = FakeOwnerResolver()
    client, store = build_test_client(
        tmp_path,
        owner_resolver=owner_resolver,
        owner_resolver_mode=UPLOAD_OWNER_RESOLVER_VERIFY,
    )

    response = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.md",
            "content_type": "text/markdown",
            "content_text": "hello from verified owner upload",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
            "ownership_ref": OWNER_REF,
        },
        headers=auth_headers(),
    )

    assert response.status_code == 202
    payload = response.json()
    assert len(owner_resolver.calls) == 1
    assert owner_resolver.calls[0]["ownership_ref"] == OWNER_REF
    assert owner_resolver.calls[0]["ensure"] is False
    assert store.get_document(payload["document_id"]) == payload


def test_upload_registration_endpoint_builds_default_resolver_when_env_mode_enabled(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    owner_resolver = FakeOwnerResolver()
    monkeypatch.setenv("NEX_CX_UPLOAD_OWNER_RESOLVER_MODE", UPLOAD_OWNER_RESOLVER_VERIFY)
    monkeypatch.setattr(
        cx_ingestion,
        "build_default_subject_registry_resolver",
        lambda *, caller_service_id: owner_resolver,
    )

    client, store = build_test_client(tmp_path)
    response = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.md",
            "content_type": "text/markdown",
            "content_text": "hello from env verified upload",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
            "ownership_ref": OWNER_REF,
        },
        headers=auth_headers(),
    )

    assert response.status_code == 202
    payload = response.json()
    assert len(owner_resolver.calls) == 1
    assert owner_resolver.calls[0]["ensure"] is False
    assert store.get_document(payload["document_id"]) == payload


def test_upload_registration_endpoint_blocks_unresolved_owner_before_persisting(
    tmp_path: Path,
) -> None:
    client, store = build_test_client(
        tmp_path,
        owner_resolver=FailingOwnerResolver(),
        owner_resolver_mode=UPLOAD_OWNER_RESOLVER_VERIFY,
    )

    response = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.md",
            "content_type": "text/markdown",
            "content_text": "this source must not be persisted",
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
            "ownership_ref": OWNER_REF,
        },
        headers=auth_headers(),
    )

    assert response.status_code == 404
    assert response.json()["error_code"] == "cx.upload_owner_unresolved"
    assert store.documents == {}
    assert store.content_repository.get_source_file_by_sha256(
        sha256_text("this source must not be persisted")
    ) is None


def test_upload_registration_endpoint_materializes_base64_source_bytes(
    tmp_path: Path,
) -> None:
    client, store = build_test_client(tmp_path)
    source_bytes = b"\x00binary endpoint upload"

    response = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.bin",
            "content_type": "application/octet-stream",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
        },
        headers=auth_headers(),
    )

    assert response.status_code == 202
    payload = response.json()
    assert payload["source_sha256"] == sha256_bytes(source_bytes)
    assert payload["upload_boundary"]["payload_source"] == "content_base64"
    assert Path(payload["storage"]["source_storage_path"]).read_bytes() == source_bytes
    assert store.get_source_bytes(payload["upload_id"]) == source_bytes
    assert store.get_source_text(payload["upload_id"]) is None


def test_source_file_materialization_endpoint_is_owner_scoped_and_redacted(
    tmp_path: Path,
) -> None:
    client, _ = build_test_client(tmp_path)
    source_bytes = b"\x00endpoint receipt bytes"
    created = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.bin",
            "content_type": "application/octet-stream",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
            "tenant_id": "tenant-a",
            "owner_user_id": "user-a",
        },
        headers=auth_headers(),
    ).json()

    response = client.get(
        f"/api/v1/documents/{created['document_id']}/source-file/materialization",
        params={"tenant_id": "tenant-a", "owner_user_id": "user-a"},
        headers=auth_headers(),
    )
    wrong_owner = client.get(
        f"/api/v1/documents/{created['document_id']}/source-file/materialization",
        params={"tenant_id": "tenant-a", "owner_user_id": "other-user"},
        headers=auth_headers(),
    )
    missing_scope = client.get(
        f"/api/v1/documents/{created['document_id']}/source-file/materialization",
        headers=auth_headers(),
    )
    unauthorized = client.get(
        f"/api/v1/documents/{created['document_id']}/source-file/materialization",
        params={"tenant_id": "tenant-a", "owner_user_id": "user-a"},
    )

    payload = response.json()
    assert response.status_code == 200
    assert payload["source_file"]["checksum_verified"] is True
    assert payload["materialization"]["source_bytes_captured"] is True
    assert payload["materialization"]["local_storage_path_included"] is False
    assert "source_storage_path" not in str(payload)
    assert str(tmp_path) not in str(payload)
    assert "endpoint receipt bytes" not in str(payload)
    assert wrong_owner.status_code == 404
    assert wrong_owner.json()["error_code"] == "cx.source_file_materialization_not_found"
    assert missing_scope.status_code == 400
    assert missing_scope.json()["error_code"] == (
        "cx.source_file_materialization_query_invalid"
    )
    assert unauthorized.status_code == 401


def test_source_file_materialization_endpoint_maps_repository_unavailable(
    tmp_path: Path,
) -> None:
    app = build_service_app(SERVICE_SPECS["nex-cx"])
    register_ingestion_routes(
        app,
        store=ContentIngestionStore(
            content_repository=FailingDocumentDetailRepository(),
        ),
        storage_config=storage_config(tmp_path),
    )
    client = TestClient(app)

    response = client.get(
        "/api/v1/documents/doc-001/source-file/materialization",
        params={"tenant_id": "tenant-a", "owner_user_id": "user-a"},
        headers=auth_headers(),
    )

    assert response.status_code == 503
    assert response.json()["error_code"] == "cx.content_repository_unavailable"
    assert response.json()["retryable"] is True


def test_payload_source_kind_labels_upload_boundary_sources() -> None:
    assert payload_source_kind(content_text="hello", source_bytes=None) == "content_text"
    assert payload_source_kind(content_text=None, source_bytes=b"hello") == "content_base64"
    assert payload_source_kind(content_text=None, source_bytes=None) == "precomputed_hash"


def test_upload_registration_endpoint_returns_existing_owner_duplicate(
    tmp_path: Path,
) -> None:
    client, store = build_test_client(tmp_path)
    body = {
        "filename": "source.md",
        "content_type": "text/markdown",
        "content_text": "same source bytes",
        "tenant_id": "tenant-a",
        "owner_user_id": "user-a",
    }
    first = client.post("/api/v1/documents/uploads", json=body, headers=auth_headers()).json()

    response = client.post(
        "/api/v1/documents/uploads",
        json={**body, "filename": "renamed.md"},
        headers=auth_headers(),
    )

    assert response.status_code == 200
    duplicate = response.json()
    assert duplicate["document_id"] == first["document_id"]
    assert duplicate["dedupe"]["status"] == "ALREADY_EXISTS"
    assert duplicate["dedupe"]["existing_document_id"] == first["document_id"]
    assert len(store.documents) == 1


def test_markdown_from_source_text_preserves_markdown() -> None:
    markdown = markdown_from_source_text(
        "# Existing\n\nBody",
        filename="source.md",
        content_type="text/markdown",
    )

    assert markdown == "# Existing\n\nBody\n"


def test_markdown_from_source_text_wraps_plain_text_with_title() -> None:
    markdown = markdown_from_source_text(
        "Plain extracted body",
        filename="source.txt",
        content_type="text/plain",
    )

    assert markdown == "# source.txt\n\nPlain extracted body\n"


def test_markdown_from_source_text_keeps_existing_trailing_newline() -> None:
    markdown = markdown_from_source_text(
        "Plain extracted body\n",
        filename="source.txt",
        content_type="text/plain",
    )

    assert markdown == "# source.txt\n\nPlain extracted body\n"


def test_markdown_from_source_text_handles_blank_source() -> None:
    assert markdown_from_source_text("  ", filename="empty.pdf", content_type="application/pdf") == (
        "# empty.pdf\n\n"
    )


def test_write_extracted_markdown_creates_parent_directory(tmp_path: Path) -> None:
    path = tmp_path / "nested" / "document.md"

    write_extracted_markdown(path, "# Written\n")

    assert path.read_text(encoding="utf-8") == "# Written\n"


def test_run_text_extraction_job_writes_markdown_and_updates_state(tmp_path: Path) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {
            "filename": "source.txt",
            "content_type": "text/plain",
            "content_text": "hello extraction",
        },
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_text="hello extraction")

    result = run_text_extraction_job(
        document["extraction"]["job_id"],
        store=store,
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    markdown_path = Path(result["extracted_markdown_path"])
    assert result["extraction_schema_version"] == "cx_text_extraction.v1"
    assert result["status"] == "SUCCEEDED"
    assert result["markdown_preview"] == "# source.txt\n\nhello extraction\n"
    assert result["extracted_markdown_sha256"] == sha256_text(markdown_path.read_text())
    assert result["extractor"] == {
        "provider": "local_mock",
        "mode": "plain_text_to_markdown",
        "version": "slice-0072",
        "source_format": "plain_text",
    }
    assert result["source_reader"] == {
        "source_reader_schema_version": "cx_source_reader.v1",
        "source": SOURCE_READER_RUNTIME_MEMORY,
        "runtime_source_bytes_used": True,
        "fallback_used": False,
        "checksum_algorithm": "sha256",
        "storage_key_included": False,
        "local_storage_path_included": False,
        "raw_source_included": False,
    }
    assert result["warnings"] == []
    assert store.get_job(result["job_id"])["status"] == "SUCCEEDED"
    assert store.get_document(result["document_id"])["extraction"]["markdown_available"] is True
    assert store.get_extraction_result(result["document_id"]) == result


def test_run_text_extraction_job_reads_materialized_source_after_memory_eviction(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {
            "filename": "source.txt",
            "content_type": "text/plain",
            "content_text": "fallback extraction",
        },
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_text="fallback extraction")
    store.source_bytes.pop(document["upload_id"])
    store.source_texts.pop(document["upload_id"])

    result = run_text_extraction_job(
        document["extraction"]["job_id"],
        store=store,
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    assert result["status"] == "SUCCEEDED"
    assert result["source_reader"]["source"] == SOURCE_READER_MATERIALIZED_LOCAL_FILE
    assert result["source_reader"]["runtime_source_bytes_used"] is False
    assert result["source_reader"]["fallback_used"] is True
    assert result["source_reader"]["storage_key_included"] is False
    assert result["source_reader"]["local_storage_path_included"] is False
    assert "fallback extraction" in Path(result["extracted_markdown_path"]).read_text(
        encoding="utf-8"
    )
    assert str(tmp_path) not in str(result["source_reader"])


def test_source_bytes_for_extraction_falls_back_without_upload_id(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {"filename": "source.txt", "content_text": "lineage fallback"},
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_text="lineage fallback")
    document_without_upload_id = {
        key: value for key, value in document.items() if key != "upload_id"
    }

    source_bytes, reader = source_bytes_for_extraction(
        document_without_upload_id,
        store=store,
        storage_config=config,
    )

    assert source_bytes == b"lineage fallback"
    assert reader["source"] == SOURCE_READER_MATERIALIZED_LOCAL_FILE
    assert reader["fallback_used"] is True


def test_run_text_extraction_job_extracts_pdf_text(tmp_path: Path) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    source_bytes = sample_pdf_bytes()
    document = build_upload_registration(
        {
            "filename": "source.pdf",
            "content_type": "application/pdf",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
        },
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_bytes=source_bytes)

    result = run_text_extraction_job(
        document["extraction"]["job_id"],
        store=store,
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    markdown = Path(result["extracted_markdown_path"]).read_text(encoding="utf-8")
    assert result["extractor"]["mode"] == "pdf_to_markdown"
    assert result["extractor"]["source_format"] == "pdf"
    assert result["warnings"] == []
    assert "Slice 0285 PDF ingestion text" in markdown
    assert "Mock extraction placeholder." not in markdown


def test_run_text_extraction_job_extracts_docx_text(tmp_path: Path) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    source_bytes = sample_docx_bytes()
    document = build_upload_registration(
        {
            "filename": "source.docx",
            "content_type": (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
        },
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_bytes=source_bytes)

    result = run_text_extraction_job(
        document["extraction"]["job_id"],
        store=store,
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    markdown = Path(result["extracted_markdown_path"]).read_text(encoding="utf-8")
    assert result["extractor"]["mode"] == "docx_to_markdown"
    assert result["extractor"]["source_format"] == "docx"
    assert result["warnings"] == []
    assert "Slice 0286 DOCX ingestion title" in markdown
    assert "Slice 0286 DOCX ingestion body" in markdown
    assert "Mock extraction placeholder." not in markdown


def test_run_text_extraction_job_uses_remaining_binary_placeholder_adapter(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    source_bytes = b"private bytes"
    document = build_upload_registration(
        {
            "filename": "source.xlsx",
            "content_type": (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ),
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
        },
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_bytes=source_bytes)

    result = run_text_extraction_job(
        document["extraction"]["job_id"],
        store=store,
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )

    markdown = Path(result["extracted_markdown_path"]).read_text(encoding="utf-8")
    assert result["extractor"]["mode"] == "binary_document_placeholder_to_markdown"
    assert result["extractor"]["source_format"] == "xlsx"
    assert result["warnings"] == ["mock_binary_extraction_placeholder:xlsx"]
    assert "private bytes" not in markdown
    assert "Mock extraction placeholder." in markdown


def test_run_text_extraction_job_reports_unsupported_adapter_source(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    source_bytes = b"\x00\x01binary"
    document = build_upload_registration(
        {
            "filename": "source.bin",
            "content_type": "application/octet-stream",
            "content_base64": base64.b64encode(source_bytes).decode("ascii"),
        },
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_bytes=source_bytes)

    with pytest.raises(IngestionError) as exc:
        run_text_extraction_job(
            document["extraction"]["job_id"],
            store=store,
            storage_config=config,
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )

    assert exc.value.status_code == 415
    assert exc.value.error_code == "cx.extractor_source_type_unsupported"


def test_run_text_extraction_job_reports_unknown_job(tmp_path: Path) -> None:
    with pytest.raises(IngestionError) as exc:
        run_text_extraction_job(
            "missing",
            store=ContentIngestionStore(),
            storage_config=storage_config(tmp_path),
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )

    assert exc.value.error_code == "cx.ingestion_job_not_found"


def test_run_text_extraction_job_reports_missing_document(tmp_path: Path) -> None:
    store = ContentIngestionStore()
    job = build_ingestion_job(
        document_id="missing-doc",
        upload_id="upload-001",
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
        created_at="2026-08-02T00:00:00Z",
    )
    store.jobs[job["job_id"]] = job

    with pytest.raises(IngestionError) as exc:
        run_text_extraction_job(
            job["job_id"],
            store=store,
            storage_config=storage_config(tmp_path),
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )

    assert exc.value.error_code == "cx.document_not_found"


def test_run_text_extraction_job_reports_missing_source_text(tmp_path: Path) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {"filename": "source.pdf", "source_sha256": "a" * 64, "size_bytes": 10},
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document)

    with pytest.raises(IngestionError) as exc:
        run_text_extraction_job(
            document["extraction"]["job_id"],
            store=store,
            storage_config=config,
            request_id=REQUEST_ID,
            trace_id=TRACE_ID,
        )

    assert exc.value.status_code == 409
    assert exc.value.error_code == "cx.source_content_unavailable"


def test_source_bytes_for_extraction_collapses_unverified_metadata_only_source(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {"filename": "source.pdf", "source_sha256": "a" * 64, "size_bytes": 10},
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document)

    with pytest.raises(IngestionError) as direct_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )
    with pytest.raises(IngestionError) as public_exc:
        source_bytes_for_extraction(document, store=store, storage_config=config)

    assert direct_exc.value.error_code == "cx.source_file_not_verified"
    assert public_exc.value.error_code == "cx.source_content_unavailable"


def test_source_bytes_for_extraction_collapses_missing_source_file_metadata(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {"filename": "source.txt", "content_text": "metadata source"},
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_text="metadata source")
    refs = store.get_content_ref(document["document_id"])
    store.source_bytes.pop(document["upload_id"])
    store.content_repository.source_files.pop(refs["source_file_id"])

    with pytest.raises(IngestionError) as exc:
        source_bytes_for_extraction(document, store=store, storage_config=config)

    assert exc.value.error_code == "cx.source_content_unavailable"


def test_read_verified_materialized_source_bytes_rejects_missing_and_corrupt_file(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {"filename": "source.txt", "content_text": "durable source"},
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_text="durable source")
    source_path = Path(document["storage"]["source_storage_path"])
    source_path.unlink()

    with pytest.raises(IngestionError) as missing_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )
    source_path.write_text("durable sourcE", encoding="utf-8")
    with pytest.raises(IngestionError) as checksum_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )
    source_path.write_text("x", encoding="utf-8")
    source_file = store.content_repository.get_source_file(
        store.get_content_ref(document["document_id"])["source_file_id"]
    )
    source_file["source_sha256"] = sha256_bytes(b"x")
    with pytest.raises(IngestionError) as size_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )

    assert missing_exc.value.error_code == "cx.source_file_missing"
    assert checksum_exc.value.error_code == "cx.source_checksum_mismatch"
    assert size_exc.value.error_code == "cx.source_size_mismatch"


def test_read_verified_materialized_source_bytes_rejects_unsafe_metadata(
    tmp_path: Path,
) -> None:
    store = ContentIngestionStore()
    config = storage_config(tmp_path)
    document = build_upload_registration(
        {"filename": "source.txt", "content_text": "safe source"},
        storage_config=config,
        request_id=REQUEST_ID,
        trace_id=TRACE_ID,
    )
    store.save_upload_registration(document, source_text="safe source")
    refs = store.get_content_ref(document["document_id"])
    source_file = store.content_repository.get_source_file(refs["source_file_id"])

    source_file["storage_backend"] = "s3"
    with pytest.raises(IngestionError) as backend_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )
    source_file["storage_backend"] = "local_filesystem"
    source_file["storage_key"] = ""
    with pytest.raises(IngestionError) as blank_key_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )
    source_file["storage_key"] = "../source.txt"
    with pytest.raises(IngestionError) as key_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )
    source_file["storage_key"] = "escape/source.txt"
    outside = tmp_path / "outside-source-root"
    outside.mkdir()
    escape = config.source_root / "escape"
    escape.symlink_to(outside, target_is_directory=True)
    with pytest.raises(IngestionError) as escape_exc:
        read_verified_materialized_source_bytes(
            document,
            store=store,
            storage_config=config,
        )
    store.source_bytes.pop(document["upload_id"])
    store.document_content_refs.pop(document["document_id"])
    with pytest.raises(IngestionError) as lineage_exc:
        source_bytes_for_extraction(document, store=store, storage_config=config)

    assert backend_exc.value.error_code == "cx.source_reader_backend_unsupported"
    assert blank_key_exc.value.error_code == "cx.source_storage_key_invalid"
    assert key_exc.value.error_code == "cx.source_storage_key_invalid"
    assert escape_exc.value.error_code == "cx.source_storage_key_invalid"
    assert lineage_exc.value.error_code == "cx.source_content_unavailable"


def test_store_save_extraction_result_requires_existing_state() -> None:
    with pytest.raises(IngestionError) as exc:
        ContentIngestionStore().save_extraction_result(
            {
                "document_id": "missing-doc",
                "job_id": "missing-job",
                "updated_at": "2026-08-02T00:00:00Z",
            }
        )

    assert exc.value.error_code == "cx.ingestion_state_not_found"


def test_upload_registration_endpoint_returns_problem_for_bad_filename(
    tmp_path: Path,
) -> None:
    client, _ = build_test_client(tmp_path)

    response = client.post(
        "/api/v1/documents/uploads",
        json={"filename": "../source.md", "content_text": "hello"},
        headers=auth_headers(),
    )

    assert response.status_code == 422
    assert response.json()["error_code"] == "cx.upload_filename_invalid"


def test_document_and_job_can_be_read_back(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)
    created = client.post(
        "/api/v1/documents/uploads",
        json={"filename": "source.md", "content_text": "hello"},
        headers=auth_headers(),
    ).json()

    document_response = client.get(
        f"/api/v1/documents/{created['document_id']}",
        params={"tenant_id": "local-tenant", "owner_user_id": "local-user"},
        headers=auth_headers(),
    )
    job_response = client.get(
        f"/api/v1/jobs/{created['extraction']['job_id']}",
        headers=auth_headers(),
    )

    assert document_response.status_code == 200
    detail = document_response.json()
    assert detail["projection_schema_version"] == "cx_document_detail_projection.v1"
    assert detail["source"] == {
        "source_kind": "memory",
        "database_env": None,
        "redacted_database_url": None,
    }
    assert detail["filters"]["document_ref"] == {
        "type": "cx.document",
        "id": created["document_id"],
    }
    assert detail["document"]["document_id"] == created["document_id"]
    assert detail["document"]["upload"]["dedupe_status"] == "CREATED"
    assert detail["document"]["extraction"] == {
        "available": True,
        "job_id": created["extraction"]["job_id"],
        "status": "PENDING",
        "markdown_available": False,
    }
    assert detail["metadata"]["not_found_and_not_authorized_collapsed"] is True
    assert "source_storage_path" not in str(detail)
    assert "hello" not in str(detail)
    assert job_response.status_code == 200
    assert job_response.json()["job_id"] == created["extraction"]["job_id"]


def test_document_detail_read_collapses_wrong_owner_and_validates_query(
    tmp_path: Path,
) -> None:
    client, _ = build_test_client(tmp_path)
    created = client.post(
        "/api/v1/documents/uploads",
        json={"filename": "source.md", "content_text": "owner scoped source"},
        headers=auth_headers(),
    ).json()

    wrong_owner = client.get(
        f"/api/v1/documents/{created['document_id']}",
        params={"tenant_id": "local-tenant", "owner_user_id": "other-user"},
        headers=auth_headers(),
    )
    invalid_owner = client.get(
        f"/api/v1/documents/{created['document_id']}",
        params={"owner_user_id": " "},
        headers=auth_headers(),
    )
    missing_owner_scope = client.get(
        f"/api/v1/documents/{created['document_id']}",
        headers=auth_headers(),
    )

    assert wrong_owner.status_code == 404
    assert wrong_owner.json()["error_code"] == "cx.document_not_found"
    assert "owner scoped source" not in str(wrong_owner.json())
    assert invalid_owner.status_code == 400
    assert invalid_owner.json()["error_code"] == "cx.document_detail_query_invalid"
    assert missing_owner_scope.status_code == 400
    assert missing_owner_scope.json()["error_code"] == "cx.document_detail_query_invalid"


def test_document_detail_read_maps_repository_unavailable(tmp_path: Path) -> None:
    app = build_service_app(SERVICE_SPECS["nex-cx"])
    register_ingestion_routes(
        app,
        store=ContentIngestionStore(
            content_repository=FailingDocumentDetailRepository(),
        ),
        storage_config=storage_config(tmp_path),
    )
    client = TestClient(app)

    response = client.get(
        "/api/v1/documents/doc-001",
        params={"tenant_id": "tenant-a", "owner_user_id": "user-a"},
        headers=auth_headers(),
    )

    assert response.status_code == 503
    assert response.json()["error_code"] == "cx.content_repository_unavailable"
    assert response.json()["retryable"] is True


def test_run_job_endpoint_materializes_markdown_and_extraction_readback(
    tmp_path: Path,
) -> None:
    client, _ = build_test_client(tmp_path)
    created = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.txt",
            "content_type": "text/plain",
            "content_text": "hello endpoint",
        },
        headers=auth_headers(),
    ).json()

    run_response = client.post(
        f"/api/v1/jobs/{created['extraction']['job_id']}/run",
        headers=auth_headers(),
    )
    read_response = client.get(
        f"/api/v1/documents/{created['document_id']}/extraction",
        headers=auth_headers(),
    )

    assert run_response.status_code == 200
    result = run_response.json()
    assert Path(result["extracted_markdown_path"]).read_text(encoding="utf-8") == (
        "# source.txt\n\nhello endpoint\n"
    )
    assert read_response.status_code == 200
    assert read_response.json()["extracted_markdown_sha256"] == result["extracted_markdown_sha256"]


def test_run_job_endpoint_requires_service_claim(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)

    response = client.post("/api/v1/jobs/missing/run")

    assert response.status_code == 401
    assert response.json()["error_code"] == "AUTHORIZATION_HEADER_MISSING"


def test_extraction_read_requires_service_claim(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)

    response = client.get("/api/v1/documents/missing/extraction")

    assert response.status_code == 401
    assert response.json()["error_code"] == "AUTHORIZATION_HEADER_MISSING"


def test_run_job_endpoint_reports_missing_source_content(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)
    created = client.post(
        "/api/v1/documents/uploads",
        json={
            "filename": "source.pdf",
            "content_type": "application/pdf",
            "source_sha256": "a" * 64,
            "size_bytes": 10,
        },
        headers=auth_headers(),
    ).json()

    response = client.post(
        f"/api/v1/jobs/{created['extraction']['job_id']}/run",
        headers=auth_headers(),
    )

    assert response.status_code == 409
    assert response.json()["error_code"] == "cx.source_content_unavailable"


def test_extraction_read_reports_not_found(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)

    response = client.get(
        "/api/v1/documents/missing/extraction",
        headers=auth_headers(),
    )

    assert response.status_code == 404
    assert response.json()["error_code"] == "cx.extraction_result_not_found"


def test_document_read_requires_service_claim(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)

    response = client.get("/api/v1/documents/missing")

    assert response.status_code == 401
    assert response.json()["error_code"] == "AUTHORIZATION_HEADER_MISSING"


def test_job_read_requires_service_claim(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)

    response = client.get("/api/v1/jobs/missing")

    assert response.status_code == 401
    assert response.json()["error_code"] == "AUTHORIZATION_HEADER_MISSING"


def test_document_and_job_reads_return_not_found(tmp_path: Path) -> None:
    client, _ = build_test_client(tmp_path)

    document_response = client.get(
        "/api/v1/documents/missing",
        params={"tenant_id": "local-tenant", "owner_user_id": "local-user"},
        headers=auth_headers(),
    )
    job_response = client.get("/api/v1/jobs/missing", headers=auth_headers())

    assert document_response.status_code == 404
    assert document_response.json()["error_code"] == "cx.document_not_found"
    assert job_response.status_code == 404
    assert job_response.json()["error_code"] == "cx.ingestion_job_not_found"
